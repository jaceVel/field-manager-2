from datetime import date, timedelta, timezone as _dt_timezone
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from .models import Job, DailyReport, SurveyFile, PersonnelName, ReportFile, Activity, ActivityCategory, ActivityType, PSSQCPreset, DiaryTemplate, JobPersonnel, JobPersonnelRole, JobVehicle, JobEquipment, Personnel, Vehicle, Equipment, Skill, JobSkillRequirement, PersonnelScheduleEntry, ScheduleStatus, SupervisorOption, SupervisorActivityTemplate, ReportPhoto, JourneyManagementPlan, JourneyPersonnel, ToolboxMeeting, ToolboxAttendee, ToolboxPhoto, JobLocation, JobMusterPoint, ToolboxTopicTemplate, JobFieldOption, VehicleAllocation, EquipmentAllocation, JobEquipmentVehicleLink, Take5Record, Take5Hazard, Take5Control, ToolboxV2Meeting, ToolboxV2Attendee, ToolboxV2Photo, JourneyV2Plan, JourneyV2Personnel, JobPhoto, TaskSafetyObservation, InfieldAudit, JSA
from .forms import JobForm, DailyReportForm, SurveyFileForm
import os
import io
import re
import json
import folium
from pyproj import Transformer
import pandas as pd

_DATUM_EPSG_BASE = {
    'mga2020':  7800,   # GDA2020 / MGA      e.g. zone 54 → EPSG:7854
    'itrf2014': 7800,   # ITRF2014 ≈ GDA2020 e.g. zone 54 → EPSG:7854
    'mga94':  28300,    # GDA94 / MGA        e.g. zone 54 → EPSG:28354
    'agd84':  20300,    # AGD84 / AMG        e.g. zone 54 → EPSG:20354
    'agd66':  20200,    # AGD66 / AMG        e.g. zone 54 → EPSG:20254
    'wgs84':  32700,    # WGS84 / UTM South  e.g. zone 54 → EPSG:32754
}

def _datum_epsg(datum, zone):
    base = _DATUM_EPSG_BASE.get(datum, 28300)
    return f'EPSG:{base + int(zone)}'


def _resolve_epsg(job, datum_key, zone):
    """Return job.survey_epsg if set, otherwise derive from file headers."""
    if job and job.survey_epsg:
        return job.survey_epsg
    return _datum_epsg(datum_key, zone)


def _load_epsg_list():
    """Load Australian EPSG code list from the data file."""
    data_path = os.path.join(os.path.dirname(__file__), 'data', 'australian_epsg.json')
    try:
        with open(data_path, 'r') as f:
            return json.load(f)
    except Exception:
        return []


def _get_job_tz(job):
    """Return a timezone object for the job's deployment file timezone setting."""
    from zoneinfo import ZoneInfo
    if job.timezone == 'custom' and job.utc_offset_custom is not None:
        return _dt_timezone(timedelta(hours=float(job.utc_offset_custom)))
    try:
        return ZoneInfo(job.timezone)
    except Exception:
        return ZoneInfo('UTC')


def _read_csv(path, **kw):
    """Drop-in for pd.read_csv that always strips column whitespace and coerces
    any 'line' column (case-insensitive) to integer."""
    df = pd.read_csv(path, **kw)
    df.columns = df.columns.str.strip()
    col_lower = {c.lower(): c for c in df.columns}
    if 'line' in col_lower:
        col = col_lower['line']
        df[col] = pd.to_numeric(df[col], errors='coerce').astype('Int64')
    return df


def _read_survey_csv(path):
    """Read a survey CSV and normalise column names to Line/Point/X/Y/Z."""
    df = _read_csv(path)
    col_lower = {c.lower(): c for c in df.columns}
    for expected, variants in [
        ('Line',  ['line']),
        ('Point', ['point', 'station']),
        ('X',     ['x', 'easting']),
        ('Y',     ['y', 'northing']),
        ('Z',     ['z', 'elevation', 'elev']),
    ]:
        if expected not in df.columns:
            for v in variants:
                if v in col_lower:
                    df = df.rename(columns={col_lower[v]: expected})
                    break
    return df


def _parse_sps21(file_field):
    """Parse an SPS 2.1 FileField (rps_file or sps_file) into (df, datum_key, zone).
    Returns (empty_df, 'mga2020', 55) on failure.
    df columns: Line, Point, X (easting), Y (northing), Z (elevation).
    """
    _DATUM_MAP = {
        'gda2020': 'mga2020', 'mga2020': 'mga2020',
        'itrf2014': 'itrf2014',
        'gda94': 'mga94', 'mga94': 'mga94',
        'agd84': 'agd84', 'agd66': 'agd66',
        'wgs84': 'wgs84',
    }
    datum_key = 'mga2020'
    zone = 55
    records = []
    try:
        file_field.open('rb')
        for raw in file_field:
            line = raw.decode('utf-8', errors='ignore').rstrip('\r\n')
            if not line:
                continue
            if line.startswith('H12'):
                text = line.lower()
                for k in _DATUM_MAP:
                    if k in text:
                        datum_key = _DATUM_MAP[k]
                        break
                m = re.search(r'zone\s+(\d+)', text)
                if m:
                    zone = int(m.group(1))
            elif not line.startswith('H') and len(line) >= 65:
                try:
                    records.append({
                        'Line':  float(line[1:11]),
                        'Point': float(line[11:21]),
                        'X':     float(line[46:55]),
                        'Y':     float(line[55:65]),
                        'Z':     float(line[65:71]) if line[65:71].strip() else None,
                    })
                except (ValueError, IndexError):
                    pass
        file_field.close()
    except Exception:
        pass
    df = pd.DataFrame(records) if records else pd.DataFrame(columns=['Line', 'Point', 'X', 'Y', 'Z'])
    return df, datum_key, zone


def _norm_combined(val):
    """Convert a whole-number float string like '1000100.0' to '1000100' before char-splitting."""
    try:
        f = float(val)
        if f == int(f):
            return str(int(f))
    except (ValueError, TypeError):
        pass
    return val


def _get_line_stn(row, line_col, stn_col, combined, split_chars):
    """Extract (line, stn) strings from a DataFrame row, handling combined columns."""
    if combined and stn_col and split_chars:
        val  = _norm_combined(str(row.get(stn_col, '')))
        line = val[:-split_chars] if len(val) > split_chars else val
        stn  = val[-split_chars:]  if len(val) > split_chars else ''
    else:
        line = str(row.get(line_col, '')) if line_col else ''
        stn  = str(row.get(stn_col,  '')) if stn_col  else ''
    return line.strip(), stn.strip()


from playwright.sync_api import sync_playwright


def _map_to_png(html_path, png_path, width=900, height=500, center=None, zoom=None, extra_wait=0):
    """Screenshot a Folium HTML map file to PNG using Playwright."""
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={'width': width, 'height': height})
            page.goto(f'file:///{html_path.replace(chr(92), "/")}')
            page.wait_for_timeout(1500 + extra_wait)  # wait for tiles to load
            if center and zoom is not None:
                lat, lng = center
                page.evaluate(f"""() => {{
                    for (var k in window) {{
                        try {{
                            if (window[k] && typeof window[k].setView === 'function' && typeof window[k].getZoom === 'function') {{
                                window[k].setView([{lat}, {lng}], {zoom}, {{animate: false}});
                                break;
                            }}
                        }} catch(e) {{}}
                    }}
                }}""")
                page.wait_for_timeout(800 + extra_wait)  # wait for tiles at new view
            page.screenshot(path=png_path, full_page=False)
            browser.close()
        return True
    except Exception:
        return False


def _add_map_legend(m, items):
    """Inject a fixed-position color legend into a Folium map. items = [(color, label), ...]"""
    rows = ''
    for color, label in items:
        border = ';border:1.5px solid #888' if color == '#ffffff' else ''
        rows += (f'<div style="display:flex;align-items:center;gap:5px;margin-bottom:2px;">'
                 f'<span style="display:inline-block;width:11px;height:11px;border-radius:50%;'
                 f'background:{color}{border};flex-shrink:0;"></span>'
                 f'<span>{label}</span></div>')
    html = (f'<div style="position:fixed;bottom:12px;left:12px;z-index:9999;'
            f'background:rgba(255,255,255,0.88);padding:7px 10px;border-radius:5px;'
            f'font-family:Arial,sans-serif;font-size:11px;line-height:1.5;'
            f'box-shadow:0 1px 4px rgba(0,0,0,0.3);">{rows}</div>')
    m.get_root().html.add_child(folium.Element(html))


def _parse_map_views(request):
    """Parse map center/zoom query params from request. Returns dict or None."""
    views = {}
    for key in ('progress', 'deployment', 'pss'):
        try:
            lat  = float(request.GET[f'{key}_lat'])
            lng  = float(request.GET[f'{key}_lng'])
            zoom = int(request.GET[f'{key}_zoom'])
            views[key] = (lat, lng, zoom)
        except (KeyError, ValueError):
            pass
    views['deployment_mode'] = request.GET.get('deployment_mode', 'alldays')
    return views or None


# --- Jobs ---

def home(request):
    return render(request, 'reports/home.html', {
        'job_count': Job.objects.count(),
    })


def qhse_take5_stats(request):
    from django.db.models import Count
    import json as _json
    from collections import Counter

    by_job = list(
        Take5Record.objects
        .values('job__pk', 'job__job_number', 'job__project_name')
        .annotate(records=Count('id'), hazards=Count('hazards'))
        .order_by('-records')
    )
    by_person = list(
        Take5Record.objects
        .values('submitted_by')
        .annotate(cnt=Count('id'))
        .order_by('-cnt')[:10]
    )
    total_records = Take5Record.objects.count()
    total_hazards = Take5Hazard.objects.count()
    recent = Take5Record.objects.select_related('job').prefetch_related('hazards')[:20]

    # Full task → hazard(s) → control(s) tree
    all_records = (
        Take5Record.objects
        .select_related('job')
        .prefetch_related('hazards__controls')
        .order_by('-submitted_at')
    )

    # Top hazards and controls for chart
    hazard_counts = Counter(
        h.hazard.strip() for h in Take5Hazard.objects.all() if h.hazard.strip()
    )
    control_counts = Counter(
        c.control.strip()
        for h in Take5Hazard.objects.prefetch_related('controls').all()
        for c in h.controls.all()
        if c.control.strip()
    )
    top_hazards = hazard_counts.most_common(12)
    top_controls = control_counts.most_common(12)

    return render(request, 'reports/qhse_take5_stats.html', {
        'by_job': by_job,
        'by_person': by_person,
        'total_records': total_records,
        'total_hazards': total_hazards,
        'jobs_count': len(by_job),
        'recent': recent,
        'hazard_labels': _json.dumps([h[0] for h in top_hazards]),
        'hazard_data': _json.dumps([h[1] for h in top_hazards]),
        'control_labels': _json.dumps([c[0] for c in top_controls]),
        'control_data': _json.dumps([c[1] for c in top_controls]),
        'all_records': all_records,
    })


def qhse_overview(request):
    from django.db.models import Count
    take5_by_job = list(
        Take5Record.objects
        .values('job__job_number', 'job__project_name', 'job__pk')
        .annotate(cnt=Count('id'))
        .order_by('-cnt')
    )
    return render(request, 'reports/qhse_overview.html', {
        'take5_total': Take5Record.objects.count(),
        'take5_jobs': len(take5_by_job),
        'take5_by_job': take5_by_job,
    })


def job_list(request):
    jobs = Job.objects.all()
    return render(request, 'reports/job_list.html', {'jobs': jobs})


def personnel_export_csv(request):
    import csv as _csv
    personnel = Personnel.objects.prefetch_related('skills').all()
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="personnel.csv"'
    writer = _csv.writer(response)
    writer.writerow(['Name', 'Phone', 'Email', 'Skills'])
    for p in personnel:
        writer.writerow([
            p.name,
            p.phone,
            p.email,
            ', '.join(s.name for s in p.skills.all()),
        ])
    return response


def personnel_list(request):
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'add_person':
            name = request.POST.get('name', '').strip()
            if name:
                Personnel.objects.create(
                    name=name,
                    phone=request.POST.get('phone', '').strip(),
                    email=request.POST.get('email', '').strip(),
                )
        elif action == 'edit_person':
            person = Personnel.objects.filter(pk=request.POST.get('person_pk')).first()
            if person:
                name = request.POST.get('name', '').strip()
                if name:
                    person.name = name
                person.phone = request.POST.get('phone', '').strip()
                person.email = request.POST.get('email', '').strip()
                person.save()
        elif action == 'delete_person':
            Personnel.objects.filter(pk=request.POST.get('person_pk')).delete()
        elif action == 'import_csv':
            import csv, io
            f = request.FILES.get('csv_file')
            if f:
                text = f.read().decode('utf-8-sig')  # strips BOM
                reader = csv.reader(io.StringIO(text))
                existing_names = set(Personnel.objects.values_list('name', flat=True))
                added, skipped = [], []
                for row in reader:
                    if not row or not row[0].strip():
                        continue
                    name  = row[0].strip()
                    phone = row[1].strip() if len(row) > 1 else ''
                    email = row[2].strip() if len(row) > 2 else ''
                    if name in existing_names:
                        skipped.append(name)
                    else:
                        Personnel.objects.create(name=name, phone=phone, email=email)
                        existing_names.add(name)
                        added.append(name)
                request.session['import_result'] = {'added': len(added), 'skipped': len(skipped)}
            return redirect('personnel_list')
        elif action == 'add_skill':
            person = Personnel.objects.filter(pk=request.POST.get('person_pk')).first()
            skill = Skill.objects.filter(pk=request.POST.get('skill_pk')).first()
            if person and skill:
                person.skills.add(skill)
                if request.headers.get('X-Fetch') == '1':
                    return JsonResponse({'ok': True, 'skill_pk': skill.pk, 'skill_name': skill.name})
                return redirect(f"{reverse('personnel_list')}?open_skill={person.pk}")
        elif action == 'remove_skill':
            person = Personnel.objects.filter(pk=request.POST.get('person_pk')).first()
            skill = Skill.objects.filter(pk=request.POST.get('skill_pk')).first()
            if person and skill:
                person.skills.remove(skill)
                if request.headers.get('X-Fetch') == '1':
                    return JsonResponse({'ok': True})
        elif action == 'create_skill':
            name = request.POST.get('skill_name', '').strip()
            if name:
                skill, created = Skill.objects.get_or_create(name=name)
                if request.headers.get('X-Fetch') == '1':
                    return JsonResponse({'ok': True, 'skill_pk': skill.pk, 'skill_name': skill.name, 'created': created})
        elif action == 'delete_skill':
            pk = request.POST.get('skill_pk')
            Skill.objects.filter(pk=pk).delete()
            if request.headers.get('X-Fetch') == '1':
                return JsonResponse({'ok': True})
        return redirect('personnel_list')
    import_result = request.session.pop('import_result', None)
    return render(request, 'reports/personnel_list.html', {
        'personnel': Personnel.objects.prefetch_related('skills').all(),
        'all_skills': Skill.objects.all(),
        'open_skill': request.GET.get('open_skill', ''),
        'import_result': import_result,
    })


def vehicle_list(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            Vehicle.objects.create(
                name=name,
                vehicle_type=request.POST.get('vehicle_type', '').strip(),
                notes=request.POST.get('notes', '').strip(),
            )
        else:
            Vehicle.objects.filter(pk=request.POST.get('delete_pk')).delete()
        return redirect('vehicle_list')
    return render(request, 'reports/vehicle_list.html', {
        'vehicles': Vehicle.objects.all(),
        'lv_vehicles': Vehicle.objects.filter(vehicle_type='Light Vehicle').order_by('name'),
        'mv_vehicles': Vehicle.objects.filter(vehicle_type='Medium Vehicle').order_by('name'),
        'hv_vehicles': Vehicle.objects.filter(vehicle_type='Heavy Vehicle').order_by('name'),
        'tt_vehicles': Vehicle.objects.filter(vehicle_type='Tilt Tray').order_by('name'),
        'vib_vehicles': Vehicle.objects.filter(vehicle_type='Vibrator').order_by('name'),
        'atv_vehicles': Vehicle.objects.filter(vehicle_type='ATV').order_by('name'),
    })


def vehicle_parse_xls(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    f = request.FILES.get('file')
    if not f:
        return JsonResponse({'error': 'No file'}, status=400)
    try:
        import io as _io
        df = pd.read_excel(_io.BytesIO(f.read()))
        def clean(v):
            return '' if (v != v or str(v) in ('nan', 'NaT', 'None')) else str(v).strip()

        def import_type(df, type_name):
            subset = df[df['AssetTypeName'].astype(str).str.strip() == type_name].copy()
            count = 0
            for _, row in subset.iterrows():
                asset_num   = clean(row.get('AssetNumber', ''))
                description = clean(row.get('AssetDescription', ''))
                category    = clean(row.get('ParentAsset', ''))
                make        = clean(row.get('ManufacturerName', ''))
                model_name  = clean(row.get('ModelNumberName', ''))
                rego        = clean(row.get('RegistrationNumber', ''))
                is_active   = bool(row.get('IsActive', True))
                name = asset_num or description
                if not name:
                    continue
                _, created = Vehicle.objects.get_or_create(
                    name=name,
                    defaults={
                        'vehicle_type': type_name,
                        'description': description,
                        'category': category,
                        'make': make,
                        'model_name': model_name,
                        'rego': rego,
                        'is_active': is_active,
                        'notes': ' '.join(filter(None, [rego, make, model_name])),
                    },
                )
                if created:
                    count += 1
            return count

        _fields = ['pk', 'name', 'description', 'category', 'make', 'model_name', 'rego', 'is_active']
        added_lv = import_type(df, 'Light Vehicle')
        added_mv = import_type(df, 'Medium Vehicle')
        added_hv = import_type(df, 'Heavy Vehicle')
        added_tt  = import_type(df, 'Tilt Tray')
        added_vib = import_type(df, 'Vibrator')

        # ATVs are identified by ParentAsset, not AssetTypeName
        atv_subset = df[df['ParentAsset'].astype(str).str.strip() == 'ATV'].copy()
        added_atv = 0
        for _, row in atv_subset.iterrows():
            asset_num   = clean(row.get('AssetNumber', ''))
            description = clean(row.get('AssetDescription', ''))
            category    = clean(row.get('ParentAsset', ''))
            make        = clean(row.get('ManufacturerName', ''))
            model_name  = clean(row.get('ModelNumberName', ''))
            rego        = clean(row.get('RegistrationNumber', ''))
            is_active   = bool(row.get('IsActive', True))
            name = asset_num or description
            if not name:
                continue
            _, created = Vehicle.objects.get_or_create(
                name=name,
                defaults={
                    'vehicle_type': 'ATV',
                    'description': description,
                    'category': category,
                    'make': make,
                    'model_name': model_name,
                    'rego': rego,
                    'is_active': is_active,
                    'notes': ' '.join(filter(None, [rego, make, model_name])),
                },
            )
            if created:
                added_atv += 1
        lv_vehicles  = list(Vehicle.objects.filter(vehicle_type='Light Vehicle').order_by('name').values(*_fields))
        mv_vehicles  = list(Vehicle.objects.filter(vehicle_type='Medium Vehicle').order_by('name').values(*_fields))
        hv_vehicles  = list(Vehicle.objects.filter(vehicle_type='Heavy Vehicle').order_by('name').values(*_fields))
        tt_vehicles  = list(Vehicle.objects.filter(vehicle_type='Tilt Tray').order_by('name').values(*_fields))
        vib_vehicles = list(Vehicle.objects.filter(vehicle_type='Vibrator').order_by('name').values(*_fields))
        atv_vehicles = list(Vehicle.objects.filter(vehicle_type='ATV').order_by('name').values(*_fields))
        return JsonResponse({
            'ok': True,
            'added': added_lv + added_mv + added_hv + added_tt + added_vib + added_atv,
            'lv_vehicles': lv_vehicles,
            'mv_vehicles': mv_vehicles,
            'hv_vehicles': hv_vehicles,
            'tt_vehicles': tt_vehicles,
            'vib_vehicles': vib_vehicles,
            'atv_vehicles': atv_vehicles,
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


def vehicle_delete_json(request, pk):
    if request.method == 'POST':
        Vehicle.objects.filter(pk=pk).delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'POST required'}, status=400)


def vehicle_add_json(request):
    if request.method == 'POST':
        import json as _json
        data = _json.loads(request.body)
        name = data.get('name', '').strip()
        if not name:
            return JsonResponse({'error': 'Name required'}, status=400)
        v, _ = Vehicle.objects.get_or_create(
            name=name,
            defaults={
                'vehicle_type': data.get('vehicle_type', 'Light Vehicle'),
                'description': data.get('description', ''),
                'category': data.get('category', ''),
                'make': data.get('make', ''),
                'model_name': data.get('model_name', ''),
                'rego': data.get('rego', ''),
                'is_active': data.get('is_active', True),
                'notes': data.get('notes', ''),
            },
        )
        return JsonResponse({'ok': True, 'pk': v.pk, 'name': v.name, 'description': v.description,
                             'category': v.category, 'make': v.make, 'model_name': v.model_name,
                             'rego': v.rego, 'is_active': v.is_active})
    return JsonResponse({'error': 'POST required'}, status=400)


def vehicle_edit_json(request, pk):
    if request.method == 'POST':
        import json as _json
        data = _json.loads(request.body)
        try:
            v = Vehicle.objects.get(pk=pk)
        except Vehicle.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        name = data.get('name', '').strip()
        if not name:
            return JsonResponse({'error': 'Name required'}, status=400)
        v.name = name
        v.description = data.get('description', v.description)
        v.category = data.get('category', v.category)
        v.make = data.get('make', v.make)
        v.model_name = data.get('model_name', v.model_name)
        v.rego = data.get('rego', v.rego)
        v.is_active = data.get('is_active', v.is_active)
        v.save()
        return JsonResponse({'ok': True, 'pk': v.pk, 'name': v.name, 'description': v.description,
                             'category': v.category, 'make': v.make, 'model_name': v.model_name,
                             'rego': v.rego, 'is_active': v.is_active})
    return JsonResponse({'error': 'POST required'}, status=400)


def equipment_list(request):
    all_eq = list(Equipment.objects.all())
    known = ('Trailer', 'Genset', 'Starlink')
    tabs = [
        {'tab': 'trailer',  'label': 'Trailers',  'type': 'Trailer',  'items': [e for e in all_eq if e.equipment_type == 'Trailer']},
        {'tab': 'genset',   'label': 'Gensets',   'type': 'Genset',   'items': [e for e in all_eq if e.equipment_type == 'Genset']},
        {'tab': 'starlink', 'label': 'Starlinks', 'type': 'Starlink', 'items': [e for e in all_eq if e.equipment_type == 'Starlink']},
        {'tab': 'other',    'label': 'Other',     'type': '',         'items': [e for e in all_eq if e.equipment_type not in known]},
    ]
    return render(request, 'reports/equipment_list.html', {'tabs': tabs})


def equipment_add_json(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    name = data.get('name', '').strip()
    if not name:
        return JsonResponse({'ok': False, 'error': 'name required'})
    eq = Equipment.objects.create(
        name=name,
        serial_number=data.get('serial_number', '').strip(),
        equipment_type=data.get('equipment_type', '').strip(),
        notes=data.get('notes', '').strip(),
        is_active=data.get('is_active', True),
    )
    return JsonResponse({'ok': True, 'pk': eq.pk, 'name': eq.name,
                         'serial_number': eq.serial_number,
                         'equipment_type': eq.equipment_type,
                         'notes': eq.notes, 'is_active': eq.is_active})


def equipment_delete_json(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    try:
        eq = Equipment.objects.get(pk=pk)
    except Equipment.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Not found'}, status=404)
    eq.delete()
    return JsonResponse({'ok': True})


def equipment_toggle_active_json(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    try:
        eq = Equipment.objects.get(pk=pk)
    except Equipment.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Not found'}, status=404)
    eq.is_active = not eq.is_active
    eq.save(update_fields=['is_active'])
    return JsonResponse({'ok': True, 'is_active': eq.is_active})


def _register_job_field_options(job):
    for field in ('recording_system', 'source_type'):
        val = (getattr(job, field) or '').strip()
        if val:
            JobFieldOption.objects.get_or_create(field=field, value=val)


def _job_form_context(form, title):
    return {
        'form': form,
        'title': title,
        'options_recording_system': list(JobFieldOption.objects.filter(field='recording_system').values_list('pk', 'value')),
        'options_source_type': list(JobFieldOption.objects.filter(field='source_type').values_list('pk', 'value')),
    }


def job_create(request):
    if request.method == 'POST':
        form = JobForm(request.POST)
        if form.is_valid():
            job = form.save()
            _register_job_field_options(job)
            return redirect('job_detail', pk=job.pk)
    else:
        form = JobForm()
    return render(request, 'reports/job_form.html', _job_form_context(form, 'New Job'))


def job_detail(request, pk):
    job = get_object_or_404(Job, pk=pk)
    return render(request, 'reports/job_detail.html', {'job': job})


def job_allocations(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    TABS = [('mob', 'Mob'), ('job', 'Job'), ('demob', 'Demob'), ('other', 'Other')]
    job_vehicles = list(job.vehicles.all())
    fleet_by_name = {v.name: v for v in Vehicle.objects.all()}
    for jv in job_vehicles:
        fv = fleet_by_name.get(jv.name)
        jv.rego = fv.rego if fv else ''
    personnel_names = [p.name for p in job.personnel.filter(is_separator=False).order_by('name')]
    job_equipment_all = list(job.equipment.all())

    tabs_data = []
    for tab_key, tab_label in TABS:
        v_allocs = list(VehicleAllocation.objects.filter(job_vehicle__job=job, tab=tab_key))
        eq_links = list(JobEquipmentVehicleLink.objects.filter(
            job_vehicle__job=job, tab=tab_key).select_related('job_equipment'))

        alloc_by_veh = {}
        for a in v_allocs:
            alloc_by_veh.setdefault(a.job_vehicle_id, []).append(
                {'pk': a.pk, 'person_name': a.person_name})

        eq_link_by_veh = {}
        for lnk in eq_links:
            eq_link_by_veh.setdefault(lnk.job_vehicle_id, []).append({
                'link_pk': lnk.pk, 'je_pk': lnk.job_equipment_id,
                'name': lnk.job_equipment.name,
                'equipment_type': lnk.job_equipment.equipment_type,
            })

        allocated_names = {a.person_name for a in v_allocs}
        assigned_eq_pks = {lnk.job_equipment_id for lnk in eq_links}

        tabs_data.append({
            'tab': tab_key,
            'label': tab_label,
            'vehicles': [{'jv': jv,
                          'allocs': alloc_by_veh.get(jv.pk, []),
                          'eq_links': eq_link_by_veh.get(jv.pk, [])}
                         for jv in job_vehicles],
            'unallocated': [n for n in personnel_names if n not in allocated_names],
            'unassigned_equipment': [e for e in job_equipment_all if e.pk not in assigned_eq_pks],
        })

    # Equipment Allocation tab data
    from collections import defaultdict
    ea_by_person = defaultdict(list)
    assigned_je_pks = set()
    for alloc in EquipmentAllocation.objects.filter(job_equipment__job=job).select_related('job_equipment'):
        ea_by_person[alloc.person_name].append({
            'alloc_pk': alloc.pk,
            'je_pk': alloc.job_equipment_id,
            'name': alloc.job_equipment.name,
            'equipment_type': alloc.job_equipment.equipment_type,
        })
        assigned_je_pks.add(alloc.job_equipment_id)
    eq_alloc_people = [
        {'name': p.name, 'eq_allocs': ea_by_person.get(p.name, [])}
        for p in job.personnel.filter(is_separator=False).order_by('name')
    ]
    unassigned_eq_ea = [e for e in job_equipment_all if e.pk not in assigned_je_pks]

    return render(request, 'reports/job_allocations.html', {
        'job': job,
        'tabs_data': tabs_data,
        'eq_alloc_people': eq_alloc_people,
        'unassigned_eq_ea': unassigned_eq_ea,
    })


def allocation_add(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    jv = get_object_or_404(JobVehicle, pk=data.get('jv_pk'), job__pk=job_pk)
    name = data.get('person_name', '').strip()
    if not name:
        return JsonResponse({'error': 'name required'}, status=400)
    tab = data.get('tab', 'job')
    alloc, _ = VehicleAllocation.objects.get_or_create(job_vehicle=jv, person_name=name, tab=tab)
    return JsonResponse({'ok': True, 'alloc_pk': alloc.pk})


def allocation_remove(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    alloc = get_object_or_404(VehicleAllocation, pk=pk)
    alloc.delete()
    return JsonResponse({'ok': True})


def equipment_allocation_add(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    je = get_object_or_404(JobEquipment, pk=data.get('je_pk'), job__pk=job_pk)
    person_name = data.get('person_name', '').strip()
    if not person_name:
        return JsonResponse({'error': 'person_name required'}, status=400)
    alloc, _ = EquipmentAllocation.objects.get_or_create(job_equipment=je, person_name=person_name)
    return JsonResponse({'ok': True, 'alloc_pk': alloc.pk})


def equipment_allocation_remove(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    alloc = get_object_or_404(EquipmentAllocation, pk=pk)
    alloc.delete()
    return JsonResponse({'ok': True})


def job_survey_files(request, pk):
    job = get_object_or_404(Job, pk=pk)

    if request.method == 'POST' and request.POST.get('form_type') == 'estimated_counts':
        rx = request.POST.get('estimated_rx_count', '').strip()
        sx = request.POST.get('estimated_sx_count', '').strip()
        job.estimated_rx_count = int(rx) if rx.isdigit() else None
        job.estimated_sx_count = int(sx) if sx.isdigit() else None
        job.save()
        return redirect('job_survey_files', pk=pk)

    survey_files = job.survey_files.all()

    rx_count = sx_count = None
    rx_lines = sx_lines = None
    rx_map_data = sx_map_data = None

    def _load_file(file_type):
        qs = survey_files.filter(file_type=file_type)
        final = qs.filter(is_final=True).first()
        latest = qs.order_by('-uploaded_at').first()
        return final, latest or final

    rx_final, rx_sf = _load_file('rx')
    sx_final, sx_sf = _load_file('sx')

    if rx_sf:
        df = _read_csv(rx_sf.file.path)
        df.columns = df.columns.str.strip()
        if rx_final:
            rx_count = len(df)
        rx_lines = {'columns': list(df.columns), 'rows': df.head(5).astype(str).values.tolist()}
        rx_map_data = {
            'pk': rx_sf.pk,
            'datum': rx_sf.datum,
            'zone': rx_sf.zone,
            'columns': list(df.columns),
            'rows': df.head(500).astype(str).values.tolist(),
            'total': len(df),
        }

    if sx_sf:
        df = _read_csv(sx_sf.file.path)
        df.columns = df.columns.str.strip()
        if sx_final:
            sx_count = len(df)
        sx_lines = {'columns': list(df.columns), 'rows': df.head(5).astype(str).values.tolist()}
        sx_map_data = {
            'pk': sx_sf.pk,
            'datum': sx_sf.datum,
            'zone': sx_sf.zone,
            'columns': list(df.columns),
            'rows': df.head(500).astype(str).values.tolist(),
            'total': len(df),
        }

    return render(request, 'reports/job_survey_files.html', {
        'job': job,
        'survey_files': survey_files,
        'zone_choices': SurveyFile.ZONE_CHOICES,
        'rx_count': rx_count,
        'sx_count': sx_count,
        'rx_lines': rx_lines,
        'sx_lines': sx_lines,
        'rx_lines_json': json.dumps(rx_lines) if rx_lines else 'null',
        'sx_lines_json': json.dumps(sx_lines) if sx_lines else 'null',
        'rx_map_json': json.dumps(rx_map_data) if rx_map_data else 'null',
        'sx_map_json': json.dumps(sx_map_data) if sx_map_data else 'null',
    })


def job_reports(request, pk):
    job = get_object_or_404(Job, pk=pk)
    reports = job.reports.filter(report_type='production').order_by('-date')
    return render(request, 'reports/job_reports.html', {
        'job': job,
        'reports': reports,
    })


def build_sx_from_rx(request, job_pk):
    """AJAX: run the Sx-from-Rx algorithm and return preview + CSV."""
    from math import sqrt, atan2, sin, cos, pi as _pi
    import csv as _csv

    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    job = get_object_or_404(Job, pk=job_pk)
    try:
        offset = float(request.POST.get('offset', 0))
        side   = request.POST.get('side', 'Left')   # Left / Right / Both
        datum  = request.POST.get('datum', '')
        zone   = request.POST.get('zone', '')
    except ValueError:
        return JsonResponse({'error': 'Invalid parameters'}, status=400)

    rx_qs = job.survey_files.filter(file_type='rx')
    rx_sf = rx_qs.filter(is_final=True).first() or rx_qs.order_by('-uploaded_at').first()
    if not rx_sf:
        return JsonResponse({'error': 'No Rx file found for this job.'}, status=400)

    df = _read_csv(rx_sf.file.path)
    df.columns = df.columns.str.strip()
    cols = list(df.columns)
    # Expect at least 5 columns: Line, Station, Easting, Northing, Elevation
    if len(cols) < 5:
        return JsonResponse({'error': f'Rx file must have at least 5 columns (has {len(cols)}).'}, status=400)

    def _offset_point(start, end, dist, s):
        mid_x = (start[0] + end[0]) / 2
        mid_y = (start[1] + end[1]) / 2
        dx, dy = end[0] - start[0], end[1] - start[1]
        angle = atan2(dy, dx)
        perp = angle + (_pi / 2) if s == 'left' else angle - (_pi / 2)
        return mid_x + dist * cos(perp), mid_y + dist * sin(perp)

    rows = df.values.tolist()
    out_rows = []

    for i in range(1, len(rows)):
        pt1, pt2 = rows[i - 1], rows[i]
        try:
            line1, line2 = str(pt1[0]), str(pt2[0])
            if line1 != line2:
                continue
            stn1  = float(pt1[1])
            E1, N1 = float(pt1[2]), float(pt1[3])
            E2, N2 = float(pt2[2]), float(pt2[3])
            elev1, elev2 = float(pt1[4]), float(pt2[4])
        except (ValueError, TypeError, IndexError):
            continue

        shot_stn  = round(stn1 + 0.5, 1)
        shot_elev = round((elev1 + elev2) / 2.0, 2)
        left  = _offset_point((E1, N1), (E2, N2), offset, 'left')
        right = _offset_point((E1, N1), (E2, N2), offset, 'right')

        if side in ('Left', 'Both'):
            out_rows.append([line1, shot_stn, round(left[0], 2), round(left[1], 2), shot_elev])
        if side in ('Right', 'Both'):
            right_line = line1 + '0' if side == 'Both' else line1
            out_rows.append([right_line, shot_stn, round(right[0], 2), round(right[1], 2), shot_elev])

    # Build CSV string
    import io as _io
    buf = _io.StringIO()
    w = _csv.writer(buf)
    w.writerow(cols[:5])
    w.writerows(out_rows)
    csv_str = buf.getvalue()

    return JsonResponse({
        'csv': csv_str,
        'columns': cols[:5],
        'preview': out_rows[:5],
        'total': len(out_rows),
        'datum': datum or rx_sf.datum,
        'zone': zone or rx_sf.zone,
    })


def sx_save_generated(request, job_pk):
    """Save a generated Sx CSV (from build_sx_from_rx) as a SurveyFile."""
    import io as _io
    from django.core.files.base import ContentFile

    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    job = get_object_or_404(Job, pk=job_pk)
    csv_str = request.POST.get('csv', '')
    datum   = request.POST.get('datum', 'mga2020')
    zone    = request.POST.get('zone', '52')

    if not csv_str:
        return JsonResponse({'error': 'No CSV data'}, status=400)

    from .models import SurveyFile as _SF
    sf = _SF(job=job, file_type='sx', datum=datum, zone=zone)
    filename = f'sx_built_from_rx_{job.job_number}.csv'
    sf.file.save(filename, ContentFile(csv_str.encode('utf-8')), save=True)

    return JsonResponse({'ok': True, 'redirect': reverse('job_survey_files', kwargs={'pk': job_pk})})


def survey_file_upload(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    datum = ''
    zone = ''
    if request.method == 'POST':
        form = SurveyFileForm(request.POST, request.FILES)
        datum = request.POST.get('datum', '')
        zone = request.POST.get('zone', '')
        if form.is_valid():
            sf = form.save(commit=False)
            sf.job = job
            sf.save()
    from urllib.parse import urlencode
    qs = urlencode({'datum': datum, 'zone': zone})
    return redirect(f"{reverse('job_survey_files', kwargs={'pk': job_pk})}?{qs}")


def survey_file_delete(request, pk):
    sf = get_object_or_404(SurveyFile, pk=pk)
    job_pk = sf.job.pk
    sf.file.delete()
    sf.delete()
    return redirect('job_survey_files', pk=job_pk)


def survey_file_toggle_final(request, pk):
    sf = get_object_or_404(SurveyFile, pk=pk)
    sf.is_final = not sf.is_final
    sf.save()
    return redirect('job_survey_files', pk=sf.job.pk)


def generate_xps(request, job_pk):
    if request.method != 'POST':
        return HttpResponse('Method not allowed', status=405)
    job = get_object_or_404(Job, pk=job_pk)

    rx_qs = job.survey_files.filter(file_type='rx')
    sx_qs = job.survey_files.filter(file_type='sx')
    rx_sf = rx_qs.filter(is_final=True).first() or rx_qs.order_by('-uploaded_at').first()
    sx_sf = sx_qs.filter(is_final=True).first() or sx_qs.order_by('-uploaded_at').first()
    if not rx_sf or not sx_sf:
        return HttpResponse('Both Rx and Sx files are required to generate XPS.', status=400)

    station_increment = int(request.POST.get('station_increment') or 1)
    live_spread       = int(request.POST.get('live_spread') or 0)

    rx_df = _read_csv(rx_sf.file.path)
    rx_df.columns = rx_df.columns.str.strip()
    sx_df = _read_csv(sx_sf.file.path)
    sx_df.columns = sx_df.columns.str.strip()

    rx_combined   = request.POST.get('rx_combined')   == '1'
    rx_split_chars = int(request.POST.get('rx_split_chars') or 0)
    rx_line_col   = request.POST.get('rx_line_col', '')
    rx_stn_col    = request.POST.get('rx_stn_col',  '')
    sx_combined   = request.POST.get('sx_combined')   == '1'
    sx_split_chars = int(request.POST.get('sx_split_chars') or 0)
    sx_line_col   = request.POST.get('sx_line_col', '')
    sx_stn_col    = request.POST.get('sx_stn_col',  '')

    def _to_int(val):
        return int(float(str(val).strip()))

    # Build receiver existence dict
    rx_dict = {}
    for _, row in rx_df.iterrows():
        line, stn = _get_line_stn(row, rx_line_col, rx_stn_col, rx_combined, rx_split_chars)
        try:
            rx_dict[(_to_int(line), _to_int(stn))] = True
        except (ValueError, TypeError):
            pass

    if not rx_dict:
        return HttpResponse(
            f'No receivers loaded. Check Rx column mapping.\n'
            f'Rx columns: {list(rx_df.columns)}\n'
            f'rx_line_col={rx_line_col!r} rx_stn_col={rx_stn_col!r} combined={rx_combined}',
            status=400, content_type='text/plain'
        )

    # Build source list
    sources = []
    for _, row in sx_df.iterrows():
        line, stn = _get_line_stn(row, sx_line_col, sx_stn_col, sx_combined, sx_split_chars)
        sources.append((line, stn))

    out_lines = ['H00 SPS format version num.     SPS 2.1;']
    ffid   = 1
    offset = live_spread // 2

    for src_line, src_stn in sources:
        try:
            shot_point   = float(str(src_stn).strip())
            src_line_int = _to_int(src_line)
            src_stn_int  = shot_point if shot_point != int(shot_point) else int(shot_point)
        except (ValueError, TypeError):
            ffid += 1
            continue

        low_stn  = int(shot_point - (offset - 1))
        high_stn = int(shot_point + offset)

        stations = sorted(
            tstn for tstn in range(low_stn, high_stn + 1)
            if (src_line_int, tstn) in rx_dict
        )
        if not stations:
            ffid += 1
            continue

        # Group into continuous runs
        sets, current = [], [stations[0]]
        for j in range(1, len(stations)):
            if stations[j] == stations[j - 1] + station_increment:
                current.append(stations[j])
            else:
                sets.append(current)
                current = [stations[j]]
        sets.append(current)

        from_ch = 1
        to_ch   = 0
        for s in sets:
            to_ch += len(s)
            out_lines.append(
                'X'
                + str(ffid).rjust(14)
                + '10'
                + str(src_line_int).rjust(10)
                + str(src_stn_int).rjust(10)
                + '1'
                + str(from_ch).rjust(5)
                + str(to_ch).rjust(5)
                + '1'
                + str(src_line_int).rjust(10)
                + str(min(s)).rjust(10)
                + str(max(s)).rjust(10)
                + '1'
            )
            from_ch += len(s)
        ffid += 1

    if len(out_lines) == 1:
        return HttpResponse(
            f'XPS generated 0 records.\n'
            f'Rx dict size: {len(rx_dict)}\n'
            f'Sources: {len(sources)}\n'
            f'live_spread={live_spread} station_increment={station_increment}\n'
            f'Sample rx keys: {list(rx_dict.keys())[:5]}\n'
            f'Sample sources: {sources[:5]}',
            status=400, content_type='text/plain'
        )

    content  = '\n'.join(out_lines) + '\n'
    date_str = date.today().strftime('%Y%m%d')
    job_slug = re.sub(r'[^\w]+', '_', f"{job.job_number}_{job.project_name}").strip('_')
    filename = f'{job_slug}_{date_str}.xps'
    response = HttpResponse(content, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def xps_preview_json(request, job_pk):
    """Return the XPS source→receiver mapping as JSON for map interaction."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    job = get_object_or_404(Job, pk=job_pk)

    rx_qs = job.survey_files.filter(file_type='rx')
    sx_qs = job.survey_files.filter(file_type='sx')
    rx_sf = rx_qs.filter(is_final=True).first() or rx_qs.order_by('-uploaded_at').first()
    sx_sf = sx_qs.filter(is_final=True).first() or sx_qs.order_by('-uploaded_at').first()
    if not rx_sf or not sx_sf:
        return JsonResponse({'error': 'Both Rx and Sx files required'}, status=400)

    station_increment = int(request.POST.get('station_increment') or 1)
    live_spread       = int(request.POST.get('live_spread') or 0)

    rx_df = _read_csv(rx_sf.file.path)
    sx_df = _read_csv(sx_sf.file.path)

    rx_combined    = request.POST.get('rx_combined')    == '1'
    rx_split_chars = int(request.POST.get('rx_split_chars') or 0)
    rx_line_col    = request.POST.get('rx_line_col', '')
    rx_stn_col     = request.POST.get('rx_stn_col',  '')
    sx_combined    = request.POST.get('sx_combined')    == '1'
    sx_split_chars = int(request.POST.get('sx_split_chars') or 0)
    sx_line_col    = request.POST.get('sx_line_col', '')
    sx_stn_col     = request.POST.get('sx_stn_col',  '')

    def _to_int(val):
        return int(float(str(val).strip()))

    rx_dict = {}
    for _, row in rx_df.iterrows():
        line, stn = _get_line_stn(row, rx_line_col, rx_stn_col, rx_combined, rx_split_chars)
        try:
            rx_dict[(_to_int(line), _to_int(stn))] = True
        except (ValueError, TypeError):
            pass

    offset  = live_spread // 2
    mapping = {}

    for _, row in sx_df.iterrows():
        line, stn = _get_line_stn(row, sx_line_col, sx_stn_col, sx_combined, sx_split_chars)
        try:
            shot_point   = float(str(stn).strip())
            src_line_int = _to_int(line)
            src_stn_val  = shot_point if shot_point != int(shot_point) else int(shot_point)
        except (ValueError, TypeError):
            continue

        low_stn  = int(shot_point - (offset - 1))
        high_stn = int(shot_point + offset)

        stations = sorted(
            t for t in range(low_stn, high_stn + 1)
            if (src_line_int, t) in rx_dict
        )
        if not stations:
            continue

        sets, current = [], [stations[0]]
        for j in range(1, len(stations)):
            if stations[j] == stations[j - 1] + station_increment:
                current.append(stations[j])
            else:
                sets.append(current)
                current = [stations[j]]
        sets.append(current)

        key = f'{src_line_int}:{src_stn_val}'
        mapping[key] = [[src_line_int, min(s), max(s)] for s in sets]

    return JsonResponse({'mapping': mapping})


def generate_sps_rps(request, job_pk):
    if request.method != 'POST':
        return HttpResponse('Method not allowed', status=405)
    job = get_object_or_404(Job, pk=job_pk)
    file_type = request.POST.get('file_type', 'rx')
    qs = job.survey_files.filter(file_type=file_type)
    sf = qs.filter(is_final=True).first() or qs.order_by('-uploaded_at').first()
    if not sf:
        return HttpResponse('No survey file found', status=400)

    df = _read_csv(sf.file.path)
    df.columns = df.columns.str.strip()

    combined    = request.POST.get('combined') == '1'
    split_chars  = int(request.POST.get('split_chars') or 0)
    line_col     = request.POST.get('line_col', '')
    stn_col      = request.POST.get('stn_col', '')
    e_col        = request.POST.get('e_col', '')
    n_col        = request.POST.get('n_col', '')
    elev_col     = request.POST.get('elev_col', '')

    prefix = 'R' if file_type == 'rx' else 'S'
    ext    = 'rps' if file_type == 'rx' else 'sps'

    out_lines = ['H00 SPS format version num:     SPS 2.1;']
    for _, row in df.iterrows():
        if combined and stn_col and split_chars:
            # stn_col holds the combined Line+Station value; split from right
            raw      = str(row.get(stn_col, ''))
            try:
                f = float(raw)
                if f == int(f):
                    raw = str(int(f))
            except (ValueError, TypeError):
                pass
            val      = raw
            line_val = val[:-split_chars] if len(val) > split_chars else val
            stn_val  = val[-split_chars:]  if len(val) > split_chars else ''
        else:
            line_val = str(row.get(line_col, '')) if line_col else ''
            stn_val  = str(row.get(stn_col,  '')) if stn_col  else ''

        if file_type == 'rx':
            try:
                stn_val = str(int(float(stn_val)))
            except (ValueError, TypeError):
                pass

        try:
            e_val = float(row.get(e_col, 0)) if e_col else 0.0
        except (ValueError, TypeError):
            e_val = 0.0
        try:
            n_val = float(row.get(n_col, 0)) if n_col else 0.0
        except (ValueError, TypeError):
            n_val = 0.0
        try:
            elev_val = float(row.get(elev_col, 0)) if elev_col else 0.0
        except (ValueError, TypeError):
            elev_val = 0.0

        line_str = line_val.rjust(10)
        stn_str  = stn_val.rjust(10)
        e_str    = f'{e_val:.2f}'.rjust(31)
        n_str    = f'{n_val:.2f}'.strip().rjust(10)
        elev_str = f'{elev_val:.2f}'.rjust(6)
        out_lines.append(f'{prefix}{line_str}{stn_str}  1{e_str}{n_str}{elev_str}123123456')

    content  = '\n'.join(out_lines) + '\n'
    date_str = date.today().strftime('%Y%m%d')
    job_slug = re.sub(r'[^\w]+', '_', f"{job.job_number}_{job.project_name}").strip('_')
    filename = f'{job_slug}_{date_str}.{ext}'
    response = HttpResponse(content, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def survey_file_map_data(request, pk):
    sf = get_object_or_404(SurveyFile, pk=pk)
    df = _read_csv(sf.file.path)
    df.columns = df.columns.str.strip()
    return JsonResponse({
        'pk': sf.pk,
        'datum': sf.datum,
        'zone': sf.zone,
        'columns': list(df.columns),
        'rows': df.astype(str).values.tolist(),
        'total': len(df),
    })


def survey_file_inspect(request, pk):
    sf = get_object_or_404(SurveyFile, pk=pk)
    try:
        df_full = _read_survey_csv(sf.file.path)
        sample = df_full.head(20).astype(str)
        rows = sample.values.tolist()
        result = {'columns': list(sample.columns), 'rows': rows}
        # Per-line stats if Line and Point columns exist
        if 'Line' in df_full.columns and 'Point' in df_full.columns:
            df_full['Point'] = pd.to_numeric(df_full['Point'], errors='coerce')
            line_stats = []
            for line, grp in df_full.groupby('Line', sort=True):
                pts = grp['Point'].dropna()
                line_stats.append({
                    'line': str(line),
                    'count': int(len(grp)),
                    'min': str(int(pts.min())) if len(pts) else '—',
                    'max': str(int(pts.max())) if len(pts) else '—',
                })
            result['line_stats'] = line_stats
        return JsonResponse(result)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


def job_edit(request, pk):
    job = get_object_or_404(Job, pk=pk)
    if request.method == 'POST':
        form = JobForm(request.POST, instance=job)
        if form.is_valid():
            job = form.save()
            _register_job_field_options(job)
            return redirect('job_detail', pk=job.pk)
    else:
        form = JobForm(instance=job)
    return render(request, 'reports/job_form.html', _job_form_context(form, f'Edit Job {job.job_number}'))


def job_delete(request, pk):
    job = get_object_or_404(Job, pk=pk)
    if request.method == 'POST':
        job.delete()
        return redirect('job_list')
    return render(request, 'reports/job_confirm_delete.html', {'job': job})


def job_field_option_delete(request, pk):
    if request.method == 'POST':
        JobFieldOption.objects.filter(pk=pk).delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'POST required'}, status=400)


# --- Reports ---

def report_delete(request, pk):
    report = get_object_or_404(DailyReport, pk=pk)
    job_pk = report.job.pk
    report_type = report.report_type
    if request.method == 'POST':
        report.delete()
        return redirect('report_create', job_pk=job_pk, report_type=report_type)
    return redirect('report_create', job_pk=job_pk, report_type=report_type)


def report_set_status(request, pk):
    report = get_object_or_404(DailyReport, pk=pk)
    if request.method == 'POST':
        new_status = request.POST.get('status', '').strip()
        if new_status in ('draft', 'submitted', 'approved'):
            report.status = new_status
            report.save(update_fields=['status'])
    return redirect(request.META.get('HTTP_REFERER', '/'))


def report_create(request, job_pk, report_type):
    job = get_object_or_404(Job, pk=job_pk)
    duplicate_error = None
    if request.method == 'POST':
        form = DailyReportForm(request.POST)
        if form.is_valid():
            date = form.cleaned_data['date']
            if job.reports.filter(report_type=report_type, date=date).exists():
                duplicate_error = f"A {report_type} report for {date.strftime('%d %b %Y')} already exists."
            else:
                report = form.save(commit=False)
                report.job = job
                report.report_type = report_type
                report.save()
                return redirect('report_detail', pk=report.pk)
    else:
        form = DailyReportForm()
    existing = job.reports.filter(report_type=report_type).order_by('-date')
    return render(request, 'reports/report_form.html', {
        'form': form,
        'job': job,
        'report_type': report_type.capitalize(),
        'existing_reports': existing,
        'duplicate_error': duplicate_error,
    })


def report_detail(request, pk):
    report = get_object_or_404(DailyReport, pk=pk)
    observers = PersonnelName.objects.filter(role='observer')
    operators = PersonnelName.objects.filter(role='operator')

    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        if form_type == 'diary':
            report.diary = request.POST.get('diary', '').replace('\r\n', '\n').replace('\r', '\n')
            report.save()
            return redirect(f'/reports/{pk}/?tab=diary')

        elif form_type == 'pss_settings':
            fields = [
                'pss_force_avg_green', 'pss_force_avg_amber',
                'pss_force_max_green', 'pss_force_max_amber',
                'pss_phase_avg_green', 'pss_phase_avg_amber',
                'pss_phase_max_green', 'pss_phase_max_amber',
                'pss_thd_avg_green',  'pss_thd_avg_amber',
                'pss_thd_max_green',  'pss_thd_max_amber',
            ]
            for f in fields:
                val = request.POST.get(f, '').strip()
                try:
                    setattr(report.job, f, float(val))
                except ValueError:
                    pass
            report.job.save()
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'map_colors':
            for field in (
                'progress_color_today', 'progress_color_prev', 'progress_color_planned',
                'deployment_color_today', 'deployment_color_prev', 'deployment_color_planned',
            ):
                val = request.POST.get(field, '').strip()
                if val.startswith('#') and len(val) == 7:
                    setattr(report.job, field, val)
            report.job.show_map_overlay = 'show_map_overlay' in request.POST
            report.job.save()
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'finish_settings':
            try:
                report.job.finish_days_per_week  = max(1, min(7, int(request.POST.get('finish_days_per_week', 7))))
                report.job.finish_rolling_window = max(1, min(60, int(request.POST.get('finish_rolling_window', 7))))
            except (ValueError, TypeError):
                pass
            report.job.finish_show_linear        = 'finish_show_linear'        in request.POST
            report.job.finish_show_calendar      = 'finish_show_calendar'      in request.POST
            report.job.finish_show_rolling       = 'finish_show_rolling'       in request.POST
            report.job.finish_include_in_report  = 'finish_include_in_report'  in request.POST
            report.job.save()
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'export_filename':
            tpl = request.POST.get('export_filename_template', '').strip()
            if report.report_type == 'supervisor':
                if tpl:
                    report.job.supervisor_filename_template = tpl
                report.job.save(update_fields=['supervisor_filename_template'])
            else:
                if tpl:
                    report.job.export_filename_template = tpl
                report.job.export_save_to_disk = 'export_save_to_disk' in request.POST
                report.job.export_save_path = request.POST.get('export_save_path', '').strip()
                report.job.save(update_fields=['export_filename_template', 'export_save_to_disk', 'export_save_path'])
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'signature_settings':
            report.job.include_signatures = 'include_signatures' in request.POST
            report.job.save(update_fields=['include_signatures'])
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'shot_chart_toggle':
            report.include_shot_chart = request.POST.get('include_shot_chart') == '1'
            report.save(update_fields=['include_shot_chart'])
            from django.http import JsonResponse
            return JsonResponse({'ok': True})

        elif form_type == 'avg_days':
            checked_ids = set(request.POST.getlist('include_in_avg'))
            for r in report.job.reports.all():
                r.include_in_avg = str(r.pk) in checked_ids
                r.save(update_fields=['include_in_avg'])
            return redirect(f'/reports/{pk}/?tab=settings')

        elif form_type == 'activities':
            report.activities.all().delete()
            categories = request.POST.getlist('category')
            types = request.POST.getlist('activity_type')
            if report.report_type == 'supervisor':
                contractors = request.POST.getlist('details')
                names = request.POST.getlist('notes')
                job_titles = request.POST.getlist('job_title')
                hours_list = request.POST.getlist('hours')
                for i, contractor in enumerate(contractors):
                    Activity.objects.create(
                        report=report,
                        start_time='00:00',
                        end_time='00:00',
                        category=categories[i] if i < len(categories) else '',
                        activity_type=types[i] if i < len(types) else 'Work',
                        details=contractor,
                        notes=names[i] if i < len(names) else '',
                        job_title=job_titles[i] if i < len(job_titles) else '',
                        hours=hours_list[i] if i < len(hours_list) else '',
                        order=i,
                    )
            else:
                starts = request.POST.getlist('start_time')
                ends = request.POST.getlist('end_time')
                details = request.POST.getlist('details')
                for i, start in enumerate(starts):
                    if start and i < len(ends) and ends[i]:
                        Activity.objects.create(
                            report=report,
                            start_time=start,
                            end_time=ends[i],
                            category=categories[i] if i < len(categories) else '',
                            activity_type=types[i] if i < len(types) else 'Work',
                            details=details[i] if i < len(details) else '',
                            order=i,
                        )
            return redirect(f'/reports/{pk}/?tab=activities')

        else:
            observer_names = [n.strip() for n in request.POST.getlist('observers') if n.strip()]
            for name in observer_names:
                PersonnelName.objects.get_or_create(name=name, role='observer')
            operator_names = [n.strip() for n in request.POST.getlist('operators') if n.strip()]
            for name in operator_names:
                PersonnelName.objects.get_or_create(name=name, role='operator')
            report.observers = '\n'.join(observer_names)
            report.operators = '\n'.join(operator_names)
            report.save()
            return redirect('report_detail', pk=report.pk)

    report_files = report.files.all()
    default_tab = 'files'
    active_tab = request.GET.get('tab', default_tab)
    map_bounds = request.GET.get('map_bounds', 'today')
    pss_presets = PSSQCPreset.objects.all()
    diary_templates = DiaryTemplate.objects.all()
    progress_map_html = None
    categories = ActivityCategory.objects.all()
    default_category = categories.filter(name__icontains='toolbox').first() or \
                       categories.filter(name__icontains='safety').first()
    debug_info = {}

    # Parse ObserverLog for shot stats
    shot_stats = None
    obs_file = report_files.filter(file_type='obslog').first()
    if obs_file:
        try:
            df = _read_csv(obs_file.file.path, skiprows=2)
            df.columns = df.columns.str.strip()
            # Filter out future dates if the log contains a date column
            if 'Local Date' in df.columns:
                try:
                    _obs_dates = pd.to_datetime(df['Local Date'].str.strip(), errors='coerce').dt.date
                    df = df[_obs_dates <= report.date]
                except Exception:
                    pass
            total = len(df)
            void_mask = df['Status'].str.strip() == 'Void'
            void = int(void_mask.sum())
            production = total - void
            prod_df = df[~void_mask]
            status_counts = df['Status'].str.strip().value_counts().to_dict()
            line_stats = []
            for line, grp in prod_df.groupby('Line', sort=True):
                line_stats.append({
                    'line': int(line) if line == int(line) else line,
                    'start': grp['Station'].min(),
                    'end': grp['Station'].max(),
                    'shots': len(grp),
                })
            # 15-minute interval shot counts
            shot_intervals = []
            if len(prod_df):
                try:
                    times = pd.to_datetime(prod_df['Local Time'].str.strip(), format='%H:%M:%S.%f', errors='coerce')
                    times = times.dropna()
                    if len(times):
                        first = times.min()
                        last = times.max()
                        # build bins from first to last in 15-min steps
                        slot = first.floor('15min')
                        end_slot = last.floor('15min')
                        bins = []
                        while slot <= end_slot:
                            bins.append(slot)
                            slot += pd.Timedelta(minutes=15)
                        counts = pd.cut(times, bins=[b for b in bins] + [bins[-1] + pd.Timedelta(minutes=15)],
                                        labels=[b.strftime('%H:%M') for b in bins], right=False)
                        vc = counts.value_counts().reindex(
                            [b.strftime('%H:%M') for b in bins], fill_value=0
                        )
                        shot_intervals = [
                            {'interval': lbl, 'shots': int(n)}
                            for lbl, n in vc.items()
                        ]
                except Exception:
                    pass

            shot_stats = {
                'total': total,
                'production': production,
                'void': void,
                'status_counts': status_counts,
                'first_shot_time': prod_df['Local Time'].iloc[0] if len(prod_df) else None,
                'last_shot_time': prod_df['Local Time'].iloc[-1] if len(prod_df) else None,
                'first_file': int(prod_df['File#'].iloc[0]) if len(prod_df) else None,
                'last_file': int(prod_df['File#'].iloc[-1]) if len(prod_df) else None,
                'line_stats': line_stats,
                'shot_intervals': shot_intervals,
            }
        except Exception as e:
            debug_info['obslog_error'] = str(e)

    # Parse rx_deployment
    rx_stats = None
    rx_df_cache = None
    rx_file = report_files.filter(file_type='rx_deployment').first()
    if rx_file:
        try:
            rx_df = _read_csv(rx_file.file.path)
            rx_df.columns = rx_df.columns.str.strip()
            # Filter out future dates relative to this report.
            # Use tz-aware Timestamp comparisons to avoid None/NaT issues
            # when comparing object-dtype .dt.date series with Python date objects.
            try:
                _tz_rx = _get_job_tz(report.job)
                _rx_dt = (
                    pd.to_datetime(rx_df['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                    .dt.tz_convert(_tz_rx)
                )
                # Midnight of the day AFTER report.date in local tz — keep everything strictly before it
                _cutoff = pd.Timestamp(report.date + timedelta(days=1), tz=_tz_rx)
                rx_df = rx_df[_rx_dt < _cutoff]
            except Exception:
                pass
            rx_df_cache = rx_df
            if rx_df.empty:
                debug_info['rx_deployment_error'] = (
                    f'No rows remain after date filter (report date: {report.date}). '
                    f'Check that the Deployment File Timezone matches the UTC offset used in the file.'
                )
            else:
                lines = sorted(rx_df['Line'].dropna().unique())
                line_stats = []
                for line, grp in rx_df.groupby('Line', sort=True):
                    line_stats.append({
                        'line': int(line) if line == int(line) else line,
                        'start': grp['Point'].min(),
                        'end': grp['Point'].max(),
                        'nodes': len(grp),
                    })
                rx_stats = {
                    'total': len(rx_df),
                    'lines': [int(l) if l == int(l) else l for l in lines],
                    'line_count': len(lines),
                    'point_min': int(rx_df['Point'].min()),
                    'point_max': int(rx_df['Point'].max()),
                    'status_counts': rx_df['Data_Status'].str.strip().value_counts().to_dict(),
                    'line_stats': line_stats,
                }
        except Exception as e:
            debug_info['rx_deployment_error'] = str(e)
    rx_deployed = rx_stats['total'] if rx_stats else None

    # Nodes deployed on this report's date specifically
    nodes_today = None
    nodes_today_line_stats = []
    if rx_df_cache is not None:
        try:
            _tz_nd = _get_job_tz(report.job)
            _dep_dt = (
                pd.to_datetime(rx_df_cache['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                .dt.tz_convert(_tz_nd)
            )
            _day_start = pd.Timestamp(report.date, tz=_tz_nd)
            _day_end = pd.Timestamp(report.date + timedelta(days=1), tz=_tz_nd)
            _today_df = rx_df_cache[(_dep_dt >= _day_start) & (_dep_dt < _day_end)]
            nodes_today = len(_today_df)
            for _line, _grp in _today_df.groupby('Line', sort=True):
                nodes_today_line_stats.append({
                    'line': int(_line) if _line == int(_line) else _line,
                    'start': _grp['Point'].min(),
                    'end': _grp['Point'].max(),
                    'nodes': len(_grp),
                })
        except Exception:
            pass

    # Job-level shot stats (reports up to and including this report's date)
    job_total_shots = 0
    avg_shots_total = 0
    production_days = 0
    all_job_reports = report.job.reports.filter(date__lte=report.date, report_type='production')
    # Per-report shot counts for settings UI
    report_shot_rows = []
    for r in all_job_reports.order_by('date'):
        obs = r.files.filter(file_type='obslog').first()
        day_prod = None
        if obs:
            try:
                d = _read_csv(obs.file.path, skiprows=2)
                d.columns = d.columns.str.strip()
                day_prod = int((d['Status'].str.strip() != 'Void').sum())
                if day_prod > 0:
                    job_total_shots += day_prod
                    if r.include_in_avg:
                        avg_shots_total += day_prod
                        production_days += 1
            except Exception:
                pass
        report_shot_rows.append({'report': r, 'shots': day_prod})

    # Planned shots from SPS file count (if final) or estimate
    planned_shots = None
    if report.job.sps_is_final and report.job.sps_count:
        planned_shots = report.job.sps_count
    if planned_shots is None and report.job.estimated_sx_count:
        planned_shots = report.job.estimated_sx_count

    shots_remaining = (planned_shots - job_total_shots) if planned_shots is not None else None
    daily_avg_shots = round(avg_shots_total / production_days) if production_days else None

    # Job-level node stats — derived from the cumulative rx_deployment file
    # (each file contains ALL nodes deployed so far, not just today's)
    job_total_nodes = rx_stats['total'] if rx_stats else 0
    avg_nodes_total = 0
    deployment_days = 0
    report_node_rows = []
    if rx_df_cache is not None and not rx_df_cache.empty:
        try:
            _tz2 = _get_job_tz(report.job)
            _dep_dt2 = (
                pd.to_datetime(rx_df_cache['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                .dt.tz_convert(_tz2)
            )
            _dep_local_dates = _dep_dt2.dt.date
            _nodes_by_date = rx_df_cache.groupby(_dep_local_dates).size().to_dict()
            # Map dates to reports for include_in_avg
            _report_by_date = {r.date: r for r in all_job_reports}
            for dep_date in sorted(_nodes_by_date):
                day_nodes = _nodes_by_date[dep_date]
                r = _report_by_date.get(dep_date)
                if r and r.include_in_avg:
                    avg_nodes_total += day_nodes
                    deployment_days += 1
                report_node_rows.append({
                    'report': r,
                    'nodes': day_nodes,
                    'date': dep_date,
                })
        except Exception:
            pass
    daily_avg_nodes = round(avg_nodes_total / deployment_days) if deployment_days else None

    planned_nodes = None
    if report.job.rps_is_final and report.job.rps_count:
        planned_nodes = report.job.rps_count
    if planned_nodes is None and report.job.estimated_rx_count:
        planned_nodes = report.job.estimated_rx_count

    nodes_remaining = (planned_nodes - job_total_nodes) if planned_nodes is not None else None

    # Finish date estimation
    finish_estimates = []
    if shots_remaining is not None and shots_remaining > 0:
        import math
        job = report.job
        dpw = max(1, min(7, job.finish_days_per_week))  # days per week (1-7)

        def _prod_to_calendar(prod_days):
            """Convert production days to calendar days using days-per-week setting."""
            return math.ceil(prod_days * 7 / dpw)

        def _finish_from(base_date, cal_days):
            return base_date + timedelta(days=cal_days)

        last_report_date = all_job_reports.order_by('date').last().date if all_job_reports.exists() else report.date

        # Method 1: Linear (daily avg)
        if job.finish_show_linear and daily_avg_shots:
            prod_days = math.ceil(shots_remaining / daily_avg_shots)
            cal_days = _prod_to_calendar(prod_days)
            finish_estimates.append({
                'label': 'Daily Avg',
                'date': _finish_from(last_report_date, cal_days),
                'detail': f'{prod_days} prod. days / {cal_days} cal. days',
            })

        # Method 2: Calendar rate (total shots ÷ elapsed calendar days)
        if job.finish_show_calendar and job_total_shots:
            first_report = all_job_reports.order_by('date').first()
            if first_report:
                elapsed = (last_report_date - first_report.date).days + 1
                if elapsed > 0:
                    cal_rate = job_total_shots / elapsed  # shots per calendar day
                    cal_days = math.ceil(shots_remaining / cal_rate)
                    finish_estimates.append({
                        'label': 'Calendar Rate',
                        'date': _finish_from(last_report_date, cal_days),
                        'detail': f'{cal_rate:.1f} shots/cal. day → {cal_days} cal. days',
                    })

        # Method 3: Rolling average (last N included days)
        if job.finish_show_rolling:
            window = max(1, job.finish_rolling_window)
            included = [r for r in report_shot_rows if r['report'].include_in_avg and r['shots'] and r['shots'] > 0]
            recent = included[-window:]
            if recent:
                rolling_avg = sum(r['shots'] for r in recent) / len(recent)
                prod_days = math.ceil(shots_remaining / rolling_avg)
                cal_days = _prod_to_calendar(prod_days)
                finish_estimates.append({
                    'label': f'Rolling {window}-day Avg',
                    'date': _finish_from(last_report_date, cal_days),
                    'detail': f'{rolling_avg:.0f} shots/day → {prod_days} prod. days / {cal_days} cal. days',
                })
    elif shots_remaining is not None and shots_remaining <= 0:
        finish_estimates = [{'label': 'Complete', 'date': None, 'detail': 'All planned shots fired'}]

    # Finish date estimation — Node Deployment
    node_finish_estimates = []
    if nodes_remaining is not None and nodes_remaining > 0:
        import math as _math
        job = report.job
        dpw = max(1, min(7, job.finish_days_per_week))

        def _prod_to_cal(prod_days):
            return _math.ceil(prod_days * 7 / dpw)

        def _finish_node(base_date, cal_days):
            return base_date + timedelta(days=cal_days)

        last_node_date = all_job_reports.order_by('date').last().date if all_job_reports.exists() else report.date

        if job.finish_show_linear and daily_avg_nodes:
            prod_days = _math.ceil(nodes_remaining / daily_avg_nodes)
            cal_days = _prod_to_cal(prod_days)
            node_finish_estimates.append({
                'label': 'Daily Avg',
                'date': _finish_node(last_node_date, cal_days),
                'detail': f'{prod_days} prod. days / {cal_days} cal. days',
            })

        if job.finish_show_calendar and job_total_nodes and report_node_rows:
            first_dep_date = report_node_rows[0]['date']
            elapsed = (last_node_date - first_dep_date).days + 1
            if elapsed > 0:
                cal_rate = job_total_nodes / elapsed
                cal_days = _math.ceil(nodes_remaining / cal_rate)
                node_finish_estimates.append({
                    'label': 'Calendar Rate',
                    'date': _finish_node(last_node_date, cal_days),
                    'detail': f'{cal_rate:.1f} nodes/cal. day → {cal_days} cal. days',
                })

        if job.finish_show_rolling:
            window = max(1, job.finish_rolling_window)
            included = [r for r in report_node_rows if r['nodes'] and r['nodes'] > 0 and r['report'] and r['report'].include_in_avg]
            recent = included[-window:]
            if recent:
                rolling_avg = sum(r['nodes'] for r in recent) / len(recent)
                prod_days = _math.ceil(nodes_remaining / rolling_avg)
                cal_days = _prod_to_cal(prod_days)
                node_finish_estimates.append({
                    'label': f'Rolling {window}-day Avg',
                    'date': _finish_node(last_node_date, cal_days),
                    'detail': f'{rolling_avg:.0f} nodes/day → {prod_days} prod. days / {cal_days} cal. days',
                })
    elif nodes_remaining is not None and nodes_remaining <= 0:
        node_finish_estimates = [{'label': 'Complete', 'date': None, 'detail': 'All planned nodes deployed'}]

    # Deployment map for Report tab
    deployment_map_html = None
    deployment_map_stats = {}
    if rx_df_cache is not None:
        try:
            rps_ff = report.job.rps_file if report.job.rps_file else None
            if rps_ff:
                rx_sur_df, datum_key, zone = _parse_sps21(rps_ff)
                epsg_in = _resolve_epsg(report.job, datum_key, zone)
                rx_sur_df['_line'] = rx_sur_df['Line'].round(1)
                rx_sur_df['_point'] = rx_sur_df['Point'].round(1)

                dep_df = rx_df_cache.copy()
                dep_df['_line'] = dep_df['Line'].round(1)
                dep_df['_point'] = dep_df['Point'].round(1)

                # Split into today vs previous using Deployment_Time (UTC → local date)
                job_tz = _get_job_tz(report.job)
                try:
                    _dep_dt_map = (
                        pd.to_datetime(dep_df['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                        .dt.tz_convert(job_tz)
                    )
                    dep_df['_local_date'] = _dep_dt_map.dt.date
                    _day_start_map = pd.Timestamp(report.date, tz=job_tz)
                    _day_end_map = pd.Timestamp(report.date + timedelta(days=1), tz=job_tz)
                    dep_today = dep_df[(_dep_dt_map >= _day_start_map) & (_dep_dt_map < _day_end_map)]
                    dep_prev  = dep_df[~((_dep_dt_map >= _day_start_map) & (_dep_dt_map < _day_end_map))]
                    deployment_map_stats['tz'] = report.job.timezone
                    deployment_map_stats['report_date'] = str(report.date)
                    deployment_map_stats['local_dates'] = sorted({str(d) for d in dep_df['_local_date'].unique()})
                    deployment_map_stats['today_count'] = len(dep_today)
                    deployment_map_stats['prev_count'] = len(dep_prev)
                    deployment_map_stats['split_error'] = None
                except Exception as e:
                    dep_today = dep_df
                    dep_prev  = dep_df.iloc[0:0]
                    deployment_map_stats['split_error'] = str(e)

                merged_today = dep_today.merge(rx_sur_df[['_line', '_point', 'X', 'Y']], on=['_line', '_point'], how='inner')
                merged_prev  = dep_prev.merge(rx_sur_df[['_line', '_point', 'X', 'Y']], on=['_line', '_point'], how='inner')
                merged = pd.concat([merged_today, merged_prev])
                deployment_map_stats['merged_today'] = len(merged_today)
                deployment_map_stats['merged_prev'] = len(merged_prev)

                t = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)

                # All Rx survey points as planned
                sur_lons, sur_lats = t.transform(rx_sur_df['X'].values, rx_sur_df['Y'].values)
                rx_sur_df['lat'] = sur_lats
                rx_sur_df['lon'] = sur_lons

                buf = 500
                if map_bounds == 'today':
                    _dep_bound_df = merged_today if not merged_today.empty else (merged if not merged.empty else rx_sur_df)
                else:
                    _dep_bound_df = rx_sur_df
                sw_lon_d, sw_lat_d = t.transform(_dep_bound_df['X'].min() - buf, _dep_bound_df['Y'].min() - buf)
                ne_lon_d, ne_lat_d = t.transform(_dep_bound_df['X'].max() + buf, _dep_bound_df['Y'].max() + buf)

                dm = folium.Map(tiles=None, zoom_control=True, width='100%', height=500,
                                scrollWheelZoom=False, dragging=True, doubleClickZoom=True)
                folium.TileLayer(
                    tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
                    attr='Esri', name='Satellite',
                ).add_to(dm)
                if report.job.show_map_overlay:
                    folium.TileLayer(
                        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Transportation/MapServer/tile/{z}/{y}/{x}',
                        attr='Esri', name='Roads', overlay=True, opacity=0.8,
                    ).add_to(dm)
                    folium.TileLayer(
                        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Boundaries_and_Places/MapServer/tile/{z}/{y}/{x}',
                        attr='Esri', name='Labels', overlay=True, opacity=0.9,
                    ).add_to(dm)
                dm.fit_bounds([[sw_lat_d, sw_lon_d], [ne_lat_d, ne_lon_d]])
                dm.get_root().header.add_child(
                    folium.Element('<style>body{margin:0;padding:0;}.leaflet-control-attribution{display:none!important;}</style>')
                )

                # Planned (all Rx survey points)
                planned_features = [
                    {'type': 'Feature',
                     'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]},
                     'properties': {}}
                    for _, row in rx_sur_df.iterrows()
                ]
                folium.GeoJson(
                    {'type': 'FeatureCollection', 'features': planned_features},
                    marker=folium.CircleMarker(radius=2, fill=True, fill_color=report.job.deployment_color_planned, color=report.job.deployment_color_planned, fill_opacity=0.7),
                ).add_to(dm)

                # Previous days (cyan)
                if not merged_prev.empty:
                    lons_p, lats_p = t.transform(merged_prev['X'].values, merged_prev['Y'].values)
                    merged_prev = merged_prev.copy()
                    merged_prev['lat'] = lats_p
                    merged_prev['lon'] = lons_p
                    prev_features = [
                        {'type': 'Feature',
                         'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]},
                         'properties': {}}
                        for _, row in merged_prev.iterrows()
                    ]
                    folium.GeoJson(
                        {'type': 'FeatureCollection', 'features': prev_features},
                        marker=folium.CircleMarker(radius=3, fill=True, fill_color=report.job.deployment_color_prev, color=report.job.deployment_color_prev, fill_opacity=0.9),
                    ).add_to(dm)

                # Today (orange)
                if not merged_today.empty:
                    lons_t, lats_t = t.transform(merged_today['X'].values, merged_today['Y'].values)
                    merged_today = merged_today.copy()
                    merged_today['lat'] = lats_t
                    merged_today['lon'] = lons_t
                    today_features = [
                        {'type': 'Feature',
                         'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]},
                         'properties': {}}
                        for _, row in merged_today.iterrows()
                    ]
                    folium.GeoJson(
                        {'type': 'FeatureCollection', 'features': today_features},
                        marker=folium.CircleMarker(radius=3, fill=True, fill_color=report.job.deployment_color_today, color=report.job.deployment_color_today, fill_opacity=0.9),
                    ).add_to(dm)
                import time as _time
                map_path = os.path.join(settings.MEDIA_ROOT, 'maps', f'deployment_{report.pk}.html')
                dm.save(map_path)
                deployment_map_html = f'{settings.MEDIA_URL}maps/deployment_{report.pk}.html?v={int(_time.time())}'
        except Exception as e:
            deployment_map_stats['outer_error'] = str(e)
            debug_info['deployment_map_error'] = str(e)

    # Active patch map — second deployment mode using last_line/station_in_ground
    active_patch_map_html = None
    last_line = report.last_line_in_ground
    last_station = report.last_station_in_ground
    if rx_df_cache is not None and last_line is not None and last_station is not None:
        try:
            rps_ff2 = report.job.rps_file if report.job.rps_file else None
            if rps_ff2:
                rx_sur_df2, datum_key2, zone2 = _parse_sps21(rps_ff2)
                epsg_in2 = _resolve_epsg(report.job, datum_key2, zone2)
                t2 = Transformer.from_crs(epsg_in2, 'EPSG:4326', always_xy=True)
                rx_sur_df2['_line'] = rx_sur_df2['Line'].round(1)
                rx_sur_df2['_point'] = rx_sur_df2['Point'].round(1)

                dep2 = rx_df_cache.copy()
                dep2['_line'] = dep2['Line'].round(1)
                dep2['_point'] = dep2['Point'].round(1)
                job_tz2 = _get_job_tz(report.job)
                _dep_dt2_patch = (
                    pd.to_datetime(dep2['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                    .dt.tz_convert(job_tz2)
                )
                dep2['_local_date'] = _dep_dt2_patch.dt.date
                _ds2 = pd.Timestamp(report.date, tz=job_tz2)
                _de2 = pd.Timestamp(report.date + timedelta(days=1), tz=job_tz2)
                _today_mask2 = (_dep_dt2_patch >= _ds2) & (_dep_dt2_patch < _de2)
                dep_today2 = dep2[_today_mask2]
                dep_prev2  = dep2[~_today_mask2]

                ll = float(last_line)
                ls = float(last_station)

                def _in_ground(row):
                    rl, rp = float(row['_line']), float(row['_point'])
                    if rl > ll:
                        return True
                    if rl == ll and rp >= ls:
                        return True
                    return False

                dep_in_ground = dep_prev2[dep_prev2.apply(_in_ground, axis=1)]

                merged_today2   = dep_today2.merge(rx_sur_df2[['_line', '_point', 'X', 'Y']], on=['_line', '_point'], how='inner')
                merged_inground = dep_in_ground.merge(rx_sur_df2[['_line', '_point', 'X', 'Y']], on=['_line', '_point'], how='inner')

                # All Rx survey points → latlon
                sur_lons2, sur_lats2 = t2.transform(rx_sur_df2['X'].values, rx_sur_df2['Y'].values)
                rx_sur_df2['lat'] = sur_lats2; rx_sur_df2['lon'] = sur_lons2

                if not merged_today2.empty:
                    tl2, ta2 = t2.transform(merged_today2['X'].values, merged_today2['Y'].values)
                    merged_today2 = merged_today2.copy(); merged_today2['lat'] = ta2; merged_today2['lon'] = tl2
                if not merged_inground.empty:
                    gl, ga = t2.transform(merged_inground['X'].values, merged_inground['Y'].values)
                    merged_inground = merged_inground.copy(); merged_inground['lat'] = ga; merged_inground['lon'] = gl

                _bound_df2 = merged_today2 if not merged_today2.empty else (merged_inground if not merged_inground.empty else rx_sur_df2)
                buf2 = 500
                sw_lon2, sw_lat2 = t2.transform(_bound_df2['X'].min() - buf2, _bound_df2['Y'].min() - buf2)
                ne_lon2, ne_lat2 = t2.transform(_bound_df2['X'].max() + buf2, _bound_df2['Y'].max() + buf2)

                ap_map = folium.Map(tiles=None, zoom_control=True, width='100%', height=500,
                                    scrollWheelZoom=False, dragging=True, doubleClickZoom=True)
                ap_map.fit_bounds([[sw_lat2, sw_lon2], [ne_lat2, ne_lon2]])
                ap_map.get_root().header.add_child(
                    folium.Element('<style>body{margin:0;padding:0;}.leaflet-control-attribution{display:none!important;}</style>')
                )
                folium.TileLayer(
                    tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
                    attr='Esri',
                ).add_to(ap_map)

                def _ap_layer(features, color, radius):
                    if features:
                        folium.GeoJson(
                            {'type': 'FeatureCollection', 'features': features},
                            marker=folium.CircleMarker(radius=radius, fill=True, fill_color=color, color=color, fill_opacity=0.9),
                        ).add_to(ap_map)

                def _feats2(df): return [{'type': 'Feature', 'geometry': {'type': 'Point', 'coordinates': [r['lon'], r['lat']]}, 'properties': {}} for _, r in df.iterrows()]

                _ap_layer(_feats2(rx_sur_df2), report.job.deployment_color_planned, 2)
                _ap_layer(_feats2(merged_inground), '#28a745', 4)   # in-ground: green
                _ap_layer(_feats2(merged_today2), report.job.deployment_color_today, 4)

                import time as _time2
                ap_path = os.path.join(settings.MEDIA_ROOT, 'maps', f'active_patch_{report.pk}.html')
                os.makedirs(os.path.dirname(ap_path), exist_ok=True)
                ap_map.save(ap_path)
                active_patch_map_html = f'{settings.MEDIA_URL}maps/active_patch_{report.pk}.html?v={int(_time2.time())}'
        except Exception as e:
            debug_info['active_patch_error'] = str(e)

    # Progress map for Report tab — only when this report has shot data
    if shot_stats:
        sps_ff = report.job.sps_file if report.job.sps_file else None
        if sps_ff:
            try:
                sx_df, datum_key, zone = _parse_sps21(sps_ff)
                epsg_in = _resolve_epsg(report.job, datum_key, zone)
                sx_df['_line'] = sx_df['Line'].round(1)
                sx_df['_point'] = sx_df['Point'].round(1)
                transformer = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)
                lons, lats = transformer.transform(sx_df['X'].values, sx_df['Y'].values)
                sx_df['lat'] = lats
                sx_df['lon'] = lons

                today_set, prev_set, void_set = set(), set(), set()
                for r in report.job.reports.filter(date__lte=report.date):
                    obs = r.files.filter(file_type='obslog').first()
                    if obs:
                        try:
                            o = _read_csv(obs.file.path, skiprows=2)
                            o.columns = o.columns.str.strip()
                            is_today = (r.pk == report.pk)
                            for _, row in o.iterrows():
                                key = (round(float(row['Line']), 1), round(float(row['Station']), 1))
                                if str(row['Status']).strip() == 'Void':
                                    void_set.add(key)
                                elif is_today:
                                    today_set.add(key)
                                else:
                                    prev_set.add(key)
                        except Exception:
                            pass

                planned_f, today_f, prev_f = [], [], []
                for _, row in sx_df.iterrows():
                    key = (row['_line'], row['_point'])
                    feat = {'type': 'Feature', 'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]}, 'properties': {}}
                    if key in today_set:
                        today_f.append(feat)
                    elif key in prev_set:
                        prev_f.append(feat)
                    elif key not in void_set:
                        planned_f.append(feat)

                buf = 500  # metres
                _today_bound_df = pd.DataFrame()
                if map_bounds == 'today':
                    _obs_today = report.files.filter(file_type='obslog').first()
                    if _obs_today:
                        try:
                            _ot = _read_csv(_obs_today.file.path, skiprows=2)
                            _ot.columns = _ot.columns.str.strip()
                            if 'Status' in _ot.columns:
                                _ot = _ot[_ot['Status'].astype(str).str.strip() != 'Void']
                            _ot['_line'] = _ot['Line'].round(1)
                            _ot['_point'] = _ot['Station'].round(1)
                            _today_bound_df = sx_df.merge(
                                _ot[['_line', '_point']].drop_duplicates(),
                                on=['_line', '_point'], how='inner'
                            )
                        except Exception:
                            pass
                _prog_bound_df = _today_bound_df if not _today_bound_df.empty else sx_df
                sw_lon, sw_lat = transformer.transform(_prog_bound_df['X'].min() - buf, _prog_bound_df['Y'].min() - buf)
                ne_lon, ne_lat = transformer.transform(_prog_bound_df['X'].max() + buf, _prog_bound_df['Y'].max() + buf)

                m = folium.Map(tiles=None, zoom_control=True, width='100%', height=500,
                               scrollWheelZoom=False, dragging=True, doubleClickZoom=True)
                m.fit_bounds([[sw_lat, sw_lon], [ne_lat, ne_lon]])
                m.get_root().header.add_child(
                    folium.Element('<style>body{margin:0;padding:0;}.leaflet-control-attribution{display:none!important;}</style>')
                )
                folium.TileLayer(
                    tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
                    attr='Esri', name='Satellite',
                ).add_to(m)
                if report.job.show_map_overlay:
                    folium.TileLayer(
                        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Transportation/MapServer/tile/{z}/{y}/{x}',
                        attr='Esri', name='Roads', overlay=True, opacity=0.8,
                    ).add_to(m)
                    folium.TileLayer(
                        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Boundaries_and_Places/MapServer/tile/{z}/{y}/{x}',
                        attr='Esri', name='Labels', overlay=True, opacity=0.9,
                    ).add_to(m)

                def add_layer(features, color, radius):
                    if features:
                        folium.GeoJson(
                            {'type': 'FeatureCollection', 'features': features},
                            marker=folium.CircleMarker(radius=radius, fill=True, fill_color=color, color=color, fill_opacity=0.9),
                        ).add_to(m)

                add_layer(planned_f, report.job.progress_color_planned, 2)
                add_layer(prev_f,    report.job.progress_color_prev,    4)
                add_layer(today_f,   report.job.progress_color_today,   4)
                map_path = os.path.join(settings.MEDIA_ROOT, 'maps', f'progress_{report.pk}.html')
                os.makedirs(os.path.dirname(map_path), exist_ok=True)
                m.save(map_path)
                import time as _time
                progress_map_html = f'{settings.MEDIA_URL}maps/progress_{report.pk}.html?v={int(_time.time())}'
            except Exception as e:
                debug_info['progress_map_error'] = str(e)

    # PSS QC maps
    pss_qc_maps_by_unit = []
    if True:
        pss_file = report_files.filter(file_type='pss').first()
        _sps_ff = report.job.sps_file if report.job.sps_file else None
        if pss_file and _sps_ff:
            try:
                pss_df = _read_csv(pss_file.file.path)
                pss_df.columns = pss_df.columns.str.strip()
                if 'Void' in pss_df.columns:
                    pss_df = pss_df[pss_df['Void'].isna()]
                pss_df['_line'] = pss_df['Line'].round(1)
                pss_df['_point'] = pss_df['Station'].round(1)

                # Filter to today's shots via obs log (PSS files are often cumulative)
                obs_file = report_files.filter(file_type='obslog').first()
                if obs_file:
                    try:
                        obs_df = _read_csv(obs_file.file.path, skiprows=2)
                        obs_df.columns = obs_df.columns.str.strip()
                        if 'Status' in obs_df.columns:
                            obs_df = obs_df[obs_df['Status'].astype(str).str.strip() != 'Void']
                        obs_df['_line'] = obs_df['Line'].round(1)
                        obs_df['_point'] = obs_df['Station'].round(1)
                        today_filtered = pss_df.merge(
                            obs_df[['_line', '_point']].drop_duplicates(),
                            on=['_line', '_point'], how='inner'
                        )
                        if not today_filtered.empty:
                            pss_df = today_filtered
                    except Exception as e:
                        debug_info['pss_obslog_filter_error'] = str(e)

                sx_df, datum_key, zone = _parse_sps21(_sps_ff)
                epsg_in = _resolve_epsg(report.job, datum_key, zone)
                sx_df['_line'] = sx_df['Line'].round(1)
                sx_df['_point'] = sx_df['Point'].round(1)
                t = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)

                job = report.job
                params = [
                    ('Phase Max', job.pss_phase_max_green, job.pss_phase_max_amber, False, 'phase_max', '°'),
                    ('Phase Avg', job.pss_phase_avg_green, job.pss_phase_avg_amber, False, 'phase_avg', '°'),
                    ('Force Max', job.pss_force_max_green, job.pss_force_max_amber, True,  'force_max', '%'),
                    ('Force Avg', job.pss_force_avg_green, job.pss_force_avg_amber, True,  'force_avg', '%'),
                    ('THD Max',   job.pss_thd_max_green,   job.pss_thd_max_amber,   False, 'thd_max',   '%'),
                    ('THD Avg',   job.pss_thd_avg_green,   job.pss_thd_avg_amber,   False, 'thd_avg',   '%'),
                ]

                def make_pss_map(col, g_thresh, a_thresh, higher_is_better, slug, unit, interactive=False, src_df=None):
                    if src_df is None:
                        src_df = pss_df
                    agg = src_df.groupby(['_line', '_point'])[col].apply(
                        lambda x: x.abs().max() if not higher_is_better else x.mean()
                    ).reset_index()
                    agg.columns = ['_line', '_point', 'val']
                    merged = agg.merge(sx_df[['_line', '_point', 'X', 'Y']], on=['_line', '_point'], how='inner')
                    if merged.empty:
                        return None
                    lons, lats = t.transform(merged['X'].values, merged['Y'].values)
                    merged['lat'] = lats
                    merged['lon'] = lons

                    _buf = 100
                    sw_lon, sw_lat = t.transform(merged['X'].min() - _buf, merged['Y'].min() - _buf)
                    ne_lon, ne_lat = t.transform(merged['X'].max() + _buf, merged['Y'].max() + _buf)

                    green_f, amber_f, red_f = [], [], []
                    for _, row in merged.iterrows():
                        v = row['val']
                        feat = {'type': 'Feature',
                                'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]},
                                'properties': {}}
                        if higher_is_better:
                            if v >= g_thresh:    green_f.append(feat)
                            elif v >= a_thresh:  amber_f.append(feat)
                            else:                red_f.append(feat)
                        else:
                            if v <= g_thresh:    green_f.append(feat)
                            elif v <= a_thresh:  amber_f.append(feat)
                            else:                red_f.append(feat)

                    pm = folium.Map(tiles=None, zoom_control=interactive, width='100%', height=200,
                                    scrollWheelZoom=interactive, dragging=interactive, doubleClickZoom=interactive)
                    folium.TileLayer(
                        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
                        attr='Esri', name='Satellite',
                    ).add_to(pm)
                    pm.fit_bounds([[sw_lat, sw_lon], [ne_lat, ne_lon]])
                    pm.get_root().header.add_child(
                        folium.Element('<style>body{margin:0;padding:0;}.leaflet-control-attribution{display:none!important;}</style>')
                    )
                    for feats, color in [(green_f, '#4caf50'), (amber_f, '#ff9800'), (red_f, '#e63946')]:
                        if feats:
                            folium.GeoJson(
                                {'type': 'FeatureCollection', 'features': feats},
                                marker=folium.CircleMarker(radius=3, fill=True, fill_color=color, color=color, fill_opacity=0.9),
                            ).add_to(pm)
                    path = os.path.join(settings.MEDIA_ROOT, 'maps', f'pss_{slug}_{report.pk}.html')
                    pm.save(path)
                    return f'{settings.MEDIA_URL}maps/pss_{slug}_{report.pk}.html'

                # Per-unit sets
                if 'Unit ID' in pss_df.columns:
                    import re as _re
                    for unit_val in sorted(pss_df['Unit ID'].unique()):
                        unit_df = pss_df[pss_df['Unit ID'] == unit_val]
                        safe = _re.sub(r'[^a-zA-Z0-9]', '_', str(unit_val))
                        unit_maps = {}
                        for i, (col, g, a, higher, slug, unit_lbl) in enumerate(params):
                            try:
                                url = make_pss_map(col, g, a, higher, f'{slug}_u{safe}', unit_lbl,
                                                   interactive=(i == 0), src_df=unit_df)
                                if url:
                                    unit_maps[col] = {'url': url, 'g': g, 'a': a, 'higher': higher,
                                                      'unit': unit_lbl, 'leader': (i == 0)}
                            except Exception as e:
                                debug_info[f'pss_unit_{safe}_{slug}_error'] = str(e)
                        if unit_maps:
                            pss_qc_maps_by_unit.append({'label': str(unit_val), 'slug': safe, 'maps': unit_maps})

            except Exception as e:
                debug_info['pss_outer_error'] = str(e)

    # Chargeable hours by activity type
    from datetime import datetime, date as _date
    import re as _re

    _is_supervisor = report.report_type == 'supervisor'

    def _activity_mins(a):
        """Return duration in minutes for an activity row."""
        if _is_supervisor:
            h = a.hours.strip() if a.hours else ''
            if not h:
                return 0
            try:
                return round(float(h) * 60)
            except ValueError:
                pass
            m = _re.match(r'(\d+(?:\.\d+)?)\s*h(?:ours?)?\s*(?:(\d+)\s*m(?:in)?)?', h, _re.IGNORECASE)
            if m:
                return round(float(m.group(1)) * 60) + int(m.group(2) or 0)
            m = _re.match(r'(\d+):(\d+)', h)
            if m:
                return int(m.group(1)) * 60 + int(m.group(2))
            return 0
        start = datetime.combine(_date.today(), a.start_time)
        end = datetime.combine(_date.today(), a.end_time)
        return int((end - start).total_seconds() / 60)

    type_pct = {at.name: at.chargeable_percentage for at in ActivityType.objects.all()}
    chargeable_by_type = {}
    for a in report.activities.all():
        mins = _activity_mins(a)
        pct = type_pct.get(a.activity_type, 100)
        ch_mins = round(mins * pct / 100)
        if a.activity_type not in chargeable_by_type:
            chargeable_by_type[a.activity_type] = {'total_mins': 0, 'chargeable_mins': 0, 'pct': pct}
        chargeable_by_type[a.activity_type]['total_mins'] += mins
        chargeable_by_type[a.activity_type]['chargeable_mins'] += ch_mins
    total_chargeable_mins = sum(v['chargeable_mins'] for v in chargeable_by_type.values())

    def _fmt(mins):
        return f"{mins // 60}h {mins % 60:02d}m"

    chargeable_rows = [
        {
            'type': k,
            'pct': v['pct'],
            'total': _fmt(v['total_mins']),
            'chargeable': _fmt(v['chargeable_mins']),
        }
        for k, v in chargeable_by_type.items()
    ]

    # Job-wide chargeable hours (same report type only)
    job_chargeable_by_type = {}
    for a in Activity.objects.filter(report__job=report.job, report__report_type=report.report_type):
        mins = _activity_mins(a)
        pct = type_pct.get(a.activity_type, 100)
        ch_mins = round(mins * pct / 100)
        if a.activity_type not in job_chargeable_by_type:
            job_chargeable_by_type[a.activity_type] = {'total_mins': 0, 'chargeable_mins': 0, 'pct': pct}
        job_chargeable_by_type[a.activity_type]['total_mins'] += mins
        job_chargeable_by_type[a.activity_type]['chargeable_mins'] += ch_mins
    job_total_chargeable_mins = sum(v['chargeable_mins'] for v in job_chargeable_by_type.values())
    job_chargeable_rows = [
        {
            'type': k,
            'pct': v['pct'],
            'total': _fmt(v['total_mins']),
            'chargeable': _fmt(v['chargeable_mins']),
        }
        for k, v in job_chargeable_by_type.items()
    ]

    # Pie chart data — by category
    daily_cat = {}
    for a in report.activities.all():
        mins = _activity_mins(a)
        daily_cat[a.category] = daily_cat.get(a.category, 0) + mins
    daily_pie_data = [{'label': k, 'mins': v} for k, v in daily_cat.items() if v > 0]

    job_cat = {}
    for a in Activity.objects.filter(report__job=report.job, report__report_type=report.report_type):
        mins = _activity_mins(a)
        job_cat[a.category] = job_cat.get(a.category, 0) + mins
    job_pie_data = [{'label': k, 'mins': v} for k, v in job_cat.items() if v > 0]

    return render(request, 'reports/report_detail.html', {
        'report': report,
        'observers': observers,
        'operators': operators,
        'saved_observers': (
            report.observers.splitlines() if report.observers
            else (DailyReport.objects.filter(job=report.job, date__lt=report.date)
                  .exclude(observers='').order_by('-date')
                  .values_list('observers', flat=True).first() or '').splitlines()
        ),
        'saved_operators': (
            report.operators.splitlines() if report.operators
            else (DailyReport.objects.filter(job=report.job, date__lt=report.date)
                  .exclude(operators='').order_by('-date')
                  .values_list('operators', flat=True).first() or '').splitlines()
        ),
        'report_files': report_files,
        'rps_file': report.job.rps_file,
        'rps_count': report.job.rps_count,
        'rps_is_final': report.job.rps_is_final,
        'rps_estimated': report.job.estimated_rx_count,
        'sps_file': report.job.sps_file,
        'sps_count': report.job.sps_count,
        'sps_is_final': report.job.sps_is_final,
        'sps_estimated': report.job.estimated_sx_count,
        'survey_epsg': report.job.survey_epsg,
        'epsg_list': _load_epsg_list(),
        'job_timezone': report.job.timezone,
        'job_utc_offset_custom': report.job.utc_offset_custom,
        'timezone_choices': report.job.TIMEZONE_CHOICES,
        'activities': report.activities.all(),
        'active_tab': active_tab,
        'categories': categories,
        'default_category': default_category,
        'activity_types': ActivityType.objects.all(),
        'shot_stats': shot_stats,
        'rx_deployed': rx_deployed,
        'nodes_today': nodes_today,
        'nodes_today_line_stats': nodes_today_line_stats,
        'rx_stats': rx_stats,
        'job_total_nodes': job_total_nodes or None,
        'nodes_remaining': nodes_remaining,
        'daily_avg_shots': daily_avg_shots,
        'daily_avg_nodes': daily_avg_nodes,
        'progress_map_html': progress_map_html,
        'deployment_map_html': deployment_map_html,
        'deployment_map_stats': deployment_map_stats,
        'active_patch_map_html': active_patch_map_html,
        'last_line_in_ground': report.last_line_in_ground,
        'last_station_in_ground': report.last_station_in_ground,
        'job_total_shots': job_total_shots or None,

        'shots_remaining': shots_remaining,
        'pss_presets': pss_presets,
        'diary_templates': diary_templates,
        'debug_info': debug_info,
        'map_bounds': map_bounds,
        'pss_qc_maps_by_unit': pss_qc_maps_by_unit,
        'chargeable_rows': chargeable_rows,
        'total_chargeable': _fmt(total_chargeable_mins),
        'job_chargeable_rows': job_chargeable_rows,
        'job_total_chargeable': _fmt(job_total_chargeable_mins),
        'daily_pie_data': daily_pie_data,
        'job_pie_data': job_pie_data,
        'activity_types_json': list(ActivityType.objects.values('name', 'chargeable_percentage')),
        'report_shot_rows': report_shot_rows,
        'finish_estimates': finish_estimates,
        'node_finish_estimates': node_finish_estimates,
        'photos': report.photos.all(),
    })


def _build_report_ctx(report, include_pss_units=None, map_views=None):
    """Build the shared context dict used by both PDF and Word export views."""
    report_files = report.files.all()

    shot_stats = None
    obs_file = report_files.filter(file_type='obslog').first()
    if obs_file:
        try:
            df = _read_csv(obs_file.file.path, skiprows=2)
            df.columns = df.columns.str.strip()
            void_mask = df['Status'].str.strip() == 'Void'
            prod_df = df[~void_mask]
            line_stats = []
            for line, grp in prod_df.groupby('Line', sort=True):
                line_stats.append({'line': int(line) if line == int(line) else line,
                                   'start': grp['Station'].min(), 'end': grp['Station'].max(), 'shots': len(grp)})
            shot_intervals = []
            if len(prod_df):
                try:
                    times = pd.to_datetime(prod_df['Local Time'].str.strip(), format='%H:%M:%S.%f', errors='coerce')
                    times = times.dropna()
                    if len(times):
                        slot = times.min().floor('15min')
                        end_slot = times.max().floor('15min')
                        bins = []
                        while slot <= end_slot:
                            bins.append(slot)
                            slot += pd.Timedelta(minutes=15)
                        counts = pd.cut(times, bins=[b for b in bins] + [bins[-1] + pd.Timedelta(minutes=15)],
                                        labels=[b.strftime('%H:%M') for b in bins], right=False)
                        vc = counts.value_counts().reindex([b.strftime('%H:%M') for b in bins], fill_value=0)
                        shot_intervals = [{'interval': lbl, 'shots': int(n)} for lbl, n in vc.items()]
                except Exception:
                    pass
            shot_stats = {'production': len(prod_df), 'void': int(void_mask.sum()), 'line_stats': line_stats, 'shot_intervals': shot_intervals}
        except Exception:
            pass

    rx_stats = None
    rx_file = report_files.filter(file_type='rx_deployment').first()
    if rx_file:
        try:
            rx_df = _read_csv(rx_file.file.path)
            rx_df.columns = rx_df.columns.str.strip()
            line_stats = []
            for line, grp in rx_df.groupby('Line', sort=True):
                line_stats.append({'line': int(line) if line == int(line) else line,
                                   'start': grp['Point'].min(), 'end': grp['Point'].max(), 'nodes': len(grp)})
            rx_stats = {'total': len(rx_df), 'line_stats': line_stats}
        except Exception:
            pass

    job_total_shots = avg_shots_total = production_days = 0
    _ctx_shot_rows = []
    for r in report.job.reports.filter(date__lte=report.date).order_by('date'):
        obs = r.files.filter(file_type='obslog').first()
        day_prod = None
        if obs:
            try:
                d = _read_csv(obs.file.path, skiprows=2)
                d.columns = d.columns.str.strip()
                day_prod = int((d['Status'].str.strip() != 'Void').sum())
                if day_prod > 0:
                    job_total_shots += day_prod
                    if r.include_in_avg:
                        avg_shots_total += day_prod
                        production_days += 1
            except Exception:
                pass
        _ctx_shot_rows.append({'report': r, 'shots': day_prod})

    # Node stats — cumulative rx_deployment file
    job_total_nodes = avg_nodes_total = deployment_days = 0
    _ctx_node_rows = []
    _latest_rx = report.job.reports.filter(
        date__lte=report.date, files__file_type='rx_deployment'
    ).order_by('-date').first()
    if _latest_rx:
        _rx_f = _latest_rx.files.filter(file_type='rx_deployment').first()
        if _rx_f:
            try:
                _rx_cum = _read_csv(_rx_f.file.path)
                _rx_cum.columns = _rx_cum.columns.str.strip()
                _tz3 = _get_job_tz(report.job)
                _dep_dt3 = (
                    pd.to_datetime(_rx_cum['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                    .dt.tz_convert(_tz3)
                )
                _cutoff3 = pd.Timestamp(report.date + timedelta(days=1), tz=_tz3)
                _mask = _dep_dt3 < _cutoff3
                _rx_cum = _rx_cum[_mask].copy()
                _rx_cum['_dep_date'] = _dep_dt3[_mask].dt.date.values
                job_total_nodes = len(_rx_cum)
                _nodes_by_date = _rx_cum.groupby('_dep_date').size().to_dict()
                _report_by_date = {r.date: r for r in report.job.reports.filter(date__lte=report.date)}
                for dep_date in sorted(_nodes_by_date):
                    dn = _nodes_by_date[dep_date]
                    r2 = _report_by_date.get(dep_date)
                    if r2 and r2.include_in_avg:
                        avg_nodes_total += dn
                        deployment_days += 1
                    _ctx_node_rows.append({'report': r2, 'nodes': dn, 'date': dep_date})
            except Exception:
                pass

    planned_shots = None
    if report.job.sps_is_final and report.job.sps_count:
        planned_shots = report.job.sps_count
    if planned_shots is None and report.job.estimated_sx_count:
        planned_shots = report.job.estimated_sx_count

    planned_nodes = None
    if report.job.rps_is_final and report.job.rps_count:
        planned_nodes = report.job.rps_count
    if planned_nodes is None and report.job.estimated_rx_count:
        planned_nodes = report.job.estimated_rx_count

    maps_dir = os.path.join(settings.MEDIA_ROOT, 'maps')
    os.makedirs(maps_dir, exist_ok=True)

    SATELLITE = 'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}'
    HIDE_CSS = '<style>body{margin:0;padding:0;}.leaflet-control-attribution{display:none!important;}</style>'

    _map_views = map_views or {}

    def _save_map(m, slug, center=None, zoom=None, width=900, height=500):
        html_path = os.path.join(maps_dir, f'{slug}_{report.pk}.html')
        png_path  = os.path.join(maps_dir, f'{slug}_{report.pk}.png')
        if not os.path.exists(html_path):
            m.save(html_path)
        extra_wait = 1500 if report.job.show_map_overlay else 0
        _map_to_png(html_path, png_path, width=width, height=height, center=center, zoom=zoom, extra_wait=extra_wait)
        return png_path if os.path.exists(png_path) else None

    def _base_map(bounds_sw, bounds_ne, t, sx_x, sx_y, buf=100):
        sw = t.transform(sx_x.min() - buf, sx_y.min() - buf)
        ne = t.transform(sx_x.max() + buf, sx_y.max() + buf)
        m = folium.Map(tiles=None, zoom_control=False, width='100%', height=500,
                       scrollWheelZoom=False, dragging=False, doubleClickZoom=False)
        folium.TileLayer(tiles=SATELLITE, attr='Esri').add_to(m)
        if report.job.show_map_overlay:
            folium.TileLayer(
                tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Transportation/MapServer/tile/{z}/{y}/{x}',
                attr='Esri', overlay=True, opacity=0.8,
            ).add_to(m)
            folium.TileLayer(
                tiles='https://server.arcgisonline.com/ArcGIS/rest/services/Reference/World_Boundaries_and_Places/MapServer/tile/{z}/{y}/{x}',
                attr='Esri', overlay=True, opacity=0.9,
            ).add_to(m)
        m.fit_bounds([[sw[1], sw[0]], [ne[1], ne[0]]])
        mv = m.get_name()
        m.get_root().script.add_child(folium.Element(f"setTimeout(function(){{ {mv}.zoomIn(1); }}, 300);"))
        m.get_root().header.add_child(folium.Element(HIDE_CSS))
        return m

    def _geojson_layer(m, features, color, radius=4):
        if features:
            folium.GeoJson(
                {'type': 'FeatureCollection', 'features': features},
                marker=folium.CircleMarker(radius=radius, fill=True, fill_color=color,
                                           color=color, fill_opacity=0.9),
            ).add_to(m)

    deployment_png = progress_png = None
    _sps_ff2 = report.job.sps_file if report.job.sps_file else None
    _rps_ff2 = report.job.rps_file if report.job.rps_file else None
    _deploy_mode = (_map_views or {}).get('deployment_mode', 'alldays')

    # --- Deployment map (All Days mode) ---
    rx_file2 = report_files.filter(file_type='rx_deployment').first()
    if rx_file2 and _rps_ff2 and _deploy_mode != 'patch':
        try:
            from zoneinfo import ZoneInfo
            rx_sur_df, datum_key, zone = _parse_sps21(_rps_ff2)
            epsg_in = _resolve_epsg(report.job, datum_key, zone)
            t = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)
            rx_sur_df['_line'] = rx_sur_df['Line'].round(1)
            rx_sur_df['_point'] = rx_sur_df['Point'].round(1)
            dep_df = _read_csv(rx_file2.file.path)
            dep_df.columns = dep_df.columns.str.strip()
            dep_df['_line'] = dep_df['Line'].round(1)
            dep_df['_point'] = dep_df['Point'].round(1)
            # Split today vs previous using Deployment_Time if available
            if 'Deployment_Time' in dep_df.columns:
                job_tz = _get_job_tz(report.job)
                _dep_dt_alldays = (
                    pd.to_datetime(dep_df['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                    .dt.tz_convert(job_tz)
                )
                _ds = pd.Timestamp(report.date, tz=job_tz)
                _de = pd.Timestamp(report.date + timedelta(days=1), tz=job_tz)
                _today_mask = (_dep_dt_alldays >= _ds) & (_dep_dt_alldays < _de)
                dep_today = dep_df[_today_mask]
                dep_prev  = dep_df[~_today_mask]
            else:
                dep_today = dep_df
                dep_prev  = dep_df.iloc[0:0]
            merged_today = dep_today.merge(rx_sur_df[['_line','_point','X','Y']], on=['_line','_point'], how='inner')
            merged_prev  = dep_prev.merge(rx_sur_df[['_line','_point','X','Y']], on=['_line','_point'], how='inner')
            sl, la = t.transform(rx_sur_df['X'].values, rx_sur_df['Y'].values)
            rx_sur_df['lat'] = la; rx_sur_df['lon'] = sl
            if not merged_prev.empty:
                pl, pa = t.transform(merged_prev['X'].values, merged_prev['Y'].values)
                merged_prev = merged_prev.copy(); merged_prev['lat'] = pa; merged_prev['lon'] = pl
            if not merged_today.empty:
                tl, ta = t.transform(merged_today['X'].values, merged_today['Y'].values)
                merged_today = merged_today.copy(); merged_today['lat'] = ta; merged_today['lon'] = tl
            def _feats(df): return [{'type':'Feature','geometry':{'type':'Point','coordinates':[r['lon'],r['lat']]},'properties':{}} for _,r in df.iterrows()]
            dm = _base_map(None, None, t, rx_sur_df['X'], rx_sur_df['Y'])
            _geojson_layer(dm, _feats(rx_sur_df), report.job.deployment_color_planned, 2)
            _geojson_layer(dm, _feats(merged_prev), report.job.deployment_color_prev, 3)
            _geojson_layer(dm, _feats(merged_today), report.job.deployment_color_today, 3)
            _add_map_legend(dm, [(report.job.deployment_color_today, 'Today'), (report.job.deployment_color_prev, 'Previous Days'), (report.job.deployment_color_planned, 'Planned')])
            _dv = _map_views.get('deployment')
            deployment_png = _save_map(dm, 'deployment',
                                       center=_dv[:2] if _dv else None,
                                       zoom=_dv[2] if _dv else None)
        except Exception:
            pass

    # --- Active patch map (patch mode only) ---
    active_patch_png = None
    active_patch_stats = None
    _ll = report.last_line_in_ground
    _ls = report.last_station_in_ground
    if rx_file2 and _rps_ff2 and _ll is not None and _ls is not None and _deploy_mode == 'patch':
        try:
            rx_sur_ap, dk_ap, zone_ap = _parse_sps21(_rps_ff2)
            t_ap = Transformer.from_crs(_resolve_epsg(report.job, dk_ap, zone_ap), 'EPSG:4326', always_xy=True)
            rx_sur_ap['_line'] = rx_sur_ap['Line'].round(1)
            rx_sur_ap['_point'] = rx_sur_ap['Point'].round(1)
            dep_ap = _read_csv(rx_file2.file.path)
            dep_ap.columns = dep_ap.columns.str.strip()
            dep_ap['_line'] = dep_ap['Line'].round(1)
            dep_ap['_point'] = dep_ap['Point'].round(1)
            if 'Deployment_Time' in dep_ap.columns:
                _tz_ap = _get_job_tz(report.job)
                _dep_dt_ap = (
                    pd.to_datetime(dep_ap['Deployment_Time'], dayfirst=False, errors='coerce', utc=True)
                    .dt.tz_convert(_tz_ap)
                )
                _ds_ap = pd.Timestamp(report.date, tz=_tz_ap)
                _de_ap = pd.Timestamp(report.date + timedelta(days=1), tz=_tz_ap)
                _today_mask_ap = (_dep_dt_ap >= _ds_ap) & (_dep_dt_ap < _de_ap)
                dep_today_ap = dep_ap[_today_mask_ap]
                dep_prev_ap  = dep_ap[~_today_mask_ap]
            else:
                dep_today_ap = dep_ap; dep_prev_ap = dep_ap.iloc[0:0]
            ll_f, ls_f = float(_ll), float(_ls)
            def _in_gnd(row):
                rl, rp = float(row['_line']), float(row['_point'])
                return rl > ll_f or (rl == ll_f and rp >= ls_f)
            dep_inground = dep_prev_ap[dep_prev_ap.apply(_in_gnd, axis=1)]
            m_today_ap = dep_today_ap.merge(rx_sur_ap[['_line','_point','X','Y']], on=['_line','_point'], how='inner')
            m_ing_ap   = dep_inground.merge(rx_sur_ap[['_line','_point','X','Y']], on=['_line','_point'], how='inner')

            # Build stats
            total_deployed = len(dep_ap)
            total_inground = len(dep_inground)
            total_today    = len(dep_today_ap)
            total_retrieved = len(dep_prev_ap) - total_inground
            total_planned  = len(rx_sur_ap) - total_deployed
            # Front of patch: highest line then station among in-ground nodes
            if not dep_inground.empty:
                front_row = dep_inground.sort_values(['_line', '_point'], ascending=False).iloc[0]
                front_line    = front_row['_line']
                front_station = front_row['_point']
            else:
                front_line = front_station = None
            active_patch_stats = {
                'last_line': _ll,
                'last_station': _ls,
                'front_line': front_line,
                'front_station': front_station,
                'in_ground': total_inground,
                'today': total_today,
                'retrieved': total_retrieved,
                'planned': total_planned,
                'total_deployed': total_deployed,
            }

            sl_ap, la_ap = t_ap.transform(rx_sur_ap['X'].values, rx_sur_ap['Y'].values)
            rx_sur_ap['lat'] = la_ap; rx_sur_ap['lon'] = sl_ap
            if not m_ing_ap.empty:
                gl, ga = t_ap.transform(m_ing_ap['X'].values, m_ing_ap['Y'].values)
                m_ing_ap = m_ing_ap.copy(); m_ing_ap['lat'] = ga; m_ing_ap['lon'] = gl
            if not m_today_ap.empty:
                tl, ta = t_ap.transform(m_today_ap['X'].values, m_today_ap['Y'].values)
                m_today_ap = m_today_ap.copy(); m_today_ap['lat'] = ta; m_today_ap['lon'] = tl
            def _feats_ap(df): return [{'type':'Feature','geometry':{'type':'Point','coordinates':[r['lon'],r['lat']]},'properties':{}} for _,r in df.iterrows()]
            ap_m = _base_map(None, None, t_ap, rx_sur_ap['X'], rx_sur_ap['Y'])
            _geojson_layer(ap_m, _feats_ap(rx_sur_ap), report.job.deployment_color_planned, 2)
            _geojson_layer(ap_m, _feats_ap(m_ing_ap), '#28a745', 3)
            _geojson_layer(ap_m, _feats_ap(m_today_ap), report.job.deployment_color_today, 3)
            _add_map_legend(ap_m, [(report.job.deployment_color_today, 'Today'), ('#28a745', 'In Ground'), (report.job.deployment_color_planned, 'Planned')])
            _apv = _map_views.get('active_patch') or _map_views.get('deployment')
            active_patch_png = _save_map(ap_m, 'active_patch',
                                         center=_apv[:2] if _apv else None,
                                         zoom=_apv[2] if _apv else None)
        except Exception:
            pass

    # --- Shot progress map ---
    if _sps_ff2:
        try:
            sx_df, datum_key, zone = _parse_sps21(_sps_ff2)
            epsg_in = _resolve_epsg(report.job, datum_key, zone)
            t = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)
            sx_df['_line'] = sx_df['Line'].round(1); sx_df['_point'] = sx_df['Point'].round(1)
            lons, lats = t.transform(sx_df['X'].values, sx_df['Y'].values)
            sx_df['lat'] = lats; sx_df['lon'] = lons
            today_set, prev_set, void_set = set(), set(), set()
            for r in report.job.reports.filter(date__lte=report.date):
                obs = r.files.filter(file_type='obslog').first()
                if obs:
                    try:
                        o = _read_csv(obs.file.path, skiprows=2); o.columns = o.columns.str.strip()
                        is_today = (r.pk == report.pk)
                        for _, row in o.iterrows():
                            key = (round(float(row['Line']),1), round(float(row['Station']),1))
                            if str(row['Status']).strip() == 'Void':
                                void_set.add(key)
                            elif is_today:
                                today_set.add(key)
                            else:
                                prev_set.add(key)
                    except Exception:
                        pass
            planned_f, today_f, prev_f = [], [], []
            for _, row in sx_df.iterrows():
                key = (row['_line'], row['_point'])
                feat = {'type':'Feature','geometry':{'type':'Point','coordinates':[row['lon'],row['lat']]},'properties':{}}
                if key in today_set:
                    today_f.append(feat)
                elif key in prev_set:
                    prev_f.append(feat)
                elif key not in void_set:
                    planned_f.append(feat)
            if today_f:
                sm = _base_map(None, None, t, sx_df['X'], sx_df['Y'])
                _geojson_layer(sm, planned_f, report.job.progress_color_planned, 2)
                _geojson_layer(sm, prev_f, report.job.progress_color_prev, 4)
                _geojson_layer(sm, today_f, report.job.progress_color_today, 4)
                _add_map_legend(sm, [(report.job.progress_color_today, 'Today'), (report.job.progress_color_prev, 'Previous Days'), (report.job.progress_color_planned, 'Planned')])
                _pv = _map_views.get('progress')
                progress_png = _save_map(sm, 'progress',
                                         center=_pv[:2] if _pv else None,
                                         zoom=_pv[2] if _pv else None)
        except Exception:
            pass

    # --- PSS QC maps ---
    # pss_png_maps is a list of (unit_label, [(col, png_path), ...])
    pss_png_maps = []
    pss_file = report_files.filter(file_type='pss').first()
    if pss_file and _sps_ff2:
        try:
            import re as _re
            pss_df = _read_csv(pss_file.file.path); pss_df.columns = pss_df.columns.str.strip()
            if 'Void' in pss_df.columns:
                pss_df = pss_df[pss_df['Void'].isna()]
            pss_df['_line'] = pss_df['Line'].round(1); pss_df['_point'] = pss_df['Station'].round(1)
            # Filter to today's shots via obs log (PSS files are often cumulative)
            _obs_file = report_files.filter(file_type='obslog').first()
            if _obs_file:
                try:
                    _obs_df = _read_csv(_obs_file.file.path, skiprows=2)
                    _obs_df.columns = _obs_df.columns.str.strip()
                    if 'Status' in _obs_df.columns:
                        _obs_df = _obs_df[_obs_df['Status'].astype(str).str.strip() != 'Void']
                    _obs_df['_line'] = _obs_df['Line'].round(1)
                    _obs_df['_point'] = _obs_df['Station'].round(1)
                    _today = pss_df.merge(
                        _obs_df[['_line', '_point']].drop_duplicates(),
                        on=['_line', '_point'], how='inner'
                    )
                    if not _today.empty:
                        pss_df = _today
                except Exception:
                    pass
            sx_df, datum_key, zone = _parse_sps21(_sps_ff2)
            epsg_in = _resolve_epsg(report.job, datum_key, zone)
            t = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)
            sx_df['_line'] = sx_df['Line'].round(1); sx_df['_point'] = sx_df['Point'].round(1)
            job = report.job
            params = [
                ('Phase Max', job.pss_phase_max_green, job.pss_phase_max_amber, False, 'phase_max', '\u00b0'),
                ('Phase Avg', job.pss_phase_avg_green, job.pss_phase_avg_amber, False, 'phase_avg', '\u00b0'),
                ('Force Max', job.pss_force_max_green, job.pss_force_max_amber, True,  'force_max', '%'),
                ('Force Avg', job.pss_force_avg_green, job.pss_force_avg_amber, True,  'force_avg', '%'),
                ('THD Max',   job.pss_thd_max_green,   job.pss_thd_max_amber,   False, 'thd_max',   '%'),
                ('THD Avg',   job.pss_thd_avg_green,   job.pss_thd_avg_amber,   False, 'thd_avg',   '%'),
            ]
            _psv = _map_views.get('pss')

            def _make_pss_unit_maps(src_df, unit_slug):
                unit_maps = []
                for col, g_thresh, a_thresh, higher, slug, sym in params:
                    try:
                        agg = src_df.groupby(['_line','_point'])[col].apply(
                            lambda x: x.abs().max() if not higher else x.mean()
                        ).reset_index(); agg.columns = ['_line','_point','val']
                        merged = agg.merge(sx_df[['_line','_point','X','Y']], on=['_line','_point'], how='inner')
                        if merged.empty: continue
                        lons, lats = t.transform(merged['X'].values, merged['Y'].values)
                        merged['lat'] = lats; merged['lon'] = lons
                        green_f, amber_f, red_f = [], [], []
                        for _, row in merged.iterrows():
                            v = row['val']
                            feat = {'type':'Feature','geometry':{'type':'Point','coordinates':[row['lon'],row['lat']]},'properties':{}}
                            if (v >= g_thresh if higher else v <= g_thresh): green_f.append(feat)
                            elif (v >= a_thresh if higher else v <= a_thresh): amber_f.append(feat)
                            else: red_f.append(feat)
                        # Build map directly (no zoomIn) so zoom matches QC tab
                        _buf = 100
                        _sw = t.transform(merged['X'].min() - _buf, merged['Y'].min() - _buf)
                        _ne = t.transform(merged['X'].max() + _buf, merged['Y'].max() + _buf)
                        pm = folium.Map(tiles=None, zoom_control=False, width='100%', height=200,
                                        scrollWheelZoom=False, dragging=False, doubleClickZoom=False)
                        folium.TileLayer(tiles=SATELLITE, attr='Esri').add_to(pm)
                        pm.fit_bounds([[_sw[1], _sw[0]], [_ne[1], _ne[0]]])
                        pm.get_root().header.add_child(folium.Element(HIDE_CSS))
                        _geojson_layer(pm, green_f, '#4caf50', 3)
                        _geojson_layer(pm, amber_f, '#ff9800', 3)
                        _geojson_layer(pm, red_f,   '#e63946', 3)
                        png = _save_map(pm, f'pss_{slug}_{unit_slug}',
                                       center=_psv[:2] if _psv else None,
                                       zoom=_psv[2] if _psv else None,
                                       width=450, height=200)
                        if png:
                            unit_maps.append({
                                'label': col,
                                'path': png,
                                'g': g_thresh,
                                'a': a_thresh,
                                'higher': higher,
                                'sym': sym,
                            })
                    except Exception:
                        pass
                return unit_maps

            if 'Unit ID' in pss_df.columns:
                all_units = sorted(pss_df['Unit ID'].astype(str).unique())
                if include_pss_units is not None:
                    all_units = [u for u in all_units if u in include_pss_units]
                for unit_val in all_units:
                    unit_df = pss_df[pss_df['Unit ID'].astype(str) == unit_val]
                    safe = _re.sub(r'[^a-zA-Z0-9]', '_', str(unit_val))
                    unit_maps = _make_pss_unit_maps(unit_df, safe)
                    if unit_maps:
                        pss_png_maps.append((str(unit_val), unit_maps))
            else:
                all_maps = _make_pss_unit_maps(pss_df, 'all')
                if all_maps:
                    pss_png_maps.append(('', all_maps))

        except Exception:
            pass

    # Chargeable hours
    from datetime import datetime as _dt2, date as _date2
    import re as _re2
    _type_pct = {at.name: at.chargeable_percentage for at in ActivityType.objects.all()}

    def _fmt_m(m):
        return f"{m // 60}h {m % 60:02d}m"

    _is_sv = report.report_type == 'supervisor'

    def _act_mins(a):
        if _is_sv:
            h = a.hours.strip() if a.hours else ''
            if not h:
                return 0
            try:
                return round(float(h) * 60)
            except ValueError:
                pass
            m = _re2.match(r'(\d+(?:\.\d+)?)\s*h(?:ours?)?\s*(?:(\d+)\s*m(?:in)?)?', h, _re2.IGNORECASE)
            if m:
                return round(float(m.group(1)) * 60) + int(m.group(2) or 0)
            m = _re2.match(r'(\d+):(\d+)', h)
            if m:
                return int(m.group(1)) * 60 + int(m.group(2))
            return 0
        s = _dt2.combine(_date2.today(), a.start_time)
        e = _dt2.combine(_date2.today(), a.end_time)
        return int((e - s).total_seconds() / 60)

    _daily = {}
    for a in report.activities.all():
        mins = _act_mins(a)
        pct = _type_pct.get(a.activity_type, 100)
        if a.activity_type not in _daily:
            _daily[a.activity_type] = {'total_mins': 0, 'chargeable_mins': 0, 'pct': pct}
        _daily[a.activity_type]['total_mins'] += mins
        _daily[a.activity_type]['chargeable_mins'] += round(mins * pct / 100)
    chargeable_rows = [{'type': k, 'pct': v['pct'], 'total': _fmt_m(v['total_mins']), 'chargeable': _fmt_m(v['chargeable_mins'])} for k, v in _daily.items()]
    total_chargeable = _fmt_m(sum(v['chargeable_mins'] for v in _daily.values()))

    _job = {}
    for a in Activity.objects.filter(report__job=report.job, report__report_type=report.report_type):
        mins = _act_mins(a)
        pct = _type_pct.get(a.activity_type, 100)
        if a.activity_type not in _job:
            _job[a.activity_type] = {'total_mins': 0, 'chargeable_mins': 0, 'pct': pct}
        _job[a.activity_type]['total_mins'] += mins
        _job[a.activity_type]['chargeable_mins'] += round(mins * pct / 100)
    job_chargeable_rows = [{'type': k, 'pct': v['pct'], 'total': _fmt_m(v['total_mins']), 'chargeable': _fmt_m(v['chargeable_mins'])} for k, v in _job.items()]
    job_total_chargeable = _fmt_m(sum(v['chargeable_mins'] for v in _job.values()))

    # Pie charts
    pie_charts_png = None
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as _plt

        _PIE_COLORS = [
            '#2196f3','#4caf50','#ff9800','#9c27b0','#00bcd4',
            '#ff5722','#795548','#f06292','#607d8b','#cddc39',
            '#e91e63','#009688','#ffc107','#3f51b5','#8bc34a'
        ]

        def _pie_data(activities):
            cat = {}
            for a in activities:
                m = _act_mins(a)
                cat[a.category] = cat.get(a.category, 0) + m
            return [(k, v) for k, v in cat.items() if v > 0]

        _daily_pie = _pie_data(report.activities.all())
        _job_pie   = _pie_data(Activity.objects.filter(report__job=report.job, report__report_type=report.report_type))

        if _daily_pie or _job_pie:
            # Build shared color map: same category → same color in both charts
            _seen, _cat_colors, _cidx = {}, {}, 0
            for _d in _daily_pie + _job_pie:
                if _d[0] not in _seen:
                    _cat_colors[_d[0]] = _PIE_COLORS[_cidx % len(_PIE_COLORS)]
                    _cidx += 1
                    _seen[_d[0]] = True

            _fig, _axes = _plt.subplots(1, 2, figsize=(11, 5.5))

            def _draw_pie(ax, data, title):
                if not data:
                    ax.text(0.5, 0.5, 'No data', ha='center', va='center', transform=ax.transAxes)
                    ax.axis('off')
                else:
                    labels = [d[0] for d in data]
                    values = [d[1] for d in data]
                    colors = [_cat_colors[l] for l in labels]
                    ax.pie(values, colors=colors, startangle=90)
                    _total = sum(values)
                    legend_labels = [
                        f'{l}  ({v // 60}h {v % 60:02d}m)  {v / _total * 100:.1f}%'
                        for l, v in zip(labels, values)
                    ]
                    ncols = 2 if len(legend_labels) > 5 else 1
                    ax.legend(legend_labels, loc='upper center', bbox_to_anchor=(0.5, -0.04),
                              ncol=ncols, fontsize=7.5, frameon=False)
                ax.set_title(title, fontsize=10, fontweight='bold', color='#1a3a5c', pad=10)

            _daily_title = 'Daily Activity Summary' if report.report_type == 'supervisor' else 'Daily Recording Summary'
            _draw_pie(_axes[0], _daily_pie, _daily_title)
            _draw_pie(_axes[1], _job_pie,   'Job Summary')
            _plt.subplots_adjust(top=0.92, bottom=0.28, left=0.04, right=0.96, wspace=0.3)
            _pie_path = os.path.join(maps_dir, f'pie_charts_{report.pk}.png')
            _plt.savefig(_pie_path, dpi=150, bbox_inches='tight', facecolor='white')
            _plt.close(_fig)
            if os.path.exists(_pie_path):
                pie_charts_png = _pie_path
    except Exception:
        pass

    # Shot interval chart PNG (for PDF export)
    shot_chart_png = None
    if report.include_shot_chart and shot_stats and shot_stats.get('shot_intervals'):
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as _splt
            _intervals = shot_stats['shot_intervals']
            _labels = [d['interval'] for d in _intervals]
            _counts = [d['shots'] for d in _intervals]
            _sfig, _sax = _splt.subplots(figsize=(10, 1.5))
            _sax.bar(_labels, _counts, color='#1a3a5c', width=0.8)
            _sax.set_ylabel('Shots', fontsize=9)
            _sax.tick_params(axis='x', labelsize=7, rotation=45)
            _sax.tick_params(axis='y', labelsize=8)
            _sax.yaxis.set_major_locator(__import__('matplotlib.ticker', fromlist=['MaxNLocator']).MaxNLocator(integer=True))
            _sax.set_facecolor('white')
            _sfig.patch.set_facecolor('white')
            _splt.tight_layout()
            _sc_path = os.path.join(maps_dir, f'shot_chart_{report.pk}.png')
            _sfig.savefig(_sc_path, dpi=150, bbox_inches='tight', facecolor='white')
            _splt.close(_sfig)
            if os.path.exists(_sc_path):
                shot_chart_png = _sc_path
        except Exception:
            pass

    # Logo
    import base64 as _b64
    _logo_path = os.path.join(settings.MEDIA_ROOT, 'velseis-logo.png')
    _logo_b64 = None
    if os.path.exists(_logo_path):
        with open(_logo_path, 'rb') as _lf:
            _logo_b64 = 'data:image/png;base64,' + _b64.b64encode(_lf.read()).decode()

    # Use same fallback as the main view: if this report has no personnel saved,
    # pull from the most recent previous report (same job) that has them.
    def _get_personnel(field):
        val = getattr(report, field)
        if val:
            return val.splitlines()
        prev = (DailyReport.objects.filter(job=report.job, date__lt=report.date)
                .exclude(**{field: ''}).order_by('-date')
                .values_list(field, flat=True).first())
        return prev.splitlines() if prev else []

    # Finish date estimates for export
    _finish_est = []
    _node_finish_est = []
    if report.job.finish_include_in_report:
        import math as _mth
        from datetime import timedelta as _td
        _dpw = max(1, min(7, report.job.finish_days_per_week))
        _shots_rem = (planned_shots - job_total_shots) if planned_shots is not None else None
        _nodes_rem = (planned_nodes - job_total_nodes) if planned_nodes is not None else None
        _daily_avg_shots = round(avg_shots_total / production_days) if production_days else None
        _daily_avg_nodes = round(avg_nodes_total / deployment_days) if deployment_days else None
        _last_date = report.job.reports.filter(date__lte=report.date).order_by('date').last()
        _last_date = _last_date.date if _last_date else report.date
        _first_report = report.job.reports.filter(date__lte=report.date).order_by('date').first()

        def _p2c(pd_): return _mth.ceil(pd_ * 7 / _dpw)
        def _fin(d, c): return d + _td(days=c)

        if _shots_rem is not None and _shots_rem > 0:
            if report.job.finish_show_linear and _daily_avg_shots:
                pd_ = _mth.ceil(_shots_rem / _daily_avg_shots)
                cd = _p2c(pd_)
                _finish_est.append({'label': 'Daily Avg', 'date': _fin(_last_date, cd), 'detail': f'{pd_} prod. days / {cd} cal. days'})
            if report.job.finish_show_calendar and job_total_shots and _first_report:
                el = (_last_date - _first_report.date).days + 1
                if el > 0:
                    cd = _mth.ceil(_shots_rem / (job_total_shots / el))
                    _finish_est.append({'label': 'Calendar Rate', 'date': _fin(_last_date, cd), 'detail': f'{job_total_shots/el:.1f} shots/cal. day → {cd} cal. days'})
            if report.job.finish_show_rolling:
                _win = max(1, report.job.finish_rolling_window)
                _inc = [r for r in _ctx_shot_rows if r['shots'] and r['shots'] > 0 and r['report'].include_in_avg]
                if _inc[-_win:]:
                    _ravg = sum(r['shots'] for r in _inc[-_win:]) / len(_inc[-_win:])
                    pd_ = _mth.ceil(_shots_rem / _ravg)
                    cd = _p2c(pd_)
                    _finish_est.append({'label': f'Rolling {_win}-day Avg', 'date': _fin(_last_date, cd), 'detail': f'{_ravg:.0f} shots/day → {pd_} prod. days / {cd} cal. days'})
        elif _shots_rem is not None and _shots_rem <= 0:
            _finish_est = [{'label': 'Complete', 'date': None, 'detail': 'All planned shots fired'}]

        if _nodes_rem is not None and _nodes_rem > 0:
            if report.job.finish_show_linear and _daily_avg_nodes:
                pd_ = _mth.ceil(_nodes_rem / _daily_avg_nodes)
                cd = _p2c(pd_)
                _node_finish_est.append({'label': 'Daily Avg', 'date': _fin(_last_date, cd), 'detail': f'{pd_} prod. days / {cd} cal. days'})
            if report.job.finish_show_calendar and job_total_nodes and _ctx_node_rows:
                try:
                    el = (_last_date - _ctx_node_rows[0]['date']).days + 1
                    if el > 0:
                        cd = _mth.ceil(_nodes_rem / (job_total_nodes / el))
                        _node_finish_est.append({'label': 'Calendar Rate', 'date': _fin(_last_date, cd), 'detail': f'{job_total_nodes/el:.1f} nodes/cal. day → {cd} cal. days'})
                except Exception:
                    pass
            if report.job.finish_show_rolling:
                _win = max(1, report.job.finish_rolling_window)
                _inc = [r for r in _ctx_node_rows if r['nodes'] and r['nodes'] > 0 and r['report'] and r['report'].include_in_avg]
                if _inc[-_win:]:
                    _ravg = sum(r['nodes'] for r in _inc[-_win:]) / len(_inc[-_win:])
                    pd_ = _mth.ceil(_nodes_rem / _ravg)
                    cd = _p2c(pd_)
                    _node_finish_est.append({'label': f'Rolling {_win}-day Avg', 'date': _fin(_last_date, cd), 'detail': f'{_ravg:.0f} nodes/day → {pd_} prod. days / {cd} cal. days'})
        elif _nodes_rem is not None and _nodes_rem <= 0:
            _node_finish_est = [{'label': 'Complete', 'date': None, 'detail': 'All planned nodes deployed'}]

    return {
        'report': report,
        'logo_b64': _logo_b64,
        'logo_path': _logo_path if os.path.exists(_logo_path) else None,
        'pie_charts_png': pie_charts_png,
        'saved_observers': _get_personnel('observers'),
        'saved_operators': _get_personnel('operators'),
        'shot_stats': shot_stats,
        'rx_stats': rx_stats,
        'job_total_shots': job_total_shots or None,
        'job_total_nodes': job_total_nodes or None,
        'shots_remaining': (planned_shots - job_total_shots) if planned_shots is not None else None,
        'nodes_remaining': (planned_nodes - job_total_nodes) if planned_nodes is not None else None,
        'daily_avg_shots': round(avg_shots_total / production_days) if production_days else None,
        'daily_avg_nodes': round(avg_nodes_total / deployment_days) if deployment_days else None,
        'activities': report.activities.all(),
        'deployment_png': deployment_png,
        'active_patch_png': active_patch_png,
        'active_patch_stats': active_patch_stats,
        'progress_png': progress_png,
        'shot_chart_png': shot_chart_png,
        'pss_png_maps': pss_png_maps,
        'chargeable_rows': chargeable_rows,
        'total_chargeable': total_chargeable,
        'job_chargeable_rows': job_chargeable_rows,
        'job_total_chargeable': job_total_chargeable,
        'finish_estimates': _finish_est,
        'node_finish_estimates': _node_finish_est,
        'photos': report.photos.all(),
    }


def report_pdf(request, pk):
    report = get_object_or_404(DailyReport, pk=pk)
    if 'pss_units' in request.GET:
        include_pss_units = [u.strip() for u in request.GET['pss_units'].split(',') if u.strip()]
    else:
        include_pss_units = None
    ctx = _build_report_ctx(report, include_pss_units=include_pss_units, map_views=_parse_map_views(request))
    # Convert absolute image paths to file:// URIs so Playwright can load them
    from pathlib import Path
    for key in ('progress_png', 'deployment_png', 'active_patch_png', 'pie_charts_png', 'shot_chart_png'):
        if ctx.get(key):
            ctx[key] = Path(ctx[key]).as_uri()
    if ctx.get('pss_png_maps'):
        ctx['pss_png_maps'] = [
            (label, [dict(m, path=Path(m['path']).as_uri()) for m in maps])
            for label, maps in ctx['pss_png_maps']
        ]
    html = render_to_string('reports/report_pdf.html', ctx, request=request)
    # Write to a temp file so Playwright can load it via file://
    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            page.wait_for_load_state('networkidle')
            pdf_bytes = page.pdf(
                format='A4',
                print_background=True,
                display_header_footer=True,
                header_template='<span></span>',
                footer_template=(
                    '<div style="width:100%; font-family:Helvetica,Arial,sans-serif; '
                    'font-size:8px; color:#888; text-align:center;">'
                    'Page <span class="pageNumber"></span> of <span class="totalPages"></span>'
                    '</div>'
                ),
                margin={'top': '1.5cm', 'bottom': '1.8cm', 'left': '1.5cm', 'right': '1.5cm'},
            )
            browser.close()
    finally:
        os.unlink(tmp.name)
    if report.report_type == 'supervisor':
        _tpl = report.job.supervisor_filename_template or 'Supervisors Daily Report-{date}'
    else:
        _tpl = report.job.export_filename_template or 'Daily Production Report-{date}'
    filename = _tpl.format(
        date=report.date.strftime('%d-%m-%Y'),
        job_number=report.job.job_number,
        client=report.job.client,
        project=report.job.project_name,
    ) + '.pdf'
    if report.job.export_save_to_disk and report.job.export_save_path:
        try:
            save_dir = report.job.export_save_path
            os.makedirs(save_dir, exist_ok=True)
            with open(os.path.join(save_dir, filename), 'wb') as _f:
                _f.write(pdf_bytes)
        except Exception:
            pass
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    if request.GET.get('inline'):
        response['Content-Disposition'] = f'inline; filename="{filename}"'
    else:
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def report_word(request, pk):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1a, 0x3a, 0x5c)

    def heading(doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1a3a5c')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_page_break(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_table_row(table, values, header=False, shaded=False):
        row = table.add_row()
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = str(val) if val is not None else '—'
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val) if val is not None else '—')
            run.font.size = Pt(9)
            if header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, '1a3a5c')
            elif shaded:
                set_cell_bg(cell, 'f5f5f5')
        return row

    report = get_object_or_404(DailyReport, pk=pk)
    if 'pss_units' in request.GET:
        include_pss_units = [u.strip() for u in request.GET['pss_units'].split(',') if u.strip()]
    else:
        include_pss_units = None
    ctx = _build_report_ctx(report, include_pss_units=include_pss_units, map_views=_parse_map_views(request))
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Logo + Title block
    if ctx['logo_path']:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_p.add_run().add_picture(ctx['logo_path'], height=Cm(1.44))

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _report_title = "SUPERVISOR DAILY REPORT" if report.report_type == 'supervisor' else "RECORDING CREW DAILY PRODUCTION REPORT"
    title_run = title_p.add_run(_report_title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.add_run(f"{report.job.client} — {report.job.project_name}").font.size = Pt(10)

    date_p = doc.add_paragraph()
    dr = date_p.add_run(f"{report.date.strftime('%d %B %Y')}   |   Job {report.job.job_number}")
    dr.font.size = Pt(10)
    dr.bold = True

    doc.add_paragraph()

    if report.report_type != 'supervisor':
        # Job Details
        heading(doc, 'Job Details')
        jt = doc.add_table(rows=0, cols=2)
        jt.style = 'Table Grid'
        rows_data = [
            ('Recording System', report.job.recording_system),
            ('Source Type', report.job.source_type or '—'),
        ]
        if report.job.rx_interval:
            rows_data.append(('Rx Interval', f'{report.job.rx_interval} m'))
        if report.job.sx_interval:
            rows_data.append(('Sx Interval', f'{report.job.sx_interval} m'))
        for i, (label, val) in enumerate(rows_data):
            add_table_row(jt, [label, val], shaded=(i % 2 == 1))

        # Personnel
        heading(doc, 'Personnel')
        pt = doc.add_table(rows=0, cols=2)
        pt.style = 'Table Grid'
        observers = ', '.join(ctx['saved_observers']) or '—'
        operators = ', '.join(ctx['saved_operators']) or '—'
        add_table_row(pt, ['Observer(s)', observers])
        add_table_row(pt, ['Vibe Operator(s)', operators], shaded=True)

        # Production Summary
        ss = ctx['shot_stats']
        if ss and ss.get('production'):
            heading(doc, 'Production Summary')
            st = doc.add_table(rows=0, cols=4)
            st.style = 'Table Grid'
            add_table_row(st, ['VP Today', 'Daily Avg', 'Total Job Shots', 'Shots Remaining'], header=True)
            add_table_row(st, [
                ss['production'],
                ctx['daily_avg_shots'] or '—',
                ctx['job_total_shots'] or '—',
                ctx['shots_remaining'] if ctx['shots_remaining'] is not None else '—',
            ])
            if ss.get('line_stats'):
                doc.add_paragraph()
                lt = doc.add_table(rows=0, cols=4)
                lt.style = 'Table Grid'
                add_table_row(lt, ['Line', 'Start', 'End', 'Shots'], header=True)
                for i, ls in enumerate(ss['line_stats']):
                    add_table_row(lt, [ls['line'], ls['start'], ls['end'], ls['shots']], shaded=(i % 2 == 1))
                total_row = lt.add_row()
                total_row.cells[0].text = 'Total'
                total_row.cells[0].merge(total_row.cells[2])
                total_row.cells[3].text = str(ss['production'])
                for cell in total_row.cells:
                    set_cell_bg(cell, 'ececec')
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)

        # Shot Progress Map
        if ctx['progress_png'] and os.path.exists(ctx['progress_png']):
            heading(doc, 'Shot Progress')
            doc.add_picture(ctx['progress_png'], width=Inches(6))

        # Node Deployment
        rx = ctx['rx_stats']
        if rx:
            heading(doc, 'Node Deployment')
            nt = doc.add_table(rows=0, cols=4)
            nt.style = 'Table Grid'
            add_table_row(nt, ['Nodes Today', 'Daily Avg', 'Total Job Nodes', 'Nodes Remaining'], header=True)
            add_table_row(nt, [
                rx['total'],
                ctx['daily_avg_nodes'] or '—',
                ctx['job_total_nodes'] or '—',
                ctx['nodes_remaining'] if ctx['nodes_remaining'] is not None else '—',
            ])
            if rx.get('line_stats'):
                doc.add_paragraph()
                nlt = doc.add_table(rows=0, cols=4)
                nlt.style = 'Table Grid'
                add_table_row(nlt, ['Line', 'Start', 'End', 'Nodes'], header=True)
                for i, ls in enumerate(rx['line_stats']):
                    add_table_row(nlt, [ls['line'], ls['start'], ls['end'], ls['nodes']], shaded=(i % 2 == 1))
                total_row = nlt.add_row()
                total_row.cells[0].text = 'Total'
                total_row.cells[0].merge(total_row.cells[2])
                total_row.cells[3].text = str(rx['total'])
                for cell in total_row.cells:
                    set_cell_bg(cell, 'ececec')
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)

        # Node Deployment Progress Map (one or the other, never both)
        if ctx['deployment_png'] and os.path.exists(ctx['deployment_png']):
            heading(doc, 'Node Deployment Progress')
            doc.add_picture(ctx['deployment_png'], width=Inches(6))
        elif ctx.get('active_patch_png') and os.path.exists(ctx['active_patch_png']):
            heading(doc, 'Node Deployment Progress — Active Patch')
            doc.add_picture(ctx['active_patch_png'], width=Inches(6))
            aps = ctx.get('active_patch_stats')
            if aps:
                at = doc.add_table(rows=2, cols=3)
                at.style = 'Table Grid'
                hdr = at.rows[0].cells
                for i, h in enumerate(['Rear of Patch', 'Front of Patch', 'Nodes in Ground']):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True
                row = at.rows[1].cells
                row[0].text = f"Line {aps['last_line']} / Sta {aps['last_station']}"
                row[1].text = f"Line {aps['front_line']} / Sta {aps['front_station']}" if aps['front_line'] else '—'
                row[2].text = str(aps['in_ground'])

    # Activity Details
    heading(doc, 'Activity Details')
    activities = list(ctx['activities'])
    if activities:
        at = doc.add_table(rows=0, cols=6)
        at.style = 'Table Grid'
        if report.report_type == 'supervisor':
            add_table_row(at, ['Contractor', 'Name', 'Job Title', 'Hours', 'Category', 'Type'], header=True)
            for i, a in enumerate(activities):
                add_table_row(at, [
                    a.details,
                    a.notes,
                    a.job_title,
                    a.hours,
                    a.category,
                    a.activity_type,
                ], shaded=(i % 2 == 1))
        else:
            add_table_row(at, ['Start', 'End', 'Duration', 'Category', 'Type', 'Details'], header=True)
            for i, a in enumerate(activities):
                add_table_row(at, [
                    a.start_time.strftime('%H:%M'),
                    a.end_time.strftime('%H:%M'),
                    a.duration,
                    a.category,
                    a.activity_type,
                    a.details,
                ], shaded=(i % 2 == 1))
    else:
        doc.add_paragraph('No activities recorded.').runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Diary
    if report.diary:
        heading(doc, 'Diary / Notes')
        diary_text = report.diary.replace('\r\n', '\n').replace('\r', '\n')
        doc.add_paragraph(diary_text).runs[0].font.size = Pt(9)

    # Chargeable Hours — Daily
    if ctx['chargeable_rows']:
        heading(doc, f"Chargeable Hours — {report.date.strftime('%d %b %Y')}")
        cht = doc.add_table(rows=0, cols=4)
        cht.style = 'Table Grid'
        add_table_row(cht, ['Activity Type', 'Chargeable %', 'Total Hours', 'Chargeable Hours'], header=True)
        for i, row in enumerate(ctx['chargeable_rows']):
            add_table_row(cht, [row['type'], f"{row['pct']}%", row['total'], row['chargeable']], shaded=(i % 2 == 1))
        tr = cht.add_row()
        tr.cells[0].text = 'Total Chargeable'
        tr.cells[0].merge(tr.cells[2])
        tr.cells[3].text = ctx['total_chargeable']
        for cell in tr.cells:
            set_cell_bg(cell, 'ececec')
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

    # Chargeable Hours — Job Total
    if ctx['job_chargeable_rows']:
        heading(doc, 'Chargeable Hours — Job Total')
        jcht = doc.add_table(rows=0, cols=4)
        jcht.style = 'Table Grid'
        add_table_row(jcht, ['Activity Type', 'Chargeable %', 'Total Hours', 'Chargeable Hours'], header=True)
        for i, row in enumerate(ctx['job_chargeable_rows']):
            add_table_row(jcht, [row['type'], f"{row['pct']}%", row['total'], row['chargeable']], shaded=(i % 2 == 1))
        tr = jcht.add_row()
        tr.cells[0].text = 'Total Chargeable'
        tr.cells[0].merge(tr.cells[2])
        tr.cells[3].text = ctx['job_total_chargeable']
        for cell in tr.cells:
            set_cell_bg(cell, 'ececec')
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

    # Pie Charts
    if ctx['pie_charts_png'] and os.path.exists(ctx['pie_charts_png']):
        _pie_title = 'Activity Summary' if report.report_type == 'supervisor' else 'Recording Summary'
        heading(doc, _pie_title)
        doc.add_picture(ctx['pie_charts_png'], width=Inches(6))

    # PSS QC Maps — new page, per-unit sections, 2-up layout
    if ctx['pss_png_maps'] and report.report_type != 'supervisor':
        add_page_break(doc)
        heading(doc, 'PSS QC Maps')
        for unit_label, unit_maps in ctx['pss_png_maps']:
            if unit_label:
                p = doc.add_paragraph()
                r = p.add_run(f'Unit {unit_label}')
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = NAVY
                p.paragraph_format.space_before = Pt(8)
            valid_maps = [m for m in unit_maps if m['path'] and os.path.exists(m['path'])]
            for i in range(0, len(valid_maps), 2):
                pair = valid_maps[i:i + 2]
                tbl = doc.add_table(rows=1, cols=len(pair))
                tbl.style = 'Table Grid'
                for j, m in enumerate(pair):
                    cell = tbl.rows[0].cells[j]
                    cell.paragraphs[0].add_run(m['label']).bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    img_para = cell.add_paragraph()
                    img_para.add_run().add_picture(m['path'], width=Inches(2.9 if len(pair) == 2 else 5.8))

    # Photos
    photos = list(ctx.get('photos', []))
    if photos:
        add_page_break(doc)
        heading(doc, 'Photos')
        cols = report.photo_columns
        col_width = (Inches(6.5) / cols)
        if cols == 1:
            for photo in photos:
                try:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run()
                    run.add_picture(photo.image.path, width=Inches(6.5))
                except Exception:
                    pass
                if photo.caption:
                    cap = doc.add_paragraph(photo.caption)
                    cap.runs[0].font.size = Pt(10)
                    cap.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(10)
        else:
            t = doc.add_table(rows=0, cols=2)
            row_cells = None
            for i, photo in enumerate(photos):
                if i % 2 == 0:
                    row_cells = t.add_row().cells
                cell = row_cells[i % 2]
                try:
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    run.add_picture(photo.image.path, width=col_width - Inches(0.1))
                except Exception:
                    pass
                if photo.caption:
                    cap_p = cell.add_paragraph(photo.caption)
                    cap_p.runs[0].font.size = Pt(10)
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    if report.report_type == 'supervisor':
        _tpl = report.job.supervisor_filename_template or 'Supervisors Daily Report-{date}'
    else:
        _tpl = report.job.export_filename_template or 'Daily Production Report-{date}'
    filename = _tpl.format(
        date=report.date.strftime('%d-%m-%Y'),
        job_number=report.job.job_number,
        client=report.job.client,
        project=report.job.project_name,
    ) + '.docx'
    docx_bytes = buf.read()
    if report.job.export_save_to_disk and report.job.export_save_path:
        try:
            save_dir = report.job.export_save_path
            os.makedirs(save_dir, exist_ok=True)
            with open(os.path.join(save_dir, filename), 'wb') as _f:
                _f.write(docx_bytes)
        except Exception:
            pass
    response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def report_file_upload(request, pk):
    report = get_object_or_404(DailyReport, pk=pk)
    if request.method == 'POST':
        upload_method = request.POST.get('upload_method')

        if upload_method == 'zip':
            f = request.FILES.get('zip_file')
            if f:
                import zipfile as _zipfile
                import re as _re
                from django.core.files.base import ContentFile as _ContentFile
                _FILE_PATTERNS = [
                    (_re.compile(r'/Reports/FinalCOG_\d{4}_',              _re.IGNORECASE), 'cog'),
                    (_re.compile(r'/Reports/ObserverLog_Detailed_\d{4}_',  _re.IGNORECASE), 'obslog'),
                    (_re.compile(r'/Reports/PSS_\d{4}_',                   _re.IGNORECASE), 'pss'),
                ]
                try:
                    with _zipfile.ZipFile(f) as zf:
                        for entry in zf.namelist():
                            norm = '/' + entry.replace('\\', '/')
                            fname = norm.split('/')[-1]
                            if not fname:
                                continue
                            for pattern, ftype in _FILE_PATTERNS:
                                if pattern.search(norm):
                                    data = zf.read(entry)
                                    ReportFile.objects.create(
                                        report=report,
                                        file_type=ftype,
                                        file=_ContentFile(data, name=fname),
                                        original_name=fname,
                                    )
                                    break
                except Exception:
                    pass

        elif upload_method == 'folder':
            file_map = [
                ('cog_file', 'cog'),
                ('obs_file', 'obslog'),
                ('pss_file', 'pss'),
            ]
            for field, ftype in file_map:
                f = request.FILES.get(field)
                if f:
                    ReportFile.objects.create(report=report, file_type=ftype, file=f, original_name=f.name)

        elif upload_method == 'individual':
            for ftype in ['obslog', 'pss', 'cog']:
                for f in request.FILES.getlist(ftype):
                    ReportFile.objects.create(report=report, file_type=ftype, file=f, original_name=f.name)

        rx_file = request.FILES.get('rx_deployment')
        if rx_file:
            ReportFile.objects.create(report=report, file_type='rx_deployment', file=rx_file, original_name=rx_file.name)

    return redirect(f'/reports/{pk}/?tab=files')


def report_file_delete(request, pk):
    rf = get_object_or_404(ReportFile, pk=pk)
    report_pk = rf.report.pk
    rf.file.delete()
    rf.delete()
    return redirect(f'/reports/{report_pk}/?tab=files')


def report_active_patch_save(request, pk):
    if request.method != 'POST':
        return redirect('report_detail', pk=pk)
    report = get_object_or_404(DailyReport, pk=pk)
    line_raw = request.POST.get('last_line_in_ground', '').strip()
    station_raw = request.POST.get('last_station_in_ground', '').strip()
    try:
        report.last_line_in_ground = float(line_raw) if line_raw else None
    except ValueError:
        report.last_line_in_ground = None
    try:
        report.last_station_in_ground = float(station_raw) if station_raw else None
    except ValueError:
        report.last_station_in_ground = None
    report.save(update_fields=['last_line_in_ground', 'last_station_in_ground'])
    return redirect(f'/reports/{pk}/?tab=files')


def _count_sps_records(file_obj):
    """Count data records (non-H lines) in an SPS 2.1 file."""
    count = 0
    try:
        for line in file_obj:
            decoded = line.decode('utf-8', errors='ignore').rstrip('\r\n')
            if decoded and not decoded.startswith('H'):
                count += 1
        file_obj.seek(0)
    except Exception:
        pass
    return count


def job_survey_position_upload(request, job_pk):
    """Upload or replace the job-level RPS or SPS survey position file, and save metadata."""
    if request.method != 'POST':
        return redirect('job_detail', pk=job_pk)
    job = get_object_or_404(Job, pk=job_pk)
    file_type = request.POST.get('file_type')  # 'rps' or 'sps'
    uploaded = request.FILES.get('file')
    update_fields = []

    if file_type in ('rps', 'sps'):
        file_field = f'{file_type}_file'
        count_field = f'{file_type}_count'
        is_final_field = f'{file_type}_is_final'
        estimated_field = 'estimated_rx_count' if file_type == 'rps' else 'estimated_sx_count'

        if uploaded:
            old = getattr(job, file_field)
            if old:
                old.delete(save=False)
            setattr(job, file_field, uploaded)
            setattr(job, count_field, _count_sps_records(uploaded))
            update_fields += [file_field, count_field]

        is_final = request.POST.get('is_final') == '1'
        setattr(job, is_final_field, is_final)
        update_fields.append(is_final_field)

        estimated_raw = request.POST.get('estimated_count', '').strip()
        if not is_final:
            setattr(job, estimated_field, int(estimated_raw) if estimated_raw.isdigit() else None)
        else:
            setattr(job, estimated_field, None)
        update_fields.append(estimated_field)

        job.save(update_fields=update_fields)

    report_pk = request.POST.get('report_pk')
    if report_pk:
        return redirect(f'/reports/{report_pk}/?tab=files')
    return redirect('job_detail', pk=job_pk)


def job_survey_epsg_save(request, job_pk):
    """Save survey coordinate system settings (EPSG override and timezone) for a job."""
    if request.method != 'POST':
        return redirect('job_detail', pk=job_pk)
    job = get_object_or_404(Job, pk=job_pk)
    update_fields = []
    epsg = request.POST.get('survey_epsg', '').strip()
    job.survey_epsg = epsg if epsg else None
    update_fields.append('survey_epsg')
    tz = request.POST.get('timezone', '').strip()
    if tz:
        job.timezone = tz
        update_fields.append('timezone')
    if tz == 'custom':
        try:
            offset = request.POST.get('utc_offset_custom', '').strip()
            job.utc_offset_custom = float(offset) if offset else None
            update_fields.append('utc_offset_custom')
        except (ValueError, TypeError):
            pass
    job.save(update_fields=update_fields)
    report_pk = request.POST.get('report_pk')
    if report_pk:
        return redirect(f'/reports/{report_pk}/?tab=files')
    return redirect('job_detail', pk=job_pk)


def pss_preset_save(request):
    if request.method == 'POST':
        name = request.POST.get('preset_name', '').strip()
        report_pk = request.POST.get('report_pk')
        if name and report_pk:
            report = get_object_or_404(DailyReport, pk=report_pk)
            job = report.job
            fields = [
                'pss_force_avg_green', 'pss_force_avg_amber',
                'pss_force_max_green', 'pss_force_max_amber',
                'pss_phase_avg_green', 'pss_phase_avg_amber',
                'pss_phase_max_green', 'pss_phase_max_amber',
                'pss_thd_avg_green',  'pss_thd_avg_amber',
                'pss_thd_max_green',  'pss_thd_max_amber',
            ]
            preset, _ = PSSQCPreset.objects.get_or_create(name=name)
            for f in fields:
                setattr(preset, f, getattr(job, f))
            preset.save()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def pss_preset_delete(request, pk):
    PSSQCPreset.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def diary_template_add(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        body = request.POST.get('body', '').strip()
        if name and body:
            DiaryTemplate.objects.create(name=name, body=body)
    return redirect(request.META.get('HTTP_REFERER', '/'))


def diary_template_update(request, pk):
    if request.method == 'POST':
        t = get_object_or_404(DiaryTemplate, pk=pk)
        name = request.POST.get('name', '').strip()
        body = request.POST.get('body', '').strip()
        if name and body:
            t.name = name
            t.body = body
            t.save()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def diary_template_delete(request, pk):
    if request.method == 'POST':
        DiaryTemplate.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def personnel_schedule(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)

    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'add_status':
            name = request.POST.get('name', '').strip()
            if name:
                count = ScheduleStatus.objects.count()
                ScheduleStatus.objects.get_or_create(
                    name=name,
                    defaults={
                        'color': request.POST.get('color', '#90caf9'),
                        'pattern': request.POST.get('pattern', 'solid'),
                        'order': count,
                    }
                )
        elif action == 'delete_status':
            ScheduleStatus.objects.filter(pk=request.POST.get('pk')).delete()
        return redirect('personnel_schedule', job_pk=job_pk)

    job_personnel = list(job.personnel.order_by('sort_order', 'name'))
    today = date.today()
    try:
        start = date.fromisoformat(request.GET.get('start', ''))
    except ValueError:
        start = today
    try:
        end = date.fromisoformat(request.GET.get('end', ''))
    except ValueError:
        end = today + timedelta(days=27)
    if end < start:
        end = start + timedelta(days=27)
    delta = (end - start).days + 1
    schedule_days = [start + timedelta(days=i) for i in range(min(delta, 180))]
    existing = {
        (e.job_personnel_id, e.date): (e.status, e.note)
        for e in PersonnelScheduleEntry.objects.filter(
            job_personnel__job=job,
            date__gte=schedule_days[0],
            date__lte=schedule_days[-1],
        )
    }
    statuses = list(ScheduleStatus.objects.all())
    schedule_rows = [
        {
            'person': p,
            'is_separator': p.is_separator,
            'cells': [] if p.is_separator else [(d, existing.get((p.pk, d), ('', ''))[0], existing.get((p.pk, d), ('', ''))[1]) for d in schedule_days],
        }
        for p in job_personnel
    ]
    on_job_names = set(p.name for p in job_personnel if not p.is_separator)
    all_personnel_names = list(Personnel.objects.exclude(name__in=on_job_names).values_list('name', flat=True).order_by('name'))
    return render(request, 'reports/personnel_schedule.html', {
        'job': job,
        'schedule_days': schedule_days,
        'schedule_rows': schedule_rows,
        'today': today,
        'statuses': statuses,
        'statuses_json': json.dumps([{'name': s.name, 'color': s.color, 'pattern': s.pattern} for s in statuses]),
        'range_start': start.isoformat(),
        'range_end': end.isoformat(),
        'all_personnel_names': all_personnel_names,
    })


def personnel_schedule_set(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        new_status = data.get('status', '')
        for cell in data.get('cells', []):
            jp = get_object_or_404(JobPersonnel, pk=cell['jp_pk'], job_id=job_pk)
            d = date.fromisoformat(cell['date'])
            entry, _ = PersonnelScheduleEntry.objects.get_or_create(job_personnel=jp, date=d)
            entry.status = new_status
            entry.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_set_note(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        jp = get_object_or_404(JobPersonnel, pk=data.get('jp_pk'), job_id=job_pk)
        d = date.fromisoformat(data['date'])
        note = data.get('note', '').strip()
        entry, _ = PersonnelScheduleEntry.objects.get_or_create(job_personnel=jp, date=d)
        entry.note = note
        entry.save()
        return JsonResponse({'ok': True, 'note': entry.note})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_toggle(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        jp = get_object_or_404(JobPersonnel, pk=data['jp_pk'], job_id=job_pk)
        d = date.fromisoformat(data['date'])
        new_status = data.get('status', '')
        entry, _ = PersonnelScheduleEntry.objects.get_or_create(job_personnel=jp, date=d)
        entry.status = '' if entry.status == new_status else new_status
        entry.save()
        return JsonResponse({'status': entry.status})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_reorder(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        for i, pk in enumerate(data.get('order', [])):
            JobPersonnel.objects.filter(pk=pk, job_id=job_pk).update(sort_order=i)
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_add_person(request, job_pk):
    if request.method == 'POST':
        job = get_object_or_404(Job, pk=job_pk)
        data = json.loads(request.body)
        is_separator = data.get('is_separator', False)
        max_order = job.personnel.count()
        if is_separator:
            jp = JobPersonnel.objects.create(job=job, name='', sort_order=max_order, is_separator=True)
            return JsonResponse({'pk': jp.pk})
        name = data.get('name', '').strip()
        if not name:
            return JsonResponse({'error': 'Name required'}, status=400)
        if job.personnel.filter(name__iexact=name, is_separator=False).exists():
            return JsonResponse({'error': 'Already on job'}, status=400)
        jp = JobPersonnel.objects.create(job=job, name=name, sort_order=max_order)
        return JsonResponse({'pk': jp.pk, 'name': jp.name})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_remove_person(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        jp = get_object_or_404(JobPersonnel, pk=data.get('jp_pk'), job_id=job_pk)
        jp.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'POST required'}, status=400)


def personnel_schedule_pdf(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    today = date.today()
    try:
        start = date.fromisoformat(request.GET.get('start', ''))
    except ValueError:
        start = today
    try:
        end = date.fromisoformat(request.GET.get('end', ''))
    except ValueError:
        end = today + timedelta(days=27)
    if end < start:
        end = start + timedelta(days=27)
    delta = (end - start).days + 1
    schedule_days = [start + timedelta(days=i) for i in range(min(delta, 180))]
    existing = {
        (e.job_personnel_id, e.date): (e.status, e.note)
        for e in PersonnelScheduleEntry.objects.filter(
            job_personnel__job=job,
            date__gte=schedule_days[0],
            date__lte=schedule_days[-1],
        )
    }
    statuses = list(ScheduleStatus.objects.all())
    schedule_rows = [
        {
            'person': p,
            'is_separator': p.is_separator,
            'cells': [] if p.is_separator else [(d, existing.get((p.pk, d), ('', ''))[0], existing.get((p.pk, d), ('', ''))[1]) for d in schedule_days],
        }
        for p in job.personnel.all()
    ]
    html = render_to_string('reports/personnel_schedule_pdf.html', {
        'job': job,
        'schedule_days': schedule_days,
        'schedule_rows': schedule_rows,
        'statuses': statuses,
        'range_start': start,
        'range_end': end,
        'today': today,
    })
    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        pdf_bytes = pg.pdf(
            format='A3', landscape=True, print_background=True,
            margin={'top': '1cm', 'bottom': '1cm', 'left': '1cm', 'right': '1cm'},
        )
        browser.close()
    date_str = today.strftime('%Y%m%d')
    job_slug = re.sub(r'[^\w]+', '_', f"{job.job_number}_{job.project_name}").strip('_')
    filename = f'{job_slug}_schedule_{date_str}.pdf'
    disposition = 'inline' if request.GET.get('inline') else 'attachment'
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'{disposition}; filename="{filename}"'
    return response


def job_personnel_page(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    job_personnel = job.personnel.all()
    on_job_names = set(job_personnel.values_list('name', flat=True))
    requirements = {r.skill_id: r.count for r in job.skill_requirements.all()}
    name_to_jp = {jp.name: jp for jp in job_personnel.prefetch_related('roles')}
    personnel_phones = {p.name: p.phone for p in Personnel.objects.all()}
    job_personnel_list = list(job_personnel)
    for jp in job_personnel_list:
        jp.phone = personnel_phones.get(jp.name, '')
    skills_data = [
            {
                'pk': s.pk,
                'name': s.name,
                'required': requirements.get(s.pk, 0),
                'personnel': [
                    {
                        'name': p.name,
                        'jp_pk': name_to_jp[p.name].pk,
                        'role_pk': {r.role: r.pk for r in name_to_jp[p.name].roles.all()}.get(s.name),
                    }
                    for p in s.personnel.all() if p.name in on_job_names
                ],
                'assigned_count': sum(
                    1 for p in s.personnel.all()
                    if p.name in on_job_names
                    and s.name in {r.role for r in name_to_jp[p.name].roles.all()}
                ),
            }
            for s in Skill.objects.prefetch_related('personnel').all()
        ]
    if request.headers.get('X-Partial') == 'skills':
        from django.template.loader import render_to_string
        html = render_to_string('reports/job_personnel_skills_partial.html', {'skills': skills_data}, request=request)
        return HttpResponse(html)
    return render(request, 'reports/job_personnel.html', {
        'job': job,
        'personnel': job_personnel_list,
        'all_personnel': Personnel.objects.all(),
        'on_job_names': on_job_names,
        'personnel_phones': personnel_phones,
        'next_url': request.GET.get('next', ''),
        'skills': skills_data,
    })


def personnel_word(request, job_pk):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    NAVY = RGBColor(0x1a, 0x3a, 0x5c)

    job = get_object_or_404(Job, pk=job_pk)
    personnel_phones = {p.name: p.phone for p in Personnel.objects.all()}
    job_personnel = list(job.personnel.filter(is_separator=False).order_by('name').prefetch_related('roles'))
    for jp in job_personnel:
        jp.phone = personnel_phones.get(jp.name, '')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f'Personnel — {job.job_number}')
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = NAVY
    if job.project_name:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub_p.add_run(job.project_name)
        sr.font.size = Pt(10)
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    hdr_cells = tbl.rows[0].cells
    for i, txt in enumerate(['Personnel', 'Phone', 'Job Role']):
        hdr_cells[i].text = txt
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(hdr_cells[i], '1a3a5c')

    for jp in job_personnel:
        row = tbl.add_row().cells
        row[0].text = jp.name
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = jp.phone or ''
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
        roles = ', '.join(r.role for r in jp.roles.all())
        row[2].text = roles
        row[2].paragraphs[0].runs[0].font.size = Pt(10)

    # Column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    import re
    job_slug = re.sub(r'[^\w]+', '_', f"{job.job_number}_{job.project_name}").strip('_')
    filename = f'{job_slug}_personnel.docx'
    response = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def job_personnel_clear_roles(request, pk):
    p = get_object_or_404(JobPersonnel, pk=pk)
    if request.method == 'POST':
        p.roles.all().delete()
    next_url = request.POST.get('next', '')
    return redirect(next_url if next_url else reverse('job_personnel_page', kwargs={'job_pk': p.job.pk}))


def job_personnel_role_add(request, pk):
    p = get_object_or_404(JobPersonnel, pk=pk)
    if request.method == 'POST':
        role = request.POST.get('role', '').strip()
        if role:
            r, _ = JobPersonnelRole.objects.get_or_create(job_personnel=p, role=role)
            if request.headers.get('X-Fetch') == '1':
                return JsonResponse({'ok': True, 'role_pk': r.pk, 'role': r.role, 'jp_pk': p.pk, 'name': p.name})
    return redirect('job_personnel_page', job_pk=p.job.pk)


def job_personnel_role_remove(request, pk):
    r = get_object_or_404(JobPersonnelRole, pk=pk)
    job_pk = r.job_personnel.job.pk
    if request.method == 'POST':
        data = {'ok': True, 'role': r.role, 'jp_pk': r.job_personnel.pk, 'name': r.job_personnel.name}
        r.delete()
        if request.headers.get('X-Fetch') == '1':
            return JsonResponse(data)
    return redirect('job_personnel_page', job_pk=job_pk)


def job_skill_requirement_update(request, job_pk):
    if request.method == 'POST':
        job = get_object_or_404(Job, pk=job_pk)
        skill_pk = request.POST.get('skill_pk')
        count = max(0, int(request.POST.get('count', 0) or 0))
        skill = get_object_or_404(Skill, pk=skill_pk)
        JobSkillRequirement.objects.update_or_create(
            job=job, skill=skill, defaults={'count': count}
        )
        if request.headers.get('X-Fetch') == '1':
            assigned = JobPersonnelRole.objects.filter(job_personnel__job=job, role=skill.name).count()
            return JsonResponse({'ok': True, 'skill_pk': skill.pk, 'count': count, 'assigned': assigned})
    return redirect('job_personnel_page', job_pk=job_pk)


def job_personnel_add(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            jp = JobPersonnel.objects.create(
                job=job,
                name=name,
                role=request.POST.get('role', '').strip(),
                notes=request.POST.get('notes', '').strip(),
            )
            if request.headers.get('X-Fetch') == '1':
                return JsonResponse({'ok': True, 'jp_pk': jp.pk, 'name': jp.name})
    next_url = request.POST.get('next') or request.GET.get('next')
    return redirect(next_url if next_url else reverse('job_personnel_page', kwargs={'job_pk': job_pk}))


def job_personnel_delete(request, pk):
    p = get_object_or_404(JobPersonnel, pk=pk)
    job_pk = p.job.pk
    if request.method == 'POST':
        data = {'ok': True, 'jp_pk': p.pk, 'name': p.name}
        p.delete()
        if request.headers.get('X-Fetch') == '1':
            return JsonResponse(data)
    next_url = request.POST.get('next') or request.GET.get('next')
    return redirect(next_url if next_url else reverse('job_personnel_page', kwargs={'job_pk': job_pk}))


def job_vehicles_page(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    job_vehicles = list(job.vehicles.all().order_by('vehicle_type', 'name'))
    assigned_names = {jv.name for jv in job_vehicles}
    fleet_by_name = {v.name: v for v in Vehicle.objects.all()}
    for jv in job_vehicles:
        fv = fleet_by_name.get(jv.name)
        jv.description = fv.description if fv else ''
        jv.rego        = fv.rego        if fv else ''
        jv.make        = fv.make        if fv else ''
        jv.model_name  = fv.model_name  if fv else ''
    fleet_tabs = [
        {'tab': 'lv',  'label': 'LV',        'type': 'Light Vehicle',  'fleet': Vehicle.objects.filter(vehicle_type='Light Vehicle',  is_active=True).order_by('name')},
        {'tab': 'mv',  'label': 'MV',        'type': 'Medium Vehicle', 'fleet': Vehicle.objects.filter(vehicle_type='Medium Vehicle', is_active=True).order_by('name')},
        {'tab': 'hv',  'label': 'HV',        'type': 'Heavy Vehicle',  'fleet': Vehicle.objects.filter(vehicle_type='Heavy Vehicle',  is_active=True).order_by('name')},
        {'tab': 'tt',  'label': 'Tilt Tray', 'type': 'Tilt Tray',     'fleet': Vehicle.objects.filter(vehicle_type='Tilt Tray',      is_active=True).order_by('name')},
        {'tab': 'vib', 'label': 'Vibrator',  'type': 'Vibrator',       'fleet': Vehicle.objects.filter(vehicle_type='Vibrator',       is_active=True).order_by('name')},
        {'tab': 'atv', 'label': 'ATV',       'type': 'ATV',            'fleet': Vehicle.objects.filter(vehicle_type='ATV',            is_active=True).order_by('name')},
    ]
    return render(request, 'reports/job_vehicles_page.html', {
        'job': job,
        'job_vehicles': job_vehicles,
        'assigned_names': assigned_names,
        'fleet_tabs': fleet_tabs,
    })


def job_vehicle_add_json(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    v = get_object_or_404(Vehicle, pk=data.get('vehicle_pk'))
    jv, _ = JobVehicle.objects.get_or_create(
        job=job, name=v.name,
        defaults={'vehicle_type': v.vehicle_type, 'notes': v.notes},
    )
    return JsonResponse({'ok': True, 'jv_pk': jv.pk, 'name': jv.name, 'vehicle_type': jv.vehicle_type,
                         'description': v.description, 'rego': v.rego, 'make': v.make, 'model_name': v.model_name})


def job_vehicle_delete_json(request, pk):
    jv = get_object_or_404(JobVehicle, pk=pk)
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    jv.delete()
    return JsonResponse({'ok': True})


def job_vehicle_add(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            JobVehicle.objects.create(
                job=job,
                name=name,
                vehicle_type=request.POST.get('vehicle_type', '').strip(),
                notes=request.POST.get('notes', '').strip(),
            )
    return redirect('job_detail', pk=job_pk)


def job_vehicle_delete(request, pk):
    v = get_object_or_404(JobVehicle, pk=pk)
    job_pk = v.job.pk
    if request.method == 'POST':
        v.delete()
    return redirect('job_detail', pk=job_pk)


def job_equipment_add(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            JobEquipment.objects.create(
                job=job,
                name=name,
                equipment_type=request.POST.get('equipment_type', '').strip(),
                notes=request.POST.get('notes', '').strip(),
            )
    return redirect('job_detail', pk=job_pk)


def job_equipment_delete(request, pk):
    e = get_object_or_404(JobEquipment, pk=pk)
    job_pk = e.job.pk
    if request.method == 'POST':
        e.delete()
    return redirect('job_detail', pk=job_pk)


def job_equipment_page(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    job_equipment = list(job.equipment.all())
    assigned_names = {je.name for je in job_equipment}
    all_eq = list(Equipment.objects.all())
    known = ('Trailer', 'Genset', 'Starlink')
    fleet_tabs = [
        {'tab': 'trailer',  'label': 'Trailers',  'type': 'Trailer',  'fleet': [e for e in all_eq if e.equipment_type == 'Trailer']},
        {'tab': 'genset',   'label': 'Gensets',   'type': 'Genset',   'fleet': [e for e in all_eq if e.equipment_type == 'Genset']},
        {'tab': 'starlink', 'label': 'Starlinks', 'type': 'Starlink', 'fleet': [e for e in all_eq if e.equipment_type == 'Starlink']},
        {'tab': 'other',    'label': 'Other',     'type': '',         'fleet': [e for e in all_eq if e.equipment_type not in known]},
    ]
    return render(request, 'reports/job_equipment_page.html', {
        'job': job,
        'job_equipment': job_equipment,
        'fleet_tabs': fleet_tabs,
        'assigned_names': assigned_names,
    })


def job_equipment_add_json(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    job = get_object_or_404(Job, pk=job_pk)
    data = json.loads(request.body)
    eq_pk = data.get('equipment_pk')
    eq = get_object_or_404(Equipment, pk=eq_pk)
    je, created = JobEquipment.objects.get_or_create(
        job=job, name=eq.name,
        defaults={'equipment_type': eq.equipment_type, 'notes': eq.notes},
    )
    return JsonResponse({'ok': True, 'je_pk': je.pk, 'name': je.name,
                         'equipment_type': je.equipment_type, 'notes': je.notes})


def job_equipment_delete_json(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    je = get_object_or_404(JobEquipment, pk=pk)
    je.delete()
    return JsonResponse({'ok': True})


def equipment_allocation_add(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    je = get_object_or_404(JobEquipment, pk=data.get('je_pk'), job__pk=job_pk)
    name = data.get('person_name', '').strip()
    if not name:
        return JsonResponse({'error': 'name required'}, status=400)
    alloc, _ = EquipmentAllocation.objects.get_or_create(job_equipment=je, person_name=name)
    return JsonResponse({'ok': True, 'alloc_pk': alloc.pk})


def equipment_allocation_remove(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    alloc = get_object_or_404(EquipmentAllocation, pk=pk)
    alloc.delete()
    return JsonResponse({'ok': True})


def allocation_copy_tab(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    src = data.get('from_tab', '')
    dst = data.get('to_tab', '')
    valid = {'mob', 'job', 'demob', 'other'}
    if src not in valid or dst not in valid or src == dst:
        return JsonResponse({'error': 'invalid tabs'}, status=400)
    # Clear destination tab
    VehicleAllocation.objects.filter(job_vehicle__job__pk=job_pk, tab=dst).delete()
    JobEquipmentVehicleLink.objects.filter(job_vehicle__job__pk=job_pk, tab=dst).delete()
    # Copy from source tab
    for a in VehicleAllocation.objects.filter(job_vehicle__job__pk=job_pk, tab=src):
        VehicleAllocation.objects.create(job_vehicle=a.job_vehicle, person_name=a.person_name, tab=dst)
    for lnk in JobEquipmentVehicleLink.objects.filter(job_vehicle__job__pk=job_pk, tab=src):
        JobEquipmentVehicleLink.objects.create(job_equipment=lnk.job_equipment, job_vehicle=lnk.job_vehicle, tab=dst)
    return JsonResponse({'ok': True})


def _build_allocation_ctx(job, tabs=None, include_travel=True, include_equipment=True):
    """Shared context builder for allocation PDF/Word exports."""
    TABS = [('mob', 'Mob'), ('job', 'Job'), ('demob', 'Demob'), ('other', 'Other')]
    job_vehicles = list(job.vehicles.all())
    job_equipment_all = list(job.equipment.all())
    personnel_names = list(
        job.personnel.filter(is_separator=False).order_by('name').values_list('name', flat=True))

    selected_tabs = [t for t in TABS if tabs is None or t[0] in tabs]

    tabs_data = []
    for tab_key, tab_label in selected_tabs:
        v_allocs = list(VehicleAllocation.objects.filter(job_vehicle__job=job, tab=tab_key))
        eq_links = list(JobEquipmentVehicleLink.objects.filter(
            job_vehicle__job=job, tab=tab_key).select_related('job_equipment'))

        alloc_by_veh = {}
        for a in v_allocs:
            alloc_by_veh.setdefault(a.job_vehicle_id, []).append(
                {'pk': a.pk, 'person_name': a.person_name})

        eq_link_by_veh = {}
        for lnk in eq_links:
            eq_link_by_veh.setdefault(lnk.job_vehicle_id, []).append({
                'link_pk': lnk.pk,
                'je_pk': lnk.job_equipment_id,
                'name': lnk.job_equipment.name,
                'equipment_type': lnk.job_equipment.equipment_type,
            })

        tabs_data.append({
            'tab': tab_key,
            'label': tab_label,
            'vehicles': [{'jv': jv,
                          'allocs': alloc_by_veh.get(jv.pk, []),
                          'eq_links': eq_link_by_veh.get(jv.pk, [])}
                         for jv in job_vehicles],
        })

    from collections import defaultdict
    ea_by_person = defaultdict(list)
    for alloc in EquipmentAllocation.objects.filter(job_equipment__job=job).select_related('job_equipment'):
        ea_by_person[alloc.person_name].append({
            'alloc_pk': alloc.pk,
            'je_pk': alloc.job_equipment_id,
            'name': alloc.job_equipment.name,
            'equipment_type': alloc.job_equipment.equipment_type,
        })
    eq_alloc_people = [
        {'name': p, 'eq_allocs': ea_by_person.get(p, [])}
        for p in personnel_names
    ]

    return {
        'job': job,
        'tabs_data': tabs_data,
        'eq_alloc_people': eq_alloc_people,
        'include_travel': include_travel,
        'include_equipment': include_equipment,
    }


def allocation_pdf(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    page = request.GET.get('page', 'all')       # travel | equipment | all
    tab_param = request.GET.get('tab', 'all')   # mob|job|demob|other|all

    include_travel = page in ('travel', 'all')
    include_equipment = page in ('equipment', 'all')
    tabs = None if tab_param == 'all' else [t.strip() for t in tab_param.split(',')]

    ctx = _build_allocation_ctx(job, tabs=tabs,
                                include_travel=include_travel,
                                include_equipment=include_equipment)
    html = render_to_string('reports/allocation_pdf.html', ctx, request=request)

    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page_obj = browser.new_page()
            page_obj.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            page_obj.wait_for_load_state('networkidle')
            pdf_bytes = page_obj.pdf(
                format='A4',
                landscape=True,
                print_background=True,
                display_header_footer=True,
                header_template='<span></span>',
                footer_template=(
                    '<div style="width:100%; font-family:Helvetica,Arial,sans-serif; '
                    'font-size:8px; color:#888; text-align:center;">'
                    'Page <span class="pageNumber"></span> of <span class="totalPages"></span>'
                    '</div>'
                ),
                margin={'top': '1.5cm', 'bottom': '1.8cm', 'left': '1.5cm', 'right': '1.5cm'},
            )
            browser.close()
    finally:
        os.unlink(tmp.name)

    filename = f'Allocations_{job.job_number}.pdf'
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def allocation_word(request, job_pk):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    job = get_object_or_404(Job, pk=job_pk)
    page = request.GET.get('page', 'all')
    tab_param = request.GET.get('tab', 'all')

    include_travel = page in ('travel', 'all')
    include_equipment = page in ('equipment', 'all')
    tabs = None if tab_param == 'all' else [t.strip() for t in tab_param.split(',')]

    ctx = _build_allocation_ctx(job, tabs=tabs,
                                include_travel=include_travel,
                                include_equipment=include_equipment)

    NAVY = RGBColor(0x1a, 0x3a, 0x5c)
    GREEN = RGBColor(0x5b, 0x6f, 0x3a)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), 'BBBBBB')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def sec_heading(doc, text, color=NAVY):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1a3a5c')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def sub_heading(doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True

    doc = Document()
    for section in doc.sections:
        section.orientation = 1  # landscape
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    title_p = doc.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f'Allocations — {job.job_number}')
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = NAVY
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(job.client or '')
    if job.project_name:
        mr.text += f' — {job.project_name}'
    mr.font.size = Pt(10)

    # ── Travel Allocations ─────────────────────────────────
    if ctx['include_travel']:
        sec_heading(doc, 'Travel Allocations')
        for t in ctx['tabs_data']:
            sub_heading(doc, t['label'])
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, txt in enumerate(['Vehicle', 'Type', 'Personnel', 'Equipment']):
                hdr[i].text = ''
                run = hdr[i].paragraphs[0].add_run(txt)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(hdr[i], '1a3a5c')

            for vd in t['vehicles']:
                row = table.add_row().cells
                # Vehicle name
                row[0].text = ''
                r = row[0].paragraphs[0].add_run(vd['jv'].name)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = NAVY
                set_cell_bg(row[0], 'EEF3F8')
                # Vehicle type
                row[1].text = ''
                r2 = row[1].paragraphs[0].add_run(vd['jv'].vehicle_type or '')
                r2.font.size = Pt(8)
                set_cell_bg(row[1], 'EEF3F8')
                # Personnel
                row[2].text = ''
                if vd['allocs']:
                    p2 = row[2].paragraphs[0]
                    for i, a in enumerate(vd['allocs']):
                        if i > 0:
                            p2.add_run('  ')
                        run3 = p2.add_run(a['person_name'])
                        run3.font.size = Pt(9)
                else:
                    r3 = row[2].paragraphs[0].add_run('—')
                    r3.font.size = Pt(9)
                    r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                # Equipment
                row[3].text = ''
                if vd['eq_links']:
                    p3 = row[3].paragraphs[0]
                    for i, lnk in enumerate(vd['eq_links']):
                        if i > 0:
                            p3.add_run('  ')
                        txt_eq = lnk['name']
                        if lnk['equipment_type']:
                            txt_eq += f' ({lnk["equipment_type"]})'
                        run4 = p3.add_run(txt_eq)
                        run4.font.size = Pt(9)
                else:
                    r4 = row[3].paragraphs[0].add_run('—')
                    r4.font.size = Pt(9)
                    r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                for cell in row:
                    set_cell_border(cell)

    # ── Equipment Allocation ───────────────────────────────
    if ctx['include_equipment']:
        if ctx['include_travel']:
            p = doc.add_paragraph()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            p.runs[0]._r.append(br) if p.runs else p.add_run()._r.append(br)
        sec_heading(doc, 'Equipment Allocation')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, txt in enumerate(['Person', 'Equipment']):
            hdr[i].text = ''
            run = hdr[i].paragraphs[0].add_run(txt)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(hdr[i], '1a3a5c')
        for pd in ctx['eq_alloc_people']:
            row = table.add_row().cells
            row[0].text = ''
            r = row[0].paragraphs[0].add_run(pd['name'])
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = NAVY
            set_cell_bg(row[0], 'EEF3F8')
            row[1].text = ''
            if pd['eq_allocs']:
                p2 = row[1].paragraphs[0]
                for i, ea in enumerate(pd['eq_allocs']):
                    if i > 0:
                        p2.add_run('  ')
                    txt_eq = ea['name']
                    if ea['equipment_type']:
                        txt_eq += f' ({ea["equipment_type"]})'
                    run2 = p2.add_run(txt_eq)
                    run2.font.size = Pt(9)
            else:
                r2 = row[1].paragraphs[0].add_run('—')
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            for cell in row:
                set_cell_border(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f'Allocations_{job.job_number}.docx'
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def vehicle_equipment_link_add(request, job_pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    je = get_object_or_404(JobEquipment, pk=data['je_pk'], job__pk=job_pk)
    jv = get_object_or_404(JobVehicle, pk=data['jv_pk'], job__pk=job_pk)
    tab = data.get('tab', 'job')
    link, created = JobEquipmentVehicleLink.objects.get_or_create(
        job_equipment=je, tab=tab, defaults={'job_vehicle': jv})
    if not created and link.job_vehicle_id != jv.pk:
        link.job_vehicle = jv
        link.save()
    return JsonResponse({'ok': True, 'link_pk': link.pk})


def vehicle_equipment_link_remove(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    link = get_object_or_404(JobEquipmentVehicleLink, pk=pk)
    link.delete()
    return JsonResponse({'ok': True})


def browse_folder(request):
    """Open a native Windows folder picker and return the selected path as JSON."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.wm_attributes('-topmost', True)
        initial = request.GET.get('initial', '')
        path = filedialog.askdirectory(parent=root, initialdir=initial or '/', title='Select folder')
        root.destroy()
        if path:
            return JsonResponse({'path': path.replace('/', os.sep)})
        return JsonResponse({'path': ''})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def category_add(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            ActivityCategory.objects.get_or_create(name=name)
    return redirect(request.META.get('HTTP_REFERER', '/'))


def category_delete(request, pk):
    ActivityCategory.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def supervisor_activity_template_save(request, report_pk):
    if request.method == 'POST':
        import json
        name = request.POST.get('template_name', '').strip()
        if name:
            report = DailyReport.objects.get(pk=report_pk)
            rows = [
                {
                    'contractor': a.details,
                    'name': a.notes,
                    'job_title': a.job_title,
                    'hours': a.hours,
                    'category': a.category,
                    'activity_type': a.activity_type,
                }
                for a in report.activities.all()
            ]
            SupervisorActivityTemplate.objects.create(name=name, rows_json=json.dumps(rows))
    return redirect(request.META.get('HTTP_REFERER', '/'))


def supervisor_activity_template_delete(request, pk):
    SupervisorActivityTemplate.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))





def supervisor_option_add(request):
    if request.method == 'POST':
        option_type = request.POST.get('option_type', '').strip()
        name = request.POST.get('name', '').strip()
        if option_type and name:
            SupervisorOption.objects.get_or_create(option_type=option_type, name=name)
    return redirect(request.META.get('HTTP_REFERER', '/'))


def supervisor_option_delete(request, pk):
    SupervisorOption.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def activity_type_add(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            max_order = ActivityType.objects.order_by('-order').values_list('order', flat=True).first() or 0
            pct = request.POST.get('chargeable_percentage', '100').strip()
            try:
                pct = max(0, min(100, int(pct)))
            except ValueError:
                pct = 100
            ActivityType.objects.get_or_create(name=name, defaults={'order': max_order + 1, 'chargeable_percentage': pct})
    return redirect(request.META.get('HTTP_REFERER', '/'))


def activity_type_update(request, pk):
    if request.method == 'POST':
        at = get_object_or_404(ActivityType, pk=pk)
        pct = request.POST.get('chargeable_percentage', '').strip()
        try:
            at.chargeable_percentage = max(0, min(100, int(pct)))
            at.save()
        except ValueError:
            pass
    return redirect(request.META.get('HTTP_REFERER', '/'))


def activity_type_delete(request, pk):
    ActivityType.objects.filter(pk=pk).delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


def personnel_delete(request, pk):
    person = get_object_or_404(PersonnelName, pk=pk)
    person.delete()
    return redirect(request.META.get('HTTP_REFERER', '/'))


# --- Map ---


def job_progress_map(request, pk):
    job = get_object_or_404(Job, pk=pk)

    sx_file = job.survey_files.filter(file_type='sx', is_final=True).first()
    if not sx_file:
        return render(request, 'reports/job_progress_map.html', {
            'job': job, 'map_html': None,
            'error': 'No final Sx survey file marked. Go to the job page and mark the Sx file as Final.'
        })

    zone = int(sx_file.zone)
    epsg_in = _datum_epsg(sx_file.datum, zone)

    sx_df = _read_survey_csv(sx_file.file.path)
    sx_df['_line'] = sx_df['Line'].round(1)
    sx_df['_point'] = sx_df['Point'].round(1)

    transformer = Transformer.from_crs(epsg_in, 'EPSG:4326', always_xy=True)
    lons, lats = transformer.transform(sx_df['X'].values, sx_df['Y'].values)
    sx_df['lat'] = lats
    sx_df['lon'] = lons

    # Collect all shots across all reports for this job
    acquired = set()
    void = set()
    for report in job.reports.all():
        obs = report.files.filter(file_type='obslog').first()
        if obs:
            try:
                obs_df = _read_csv(obs.file.path, skiprows=2)
                obs_df.columns = obs_df.columns.str.strip()
                for _, row in obs_df.iterrows():
                    key = (round(float(row['Line']), 1), round(float(row['Station']), 1))
                    if str(row['Status']).strip() == 'Void':
                        void.add(key)
                    else:
                        acquired.add(key)
            except Exception:
                pass

    planned, done, voided = [], [], []
    for _, row in sx_df.iterrows():
        key = (row['_line'], row['_point'])
        feat = {
            'type': 'Feature',
            'geometry': {'type': 'Point', 'coordinates': [row['lon'], row['lat']]},
            'properties': {'label': f"L{int(row['_line'])} SP{row['_point']}"},
        }
        if key in acquired:
            done.append(feat)
        elif key in void:
            voided.append(feat)
        else:
            planned.append(feat)

    centre = [sx_df['lat'].mean(), sx_df['lon'].mean()]
    m = folium.Map(location=centre, zoom_start=12, tiles=None)

    folium.TileLayer(
        tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
        attr='Esri World Imagery', name='Satellite', overlay=False, control=True,
    ).add_to(m)
    folium.TileLayer('OpenStreetMap', name='Street Map', overlay=False, control=True).add_to(m)

    def geojson_layer(features, name, color, radius=3):
        if features:
            folium.GeoJson(
                {'type': 'FeatureCollection', 'features': features},
                name=name,
                marker=folium.CircleMarker(radius=radius, fill=True, fill_color=color, color=color, fill_opacity=0.85),
                tooltip=folium.GeoJsonTooltip(fields=['label'], aliases=['']),
            ).add_to(m)

    geojson_layer(planned, 'Planned', '#aaaaaa', radius=2)
    geojson_layer(done, 'Acquired', '#2196f3', radius=4)
    geojson_layer(voided, 'Void', '#e63946', radius=4)

    folium.LayerControl(position='topright').add_to(m)

    stats = {
        'planned': len(sx_df),
        'acquired': len(acquired),
        'void': len(void),
        'remaining': len(planned),
    }
    map_html = m._repr_html_()
    return render(request, 'reports/job_progress_map.html', {
        'job': job,
        'map_html': map_html,
        'stats': stats,
    })


def report_photo_upload(request, report_pk):
    from django.db.models import Max
    report = get_object_or_404(DailyReport, pk=report_pk)
    if request.method == 'POST':
        max_order = report.photos.aggregate(m=Max('order'))['m'] or 0
        for i, f in enumerate(request.FILES.getlist('images')):
            ReportPhoto.objects.create(report=report, image=f, order=max_order + i + 1)
    return redirect(f"{reverse('report_detail', args=[report_pk])}?tab=photos")


def report_photo_delete(request, pk):
    photo = get_object_or_404(ReportPhoto, pk=pk)
    report_pk = photo.report_id
    if request.method == 'POST':
        photo.image.delete(save=False)
        photo.delete()
    return redirect(f"{reverse('report_detail', args=[report_pk])}?tab=photos")


def report_photo_update(request, pk):
    photo = get_object_or_404(ReportPhoto, pk=pk)
    if request.method == 'POST':
        photo.caption = request.POST.get('caption', '').strip()
        photo.border_style = request.POST.get('border_style', 'none')
        photo.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def report_photo_reorder(request, report_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        for item in data:
            ReportPhoto.objects.filter(pk=item['pk'], report_id=report_pk).update(order=item['order'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def report_photo_columns(request, report_pk):
    if request.method == 'POST':
        report = get_object_or_404(DailyReport, pk=report_pk)
        cols = int(request.POST.get('columns', 2))
        report.photo_columns = cols if cols in (1, 2) else 2
        report.save()
        return JsonResponse({'ok': True, 'columns': report.photo_columns})
    return JsonResponse({'ok': False}, status=400)


# ── Journey Management ────────────────────────────────────────────────────────

def journey_list(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    plans = job.journey_plans.prefetch_related('personnel').all()
    v2_plans = job.journey_v2_plans.prefetch_related('personnel').all()
    return render(request, 'reports/journey_list.html', {'job': job, 'plans': plans, 'v2_plans': v2_plans})


def _save_journey(request, plan):
    """Save all fields from a journey form POST onto plan (and personnel rows)."""
    def d(field): return request.POST.get(field, '').strip() or None
    def b(field): v = request.POST.get(field); return True if v == 'Y' else (False if v == 'N' else None)

    plan.plan_number = request.POST.get('plan_number', '').strip()
    plan.departing_from = request.POST.get('departing_from', '').strip()
    plan.depart_date = d('depart_date')
    plan.depart_time = d('depart_time')
    plan.depart_contact = request.POST.get('depart_contact', '').strip()
    plan.depart_phone = request.POST.get('depart_phone', '').strip()
    plan.overnight_location = request.POST.get('overnight_location', '').strip()
    plan.overnight_arrival_date = d('overnight_arrival_date')
    plan.overnight_arrival_time = d('overnight_arrival_time')
    plan.overnight_departure_date = d('overnight_departure_date')
    plan.overnight_departure_time = d('overnight_departure_time')
    plan.arriving_at = request.POST.get('arriving_at', '').strip()
    plan.arrive_date = d('arrive_date')
    plan.arrive_time = d('arrive_time')
    plan.arrive_contact = request.POST.get('arrive_contact', '').strip()
    plan.arrive_phone = request.POST.get('arrive_phone', '').strip()
    plan.route = request.POST.get('route', '').strip()
    plan.break_journey_at = request.POST.get('break_journey_at', '').strip()
    plan.radio_channel = request.POST.get('radio_channel', '').strip()
    plan.other_instructions = request.POST.get('other_instructions', '').strip()
    plan.route_waypoints = request.POST.get('route_waypoints', '').strip()
    plan.rest_stops_json = request.POST.get('rest_stops_json', '').strip()
    plan.coordinator_name = request.POST.get('coordinator_name', '').strip()
    plan.coordinator_phone = request.POST.get('coordinator_phone', '').strip()
    plan.plan_communicated = b('plan_communicated')
    plan.before_signature = request.POST.get('before_signature', '').strip()
    plan.before_date = d('before_date')
    plan.journey_completed = b('journey_completed')
    plan.after_signature = request.POST.get('after_signature', '').strip()
    plan.after_date = d('after_date')
    plan.action_items = request.POST.get('action_items', '').strip()
    plan.include_map_in_pdf = 'include_map_in_pdf' in request.POST
    plan.map_tile_layer = request.POST.get('map_tile_layer', 'street')
    plan.save()

    plan.personnel.all().delete()
    regos = request.POST.getlist('p_rego')
    names = request.POST.getlist('p_name')
    drivers = request.POST.getlist('p_driver')
    phones = request.POST.getlist('p_phone')
    signatures = request.POST.getlist('p_signature')
    for i, name in enumerate(names):
        if not name.strip():
            continue
        JourneyPersonnel.objects.create(
            plan=plan,
            rego=regos[i].strip() if i < len(regos) else '',
            name=name.strip(),
            is_driver=(drivers[i] == 'Y') if i < len(drivers) else False,
            phone=phones[i].strip() if i < len(phones) else '',
            signature=signatures[i].strip() if i < len(signatures) else '',
            order=i,
        )


def journey_create(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        plan = JourneyManagementPlan(job=job)
        _save_journey(request, plan)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': plan.pk, 'edit_url': reverse('journey_edit', args=[plan.pk])})
        return redirect('journey_list', job_pk=job_pk)
    personnel_phones = {p.name: p.phone for p in Personnel.objects.exclude(phone='')}
    jp_qs = job.personnel.order_by('name')
    mob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=job, status__iexact='mobilise').values_list('job_personnel__name', flat=True))
    demob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=job, status__iexact='demobilise').values_list('job_personnel__name', flat=True))
    job_personnel_data = [{'name': p.name, 'phone': personnel_phones.get(p.name, ''),
                           'mob': p.name in mob_names, 'demob': p.name in demob_names} for p in jp_qs]
    return render(request, 'reports/journey_form.html', {
        'job': job,
        'plan': None,
        'job_personnel': jp_qs,
        'personnel_phones': personnel_phones,
        'job_personnel_data': job_personnel_data,
    })


def journey_edit(request, pk):
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)
    if request.method == 'POST':
        _save_journey(request, plan)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': plan.pk})
        return redirect('journey_list', job_pk=plan.job_id)
    personnel_phones = {p.name: p.phone for p in Personnel.objects.exclude(phone='')}
    jp_qs = plan.job.personnel.all()
    mob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=plan.job, status__iexact='mobilise').values_list('job_personnel__name', flat=True))
    demob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=plan.job, status__iexact='demobilise').values_list('job_personnel__name', flat=True))
    job_personnel_data = [{'name': p.name, 'phone': personnel_phones.get(p.name, ''),
                           'mob': p.name in mob_names, 'demob': p.name in demob_names} for p in jp_qs]

    # Build allocation data per tab for "From Allocations" auto-fill (exclude Vibrators)
    fleet_by_name = {v.name: v for v in Vehicle.objects.all()}
    job_vehicles_nv = list(plan.job.vehicles.exclude(vehicle_type='Vibrator'))
    # Pre-fetch trailer and starlink links for the whole job
    eq_links_for_suffix = JobEquipmentVehicleLink.objects.filter(
        job_vehicle__job=plan.job,
        job_equipment__equipment_type__in=('Trailer', 'Starlink')
    ).select_related('job_equipment')
    trailers_by_veh_tab = {}
    starlinks_by_veh_tab = {}
    for lnk in eq_links_for_suffix:
        key = (lnk.job_vehicle_id, lnk.tab)
        if lnk.job_equipment.equipment_type == 'Trailer':
            trailers_by_veh_tab.setdefault(key, []).append(lnk.job_equipment.name)
        else:
            starlinks_by_veh_tab.setdefault(key, []).append(lnk.job_equipment.name)

    alloc_by_tab = {}
    for tab_key in ('mob', 'job', 'demob', 'other'):
        vehicles = []
        for jv in job_vehicles_nv:
            fleet = fleet_by_name.get(jv.name)
            rego_label = ' '.join(filter(None, [jv.name, fleet.rego if fleet else '']))
            suffix_parts = (
                [f'({t})' for t in trailers_by_veh_tab.get((jv.pk, tab_key), [])] +
                [f'({s})' for s in starlinks_by_veh_tab.get((jv.pk, tab_key), [])]
            )
            suffix = ' '.join(suffix_parts)
            people = [{'name': a.person_name + (' ' + suffix if suffix else ''),
                       'phone': personnel_phones.get(a.person_name, '')}
                      for a in jv.allocations.filter(tab=tab_key)]
            if people:
                vehicles.append({'rego_label': rego_label, 'people': people})
        alloc_by_tab[tab_key] = vehicles

    return render(request, 'reports/journey_form.html', {
        'job': plan.job,
        'plan': plan,
        'job_personnel': jp_qs,
        'personnel_phones': personnel_phones,
        'job_personnel_data': job_personnel_data,
        'alloc_by_tab': alloc_by_tab,
        'alloc_tabs': [k for k, v in alloc_by_tab.items() if v],
    })


def journey_delete(request, pk):
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)
    job_pk = plan.job_id
    if request.method == 'POST':
        plan.delete()
    return redirect('journey_list', job_pk=job_pk)


def journey_copy(request, pk):
    src = get_object_or_404(JourneyManagementPlan, pk=pk)
    new_plan = JourneyManagementPlan(
        job=src.job,
        departing_from=src.departing_from,
        depart_contact=src.depart_contact,
        depart_phone=src.depart_phone,
        overnight_location=src.overnight_location,
        arriving_at=src.arriving_at,
        arrive_contact=src.arrive_contact,
        arrive_phone=src.arrive_phone,
        route=src.route,
        break_journey_at=src.break_journey_at,
        radio_channel=src.radio_channel,
        other_instructions=src.other_instructions,
        route_waypoints=src.route_waypoints,
        coordinator_name=src.coordinator_name,
        coordinator_phone=src.coordinator_phone,
    )
    new_plan.save()
    for p in src.personnel.all():
        JourneyPersonnel.objects.create(
            plan=new_plan,
            rego=p.rego,
            name=p.name,
            is_driver=p.is_driver,
            phone=p.phone,
            order=p.order,
        )
    return redirect('journey_edit', pk=new_plan.pk)


def _logo_b64():
    import base64 as _b64
    _logo_path = os.path.join(settings.MEDIA_ROOT, 'velseis-logo.png')
    if os.path.exists(_logo_path):
        with open(_logo_path, 'rb') as _lf:
            return 'data:image/png;base64,' + _b64.b64encode(_lf.read()).decode()
    return None


def _map_screenshot_b64(waypoints_json, tile_layer='street'):
    """Launch a dedicated Playwright session to screenshot a map and return base64 PNG data URL."""
    import json as _json, base64 as _b64
    try:
        wps = _json.loads(waypoints_json) if waypoints_json else []
    except Exception:
        return None
    if not wps:
        return None

    if tile_layer == 'satellite':
        tile_url = 'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}'
    else:
        tile_url = 'https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png'

    wps_json = _json.dumps(wps)
    map_html = f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css">
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>*{{margin:0;padding:0;box-sizing:border-box;}} #map{{width:760px;height:320px;}}</style>
</head><body>
<div id="map"></div>
<script>
var map = L.map('map', {{zoomControl:false, attributionControl:false}});
L.tileLayer('{tile_url}', {{maxZoom:19}}).addTo(map);
var wps = {wps_json};
var latlngs = wps.map(function(w){{return [w.lat, w.lon];}});
L.polyline(latlngs, {{color:'#1a3a5c', weight:3, opacity:0.9}}).addTo(map);
wps.forEach(function(w, i) {{
    var icon = L.divIcon({{
        html: '<div style="background:#1a3a5c;color:white;border-radius:50%;width:22px;height:22px;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:bold;border:2px solid white;box-shadow:0 1px 4px rgba(0,0,0,0.5)">'+(i+1)+'</div>',
        className:'', iconSize:[22,22], iconAnchor:[11,11]
    }});
    L.marker([w.lat, w.lon], {{icon:icon}}).addTo(map);
}});
map.fitBounds(latlngs, {{padding:[28,28]}});
</script>
</body></html>"""

    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={'width': 760, 'height': 320})
        page.set_content(map_html, wait_until='networkidle')
        page.wait_for_timeout(3500)
        img_bytes = page.screenshot(clip={'x': 0, 'y': 0, 'width': 760, 'height': 320})
        browser.close()
    return 'data:image/png;base64,' + _b64.b64encode(img_bytes).decode()


def job_overview_pdf(request, pk):
    job = get_object_or_404(Job, pk=pk)
    logo_b64 = _logo_b64()

    # Personnel
    personnel = list(job.personnel.filter(is_separator=False).prefetch_related('roles'))

    # Vehicles & Equipment
    vehicles = list(job.vehicles.all())
    equipment = list(job.equipment.all())

    # Daily reports — production & supervisor, ordered by date
    reports = list(job.reports.order_by('date').select_related())
    production_reports = [r for r in reports if r.report_type == 'production']
    supervisor_reports = [r for r in reports if r.report_type == 'supervisor']

    # Toolbox meetings
    toolbox_meetings = list(job.toolbox_meetings.order_by('date').prefetch_related('attendees'))

    # Journey management plans
    journey_plans = list(job.journey_plans.order_by('depart_date').prefetch_related('personnel'))

    # Survey file counts
    rx_count = sx_count = None
    for sf in job.survey_files.filter(is_final=True):
        try:
            count = len(_read_csv(sf.file.path))
            if sf.file_type == 'rx':
                rx_count = count
            elif sf.file_type == 'sx':
                sx_count = count
        except Exception:
            pass

    # Date range covered by production reports
    first_date = production_reports[0].date if production_reports else None
    last_date = production_reports[-1].date if production_reports else None

    generated_date = date.today().strftime('%d %b %Y').lstrip('0')

    html = render_to_string('reports/job_overview_pdf.html', {
        'job': job,
        'logo_b64': logo_b64,
        'personnel': personnel,
        'vehicles': vehicles,
        'equipment': equipment,
        'production_reports': production_reports,
        'supervisor_reports': supervisor_reports,
        'toolbox_meetings': toolbox_meetings,
        'journey_plans': journey_plans,
        'rx_count': rx_count,
        'sx_count': sx_count,
        'first_date': first_date,
        'last_date': last_date,
        'generated_date': generated_date,
    })
    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        pdf_bytes = pg.pdf(format='A4', print_background=True, margin={'top': '1.5cm', 'bottom': '1.5cm', 'left': '1.5cm', 'right': '1.5cm'})
        browser.close()
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    fname = f"Job_Overview_{job.job_number.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'inline; filename="{fname}"'
    return response


def _haversine_km(lat1, lon1, lat2, lon2):
    import math
    R = 6371
    d = math.radians
    dlat = d(lat2 - lat1)
    dlon = d(lon2 - lon1)
    a = math.sin(dlat/2)**2 + math.cos(d(lat1)) * math.cos(d(lat2)) * math.sin(dlon/2)**2
    return round(R * 2 * math.asin(math.sqrt(a)))


def _build_waypoints(raw_json):
    import json as _json
    try:
        wps = _json.loads(raw_json) if raw_json else []
    except Exception:
        return []
    for i, wp in enumerate(wps):
        if i + 1 < len(wps):
            nxt = wps[i + 1]
            try:
                wp['km'] = _haversine_km(wp['lat'], wp['lon'], nxt['lat'], nxt['lon'])
            except Exception:
                wp['km'] = None
        else:
            wp['km'] = None
    return wps


def journey_pdf(request, pk):
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)
    logo_b64 = _logo_b64()
    route_waypoints = _build_waypoints(plan.route_waypoints)
    route_text_parts = [p.strip() for p in plan.route.split('→') if p.strip()] if plan.route else []
    route_total_km = sum(wp['km'] for wp in route_waypoints if wp.get('km') is not None)
    # Map screenshot runs in its own Playwright session (before template rendering)
    map_b64 = None
    if plan.include_map_in_pdf and plan.route_waypoints:
        map_b64 = _map_screenshot_b64(plan.route_waypoints, plan.map_tile_layer)

    # Template rendering happens outside Playwright (avoids SynchronousOnlyOperation)
    html = render_to_string('reports/journey_pdf.html', {
        'plan': plan, 'logo_b64': logo_b64,
        'route_waypoints': route_waypoints,
        'route_text_parts': route_text_parts,
        'route_total_km': route_total_km,
        'include_map': plan.include_map_in_pdf,
        'map_b64': map_b64,
    })

    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        pdf_bytes = pg.pdf(format='A4', print_background=True, margin={'top': '1.5cm', 'bottom': '1.5cm', 'left': '1.5cm', 'right': '1.5cm'})
        browser.close()
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    date_str = plan.depart_date.strftime('%d_%m_%Y') if plan.depart_date else 'undated'
    disposition = 'attachment' if request.GET.get('download') else 'inline'
    response['Content-Disposition'] = f'{disposition}; filename="Journey_Management_Plan_{date_str}.pdf"'
    return response


def journey_preview(request, pk):
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)
    logo_b64 = _logo_b64()
    route_waypoints = _build_waypoints(plan.route_waypoints)
    route_text_parts = [p.strip() for p in plan.route.split('→') if p.strip()] if plan.route else []
    route_total_km = sum(wp['km'] for wp in route_waypoints if wp.get('km') is not None)
    html = render_to_string('reports/journey_pdf.html', {
        'plan': plan, 'logo_b64': logo_b64,
        'route_waypoints': route_waypoints,
        'route_text_parts': route_text_parts,
        'route_total_km': route_total_km,
        'include_map': plan.include_map_in_pdf,
    })
    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        pdf_bytes = pg.pdf(format='A4', print_background=True, margin={'top': '1.5cm', 'bottom': '1.5cm', 'left': '1.5cm', 'right': '1.5cm'})
        browser.close()
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="Journey_Management_Plan.pdf"'
    return response


def journey_email(request, pk):
    from django.core.mail import EmailMessage as DjangoEmail
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)

    # Build recipient list: personnel in the plan who have an email in Personnel
    names = plan.personnel.values_list('name', flat=True)
    recipients = list(Personnel.objects.filter(name__in=names).exclude(email='').values('name', 'email'))

    if request.method == 'POST':
        extra = request.POST.get('extra_emails', '').strip()
        to_addrs = [r['email'] for r in recipients]
        if extra:
            to_addrs += [e.strip() for e in extra.replace(',', '\n').splitlines() if e.strip()]

        if not to_addrs:
            return render(request, 'reports/journey_email.html', {
                'plan': plan, 'recipients': recipients, 'error': 'No email addresses available.',
            })

        logo_b64 = _logo_b64()
        route_waypoints = _build_waypoints(plan.route_waypoints)
        html = render_to_string('reports/journey_pdf.html', {
            'plan': plan, 'logo_b64': logo_b64, 'route_waypoints': route_waypoints,
        })
        with sync_playwright() as p:
            browser = p.chromium.launch()
            pg = browser.new_page()
            pg.set_content(html, wait_until='networkidle')
            pdf_bytes = pg.pdf(format='A4', margin={'top':'1.5cm','bottom':'1.5cm','left':'1.5cm','right':'1.5cm'})
            browser.close()

        route_label = ''
        if plan.departing_from:
            route_label = plan.departing_from
            if plan.arriving_at:
                route_label += f' → {plan.arriving_at}'
        date_label = plan.depart_date.strftime('%d %b %Y').lstrip('0') if plan.depart_date else ''
        subject = f'Journey Management Plan — {route_label}' + (f' ({date_label})' if date_label else '')
        body = (
            f'Please find attached the Journey Management Plan for {route_label}.\n\n'
            f'Job: {plan.job.job_number} — {plan.job.project_name}\n'
        )
        if date_label:
            body += f'Departure: {date_label}\n'

        msg = DjangoEmail(subject=subject, body=body, to=to_addrs)
        filename = f'Journey_Management_Plan_{plan.depart_date.strftime("%d_%m_%Y") if plan.depart_date else "plan"}.pdf'
        msg.attach(filename, pdf_bytes, 'application/pdf')
        msg.send()

        return render(request, 'reports/journey_email.html', {
            'plan': plan, 'recipients': recipients, 'sent': True, 'to_addrs': to_addrs,
        })

    return render(request, 'reports/journey_email.html', {
        'plan': plan, 'recipients': recipients,
    })


# ── Toolbox Meetings ───────────────────────────────────────────────────────────

TSO_QUESTIONS = [
    "Was a Haz ID (Take 5 / SLAM) completed before performing the task?",
    "Were there any hazards identified prior to the task starting?",
    "What controls are in place to minimise the risk from these hazards?",
    "Are all crew member/s trained, qualified and competent to perform the task?",
    "Are the tools and/or equipment fit for purpose and in good order?",
    "Were tools and equipment being used correctly?",
    "Was the correct Isolation Procedure followed prior to commencing the task?",
    "Is the correct PPE being used?",
    "Was there good communication between team members?",
    "Was the housekeeping of the work area satisfactory to safely complete the task?",
    "Is there a task procedure in place for the task being observed?",
    "If Yes, are observed team members following the procedure safely?",
    "If no, has there been a change to the current process, procedure or design?",
    "If there has been a change, should a task procedure be developed for the task or will a task safety analysis / risk assessment be adequate?",
    "Are manual handling techniques being performed correctly for the task? (lifting/reaching/weight/carrying)",
]


def qhse_list(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    records = job.take5_records.prefetch_related('hazards').all()
    observations = job.task_observations.all()
    audits = job.infield_audits.all()
    jsas = job.jsas.all()
    return render(request, 'reports/qhse_list.html', {
        'job': job,
        'records': records,
        'observations': observations,
        'audits': audits,
        'jsas': jsas,
    })


def task_observation_create(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    obs = TaskSafetyObservation(job=job)
    obs.save()
    return redirect('task_observation_edit', pk=obs.pk)


def task_observation_edit(request, pk):
    from datetime import datetime
    obs = get_object_or_404(TaskSafetyObservation, pk=pk)
    if request.method == 'POST':
        obs.task_being_observed = request.POST.get('task_being_observed', '').strip()
        raw_date = request.POST.get('date', '').strip()
        try:
            obs.date = date.fromisoformat(raw_date) if raw_date else None
        except ValueError:
            obs.date = None
        raw_time = request.POST.get('time', '').strip()
        from datetime import time as dtime
        try:
            obs.time = dtime.fromisoformat(raw_time) if raw_time else None
        except ValueError:
            obs.time = None
        obs.observer = request.POST.get('observer', '').strip()
        obs.location = request.POST.get('location', '').strip()
        obs.team_members = [m.strip() for m in request.POST.getlist('team_member') if m.strip()]
        checklist = {}
        for i in range(len(TSO_QUESTIONS)):
            checklist[str(i)] = {
                'answer': request.POST.get(f'q_answer_{i}', ''),
                'comment': request.POST.get(f'q_comment_{i}', '').strip(),
            }
        obs.checklist = checklist
        obs.at_risk = request.POST.get('at_risk', '').strip()
        obs.discussion = request.POST.get('discussion', '').strip()
        obs.save()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True})
        return redirect('task_observation_edit', pk=obs.pk)
    checklist_data = [(i, q, obs.checklist.get(str(i), {'answer': '', 'comment': ''}))
                      for i, q in enumerate(TSO_QUESTIONS)]
    team_members = list(obs.team_members) + [''] * (10 - len(obs.team_members))
    return render(request, 'reports/task_observation_form.html', {
        'obs': obs,
        'job': obs.job,
        'checklist_data': checklist_data,
        'team_members': team_members[:10],
    })


def task_observation_delete(request, pk):
    obs = get_object_or_404(TaskSafetyObservation, pk=pk)
    job_pk = obs.job_id
    if request.method == 'POST':
        obs.delete()
    return redirect(reverse('qhse_list', args=[job_pk]) + '?tab=observations')


def task_observation_sign(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    from datetime import datetime
    obs = get_object_or_404(TaskSafetyObservation, pk=pk)
    data = json.loads(request.body)
    sig = data.get('signature', '').strip()
    if not sig:
        return JsonResponse({'error': 'No signature data'}, status=400)
    obs.signature = sig
    try:
        obs.signed_at = datetime.fromisoformat(data.get('signed_at', ''))
    except (ValueError, TypeError):
        from django.utils import timezone
        obs.signed_at = timezone.now()
    obs.save()
    return JsonResponse({'ok': True, 'signed_at': obs.signed_at.strftime('%d %b %Y %H:%M')})


def task_observation_sign_clear(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    obs = get_object_or_404(TaskSafetyObservation, pk=pk)
    obs.signature = ''
    obs.signed_at = None
    obs.save()
    return JsonResponse({'ok': True})


def task_observation_pdf(request, pk):
    obs = get_object_or_404(TaskSafetyObservation, pk=pk)
    checklist_data = [(i, q, obs.checklist.get(str(i), {'answer': '', 'comment': ''}))
                      for i, q in enumerate(TSO_QUESTIONS)]
    team_members = [m for m in obs.team_members if m]
    html = render_to_string('reports/task_observation_pdf.html', {
        'obs': obs,
        'checklist_data': checklist_data,
        'team_members': team_members,
    }, request=request)
    from weasyprint import HTML as WeasyprintHTML
    pdf = WeasyprintHTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    resp = HttpResponse(pdf, content_type='application/pdf')
    filename = f"Task_Observation_{obs.job.job_number}_{obs.date or 'nodate'}.pdf"
    inline = request.GET.get('inline', '0') == '1'
    resp['Content-Disposition'] = f'{"inline" if inline else "attachment"}; filename="{filename}"'
    return resp


INFIELD_AUDIT_SECTIONS = [
    ('1', 'House Keeping', [
        ('1a', 'Is there evidence of rubbish in the vehicles? Is any rubbish stored in appropriate receptacles? Is it secured against animals getting into the rubbish?'),
        ('1b', 'Are load restraints being used in all vehicles?'),
        ('1c', 'Is there rubbish on the grid/laydown/access tracks (markers, food wrappers, flagging tape etc)?'),
        ('1d', 'Is there any contaminated waste? Has it been cleaned up with a spill kit (fuel, hydraulic fluid etc)? Is it secured in a vehicle for disposal off site?'),
    ]),
    ('2', 'PPE', [
        ('2a', 'Are crew wearing correct, good condition PPE as per Site Regulations: Hard Hats with Brims, Long Sleeved Shirts and long pants, Steel Capped Shoes, Eye Protection (AS/NZS 1337), Gloves and glove belt clips, Hearing Protection (if applicable)?'),
        ('2b', 'Are Crew wearing/have access to isolation locks as per Site Regulations?'),
    ]),
    ('3', 'First Aid', [
        ('3a', 'Are Tick Kits available to each person? Does the tick kit contain: Antiseptic cream, Band-aids, Tick key/tweezers, Alcohol swabs, Antimicrobial soap, Instruction sheet?'),
        ('3b', 'Is there a first aid kit in each vehicle? Has the First aid kit been checked off on all vehicle prestarts? What is the expiry date on the kit for each vehicle?'),
        ('3c', 'Is Sunscreen available in each vehicle?'),
        ('3d', 'Is there a Snakebite Kit in each vehicle? Does the Service Truck and Dog box have a trauma kit each? What are the expiry dates?'),
        ('3e', 'Have there been any incidents requiring first aid treatment in the current job?'),
        ('3f', 'Has an Emergency Response/Drill been conducted for this job? If no, perform an Emergency Drill prior to finishing job.'),
        ('3g', 'Has the drug and allergy record been filled out for each crew member?'),
    ]),
    ('4', 'Emergency Response', [
        ('4a', 'Where are the Emergency muster points for this job? Have they been identified and sign posted? Is it a client sign or a Velseis sign?'),
        ('4b', 'Are Emergency Response Plans readily available? Do the crew know where to access this information?'),
        ('4c', 'Are the Crew aware of Muster Points and EMP? Where/when was it discussed?'),
        ('4d', 'Does each service truck and Envirovibe have a Spill Kit available?'),
    ]),
    ('5', 'Electrical Equipment', [
        ('5a', 'Are electrical items currently in test and tag, including the RCD?'),
        ('5b', 'Have visual inspections been made of any electrical equipment used onsite?'),
        ('5c', 'Are there any defective batteries? Are they tagged out to take off site?'),
    ]),
]


def infield_audit_create(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    audit = InfieldAudit(job=job)
    audit.save()
    return redirect('infield_audit_edit', pk=audit.pk)


def infield_audit_edit(request, pk):
    from datetime import datetime
    audit = get_object_or_404(InfieldAudit, pk=pk)
    if request.method == 'POST':
        audit.date_commenced = _parse_date(request.POST.get('date_commenced', ''))
        audit.date_completed = _parse_date(request.POST.get('date_completed', ''))
        audit.site_location = request.POST.get('site_location', '').strip()
        audit.audit_conducted_by = request.POST.get('audit_conducted_by', '').strip()
        audit.crew_supervisor = request.POST.get('crew_supervisor', '').strip()
        checklist = {}
        for _, _, items in INFIELD_AUDIT_SECTIONS:
            for item_id, _ in items:
                checklist[item_id] = {
                    'answer': request.POST.get(f'item_answer_{item_id}', ''),
                    'comment': request.POST.get(f'item_comment_{item_id}', '').strip(),
                    'date': request.POST.get(f'item_date_{item_id}', '').strip(),
                }
        audit.checklist = checklist
        # Actions table
        observations = request.POST.getlist('action_observation')
        recommendations = request.POST.getlist('action_recommendations')
        priorities = request.POST.getlist('action_priority')
        responsibilities = request.POST.getlist('action_responsibility')
        due_dates = request.POST.getlist('action_due_date')
        completed_by = request.POST.getlist('action_completed_by')
        action_dates = request.POST.getlist('action_date')
        actions = []
        for i, obs_text in enumerate(observations):
            if any([obs_text.strip(), recommendations[i].strip() if i < len(recommendations) else '']):
                actions.append({
                    'observation': obs_text.strip(),
                    'recommendations': recommendations[i].strip() if i < len(recommendations) else '',
                    'priority': priorities[i].strip() if i < len(priorities) else '',
                    'responsibility': responsibilities[i].strip() if i < len(responsibilities) else '',
                    'due_date': due_dates[i].strip() if i < len(due_dates) else '',
                    'completed_by': completed_by[i].strip() if i < len(completed_by) else '',
                    'date': action_dates[i].strip() if i < len(action_dates) else '',
                })
        audit.actions = actions
        audit.save()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True})
        return redirect('infield_audit_edit', pk=audit.pk)

    sections_data = []
    for sec_id, sec_title, items in INFIELD_AUDIT_SECTIONS:
        rows = [(item_id, question, audit.checklist.get(item_id, {'answer': '', 'comment': '', 'date': ''}))
                for item_id, question in items]
        sections_data.append((sec_id, sec_title, rows))

    # Ensure at least 3 empty action rows
    actions = list(audit.actions)
    while len(actions) < 3:
        actions.append({'observation': '', 'recommendations': '', 'priority': '', 'responsibility': '', 'due_date': '', 'completed_by': '', 'date': ''})

    return render(request, 'reports/infield_audit_form.html', {
        'audit': audit,
        'job': audit.job,
        'sections_data': sections_data,
        'actions': actions,
    })


def _parse_date(s):
    try:
        return date.fromisoformat(s) if s else None
    except ValueError:
        return None


def infield_audit_delete(request, pk):
    audit = get_object_or_404(InfieldAudit, pk=pk)
    job_pk = audit.job_id
    if request.method == 'POST':
        audit.delete()
    return redirect(reverse('qhse_list', args=[job_pk]) + '?tab=audits')


def infield_audit_sign(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    from datetime import datetime
    audit = get_object_or_404(InfieldAudit, pk=pk)
    data = json.loads(request.body)
    sig = data.get('signature', '').strip()
    if not sig:
        return JsonResponse({'error': 'No signature data'}, status=400)
    audit.signature = sig
    try:
        audit.signed_at = datetime.fromisoformat(data.get('signed_at', ''))
    except (ValueError, TypeError):
        from django.utils import timezone
        audit.signed_at = timezone.now()
    audit.save()
    return JsonResponse({'ok': True, 'signed_at': audit.signed_at.strftime('%d %b %Y %H:%M')})


def infield_audit_sign_clear(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    audit = get_object_or_404(InfieldAudit, pk=pk)
    audit.signature = ''
    audit.signed_at = None
    audit.save()
    return JsonResponse({'ok': True})


def infield_audit_pdf(request, pk):
    audit = get_object_or_404(InfieldAudit, pk=pk)
    sections_data = []
    for sec_id, sec_title, items in INFIELD_AUDIT_SECTIONS:
        rows = [(item_id, question, audit.checklist.get(item_id, {'answer': '', 'comment': '', 'date': ''}))
                for item_id, question in items]
        sections_data.append((sec_id, sec_title, rows))
    html = render_to_string('reports/infield_audit_pdf.html', {
        'audit': audit,
        'sections_data': sections_data,
    }, request=request)
    from weasyprint import HTML as WeasyprintHTML
    pdf = WeasyprintHTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    resp = HttpResponse(pdf, content_type='application/pdf')
    filename = f"Infield_Audit_{audit.job.job_number}_{audit.date_commenced or 'nodate'}.pdf"
    inline = request.GET.get('inline', '0') == '1'
    resp['Content-Disposition'] = f'{"inline" if inline else "attachment"}; filename="{filename}"'
    return resp


# ── JSA ──────────────────────────────────────────────────────────────────────

def jsa_create(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    jsa = JSA(job=job)
    jsa.analysis = [{'job_step': '', 'hazard': '', 'control': '', 'person_responsible': '', 'managed': ''} for _ in range(14)]
    jsa.participants = [{'name': '', 'position': '', 'years_exp': '', 'signature': ''} for _ in range(8)]
    jsa.save()
    return redirect('jsa_edit', pk=jsa.pk)


def jsa_edit(request, pk):
    from datetime import datetime
    jsa = get_object_or_404(JSA, pk=pk)
    if request.method == 'POST':
        jsa.jsa_name = request.POST.get('jsa_name', '').strip()
        jsa.reference_tp = request.POST.get('reference_tp', '').strip()
        jsa.reference_tra = request.POST.get('reference_tra', '').strip()
        jsa.project = request.POST.get('project', '').strip()
        jsa.site = request.POST.get('site', '').strip()
        jsa.date = _parse_date(request.POST.get('date', ''))
        raw_time = request.POST.get('time', '').strip()
        try:
            jsa.time = datetime.strptime(raw_time, '%H:%M').time() if raw_time else None
        except ValueError:
            jsa.time = None
        jsa.job_description = request.POST.get('job_description', '').strip()
        jsa.tools_equipment = request.POST.get('tools_equipment', '').strip()
        jsa.ppe_required = request.POST.get('ppe_required', '').strip()
        jsa.permits_approvals = request.POST.get('permits_approvals', '').strip()
        jsa.approver_name = request.POST.get('approver_name', '').strip()
        jsa.approver_position = request.POST.get('approver_position', '').strip()
        jsa.approval_date = _parse_date(request.POST.get('approval_date', ''))

        # Participants
        import json as _json
        try:
            jsa.participants = _json.loads(request.POST.get('participants_json', '[]'))
        except (ValueError, TypeError):
            pass

        # Analysis rows
        try:
            jsa.analysis = _json.loads(request.POST.get('analysis_json', '[]'))
        except (ValueError, TypeError):
            pass

        # Photos (base64 data URLs from file input JS)
        p1 = request.POST.get('photo1_data', '').strip()
        p2 = request.POST.get('photo2_data', '').strip()
        if p1:
            jsa.photo1 = p1
        if p2:
            jsa.photo2 = p2

        jsa.save()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True})
        return redirect('jsa_edit', pk=jsa.pk)

    # Ensure minimum rows
    while len(jsa.analysis) < 14:
        jsa.analysis.append({'job_step': '', 'hazard': '', 'control': '', 'person_responsible': '', 'managed': ''})
    while len(jsa.participants) < 8:
        jsa.participants.append({'name': '', 'position': '', 'years_exp': '', 'signature': ''})

    return render(request, 'reports/jsa_form.html', {'jsa': jsa, 'job': jsa.job})


def jsa_delete(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    jsa = get_object_or_404(JSA, pk=pk)
    job_pk = jsa.job_id
    jsa.delete()
    return redirect('qhse_list', job_pk=job_pk)


def jsa_sign(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    import json as _json
    from datetime import datetime
    jsa = get_object_or_404(JSA, pk=pk)
    data = _json.loads(request.body)
    jsa.approver_signature = data.get('signature', '')
    try:
        jsa.approval_signed_at = datetime.fromisoformat(data.get('signed_at', ''))
    except (ValueError, TypeError):
        from django.utils import timezone
        jsa.approval_signed_at = timezone.now()
    jsa.save()
    return JsonResponse({'ok': True, 'signed_at': jsa.approval_signed_at.strftime('%d %b %Y %H:%M')})


def jsa_sign_clear(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    jsa = get_object_or_404(JSA, pk=pk)
    jsa.approver_signature = ''
    jsa.approval_signed_at = None
    jsa.save()
    return JsonResponse({'ok': True})


def jsa_pdf(request, pk):
    jsa = get_object_or_404(JSA, pk=pk)
    # Pair up participants for 2-column PDF layout
    parts = jsa.participants or []
    while len(parts) < 8:
        parts.append({'name': '', 'position': '', 'years_exp': '', 'signature': ''})
    participant_pairs = [(parts[i], parts[i + 4]) for i in range(4)]
    # Only include analysis rows that have content
    analysis_rows = [r for r in (jsa.analysis or []) if any(r.get(k) for k in ('job_step', 'hazard', 'control', 'person_responsible'))]
    html = render_to_string('reports/jsa_pdf.html', {
        'jsa': jsa,
        'participant_pairs': participant_pairs,
        'analysis_rows': analysis_rows,
    }, request=request)
    from weasyprint import HTML as WeasyprintHTML
    pdf = WeasyprintHTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    resp = HttpResponse(pdf, content_type='application/pdf')
    filename = f"JSA_{jsa.job.job_number}_{jsa.jsa_name or 'noname'}_{jsa.date or 'nodate'}.pdf"
    inline = request.GET.get('inline', '0') == '1'
    resp['Content-Disposition'] = f'{"inline" if inline else "attachment"}; filename="{filename}"'
    return resp


def toolbox_list(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    meetings = job.toolbox_meetings.prefetch_related('attendees').all()
    v2_meetings = job.toolbox_v2_meetings.prefetch_related('attendees').all()
    return render(request, 'reports/toolbox_list.html', {'job': job, 'meetings': meetings, 'v2_meetings': v2_meetings})


def _toolbox_alloc_data(job):
    """Build {tab: [{name, vehicle}]} from VehicleAllocation for a job."""
    result = {}
    for tab in ('mob', 'job', 'demob', 'other'):
        entries = []
        for jv in job.vehicles.prefetch_related('allocations').all():
            for alloc in jv.allocations.filter(tab=tab):
                entries.append({'name': alloc.person_name, 'vehicle': jv.name})
        result[tab] = entries
    return result


def _toolbox_eq_alloc_data(job):
    """Build [{person_name, equipment:[{name, equipment_type}]}] from EquipmentAllocation for a job."""
    from collections import defaultdict
    by_person = defaultdict(list)
    for alloc in EquipmentAllocation.objects.filter(job_equipment__job=job).select_related('job_equipment'):
        eq = alloc.job_equipment
        by_person[alloc.person_name].append({'name': eq.name, 'equipment_type': eq.equipment_type})
    return [{'person_name': name, 'equipment': eqs} for name, eqs in sorted(by_person.items())]


def _save_toolbox(request, meeting):
    """Save all fields from a toolbox form POST onto meeting (and attendee rows)."""
    def d(field): return request.POST.get(field, '').strip() or None

    meeting.date = d('date')
    meeting.time = d('time')
    meeting.location = request.POST.get('location', '').strip()
    meeting.supervisor = request.POST.get('supervisor', '').strip()
    meeting.days_on_job = request.POST.get('days_on_job', '').strip()
    meeting.jmp_number = request.POST.get('jmp_number', '').strip()

    meeting.yesterday_activities = request.POST.get('yesterday_activities', '').strip()
    meeting.jmp_route_discussed = 'jmp_route_discussed' in request.POST
    meeting.jmp_hours_noted = 'jmp_hours_noted' in request.POST
    meeting.jmp_contact_numbers = 'jmp_contact_numbers' in request.POST
    meeting.jmp_signed_off = 'jmp_signed_off' in request.POST
    meeting.jmp_lead_tail = 'jmp_lead_tail' in request.POST

    meeting.todays_activities = request.POST.get('todays_activities', '').strip()
    meeting.terrain_discussion = request.POST.get('terrain_discussion', '').strip()
    meeting.road_condition = request.POST.get('road_condition', '').strip()
    meeting.muster_point = request.POST.get('muster_point', '').strip()
    meeting.forecast = request.POST.get('forecast', '').strip()
    meeting.uv_index = request.POST.get('uv_index', '').strip()
    meeting.chance_of_rain = request.POST.get('chance_of_rain', '').strip()
    meeting.min_temp = request.POST.get('min_temp', '').strip()
    meeting.max_temp = request.POST.get('max_temp', '').strip()
    meeting.grass_fire = request.POST.get('grass_fire', '').strip()
    meeting.forest_fire = request.POST.get('forest_fire', '').strip()
    meeting.wind = request.POST.get('wind', '').strip()
    meeting.other_topics = request.POST.get('other_topics', '').strip()
    meeting.include_equipment_allocation = 'include_equipment_allocation' in request.POST
    meeting.save()

    if meeting.location:
        JobLocation.objects.get_or_create(job=meeting.job, name=meeting.location)
    if meeting.muster_point:
        JobMusterPoint.objects.get_or_create(job=meeting.job, name=meeting.muster_point)

    meeting.attendees.all().delete()
    names = request.POST.getlist('a_name')
    roles = request.POST.getlist('a_role')
    vehicles = request.POST.getlist('a_vehicle')
    bacs = request.POST.getlist('a_bac')
    sigs = request.POST.getlist('a_signature')
    is_seps = request.POST.getlist('a_is_sep')
    non_sep_idx = 0  # separate counter for lists that exclude separator rows
    for i, name in enumerate(names):
        is_sep = is_seps[i] == '1' if i < len(is_seps) else False
        if is_sep:
            job_role, vehicle, bac, signature = '', '', '', ''
        else:
            job_role   = roles[non_sep_idx].strip()    if non_sep_idx < len(roles)    else ''
            vehicle    = vehicles[non_sep_idx].strip() if non_sep_idx < len(vehicles) else ''
            bac        = bacs[non_sep_idx].strip()     if non_sep_idx < len(bacs)     else ''
            signature  = sigs[non_sep_idx].strip()     if non_sep_idx < len(sigs)     else ''
            non_sep_idx += 1
        ToolboxAttendee.objects.create(
            meeting=meeting,
            is_separator=is_sep,
            name=name.strip(),
            job_role=job_role,
            vehicle=vehicle,
            bac=bac,
            signature=signature,
            order=i,
        )


def toolbox_create(request, job_pk, meeting_type):
    job = get_object_or_404(Job, pk=job_pk)
    if meeting_type not in ('daily', 'jmp'):
        return redirect('toolbox_list', job_pk=job_pk)
    if request.method == 'POST':
        meeting = ToolboxMeeting(job=job, meeting_type=meeting_type)
        _save_toolbox(request, meeting)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': meeting.pk, 'edit_url': reverse('toolbox_edit', args=[meeting.pk])})
        return redirect('toolbox_list', job_pk=job_pk)
    jp = list(job.personnel.order_by('name'))
    supervisors = list(job.personnel.filter(roles__role__icontains='supervisor').distinct())
    previous = (ToolboxMeeting.objects.filter(job=job, meeting_type=meeting_type)
                .order_by('-date', '-pk').first())
    import json as _json
    prev_data = None
    if previous:
        prev_data = _json.dumps({
            'location':             previous.location or '',
            'supervisor':           previous.supervisor or '',
            'muster_point':         previous.muster_point or '',
            'terrain_discussion':   previous.terrain_discussion or '',
            'yesterday_activities': previous.todays_activities or '',
            'other_topics':         previous.other_topics or '',
            'days_on_job':          previous.days_on_job or '',
            'jmp_number':           previous.jmp_number or '',
            'road_condition':       previous.road_condition or '',
            'prev_date':            previous.date.strftime('%d/%m/%Y') if previous.date else '',
            'time':                 previous.time.strftime('%H:%M') if previous.time else '',
        })
    jmp_numbers = list(job.journey_plans.exclude(plan_number='').order_by('-created_at').values_list('plan_number', flat=True)[:20])
    personnel_roles_json = json.dumps({p.name: ', '.join(r.role for r in p.roles.all()) for p in jp})
    job_vehicle_names = list(job.vehicles.values_list('name', flat=True))
    return render(request, 'reports/toolbox_form.html', {
        'job': job,
        'meeting': None,
        'meeting_type': meeting_type,
        'job_personnel': jp,
        'job_personnel_names': {p.name for p in jp},
        'supervisors': supervisors,
        'saved_locations': job.saved_locations.all(),
        'saved_muster_points': job.saved_muster_points.all(),
        'topic_templates': ToolboxTopicTemplate.objects.all(),
        'prev_data': prev_data,
        'jmp_numbers': jmp_numbers,
        'personnel_roles_json': personnel_roles_json,
        'job_vehicle_names': job_vehicle_names,
        'alloc_data_json': json.dumps(_toolbox_alloc_data(job)),
        'eq_alloc_data_json': json.dumps(_toolbox_eq_alloc_data(job)),
    })


def toolbox_edit(request, pk):
    meeting = get_object_or_404(ToolboxMeeting, pk=pk)
    if request.method == 'POST':
        _save_toolbox(request, meeting)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': meeting.pk})
        return redirect('toolbox_edit', pk=meeting.pk)
    jp = list(meeting.job.personnel.all())
    supervisors = list(meeting.job.personnel.filter(roles__role__icontains='supervisor').distinct())
    personnel_roles_json = json.dumps({p.name: ', '.join(r.role for r in p.roles.all()) for p in jp})
    job_vehicle_names = list(meeting.job.vehicles.values_list('name', flat=True))
    return render(request, 'reports/toolbox_form.html', {
        'job': meeting.job,
        'meeting': meeting,
        'meeting_type': meeting.meeting_type,
        'job_personnel': jp,
        'job_personnel_names': {p.name for p in jp},
        'supervisors': supervisors,
        'saved_locations': meeting.job.saved_locations.all(),
        'saved_muster_points': meeting.job.saved_muster_points.all(),
        'topic_templates': ToolboxTopicTemplate.objects.all(),
        'personnel_roles_json': personnel_roles_json,
        'job_vehicle_names': job_vehicle_names,
        'alloc_data_json': json.dumps(_toolbox_alloc_data(meeting.job)),
        'eq_alloc_data_json': json.dumps(_toolbox_eq_alloc_data(meeting.job)),
        'jmp_numbers': list(meeting.job.journey_plans.exclude(plan_number='').order_by('-created_at').values_list('plan_number', flat=True)[:20]),
        'photos': meeting.photos.all(),
    })


def toolbox_working_personnel(request, job_pk):
    from django.http import JsonResponse
    job = get_object_or_404(Job, pk=job_pk)
    date_str = request.GET.get('date', '')
    if not date_str:
        return JsonResponse({'names': []})
    entries = PersonnelScheduleEntry.objects.filter(
        job_personnel__job=job, date=date_str, status__iexact='working'
    ).select_related('job_personnel')
    names = [e.job_personnel.name for e in entries]
    return JsonResponse({'names': names})


def toolbox_delete(request, pk):
    meeting = get_object_or_404(ToolboxMeeting, pk=pk)
    job_pk = meeting.job_id
    if request.method == 'POST':
        meeting.delete()
    return redirect('toolbox_list', job_pk=job_pk)


def toolbox_copy(request, pk):
    src = get_object_or_404(ToolboxMeeting, pk=pk)
    new_meeting = ToolboxMeeting(
        job=src.job,
        meeting_type=src.meeting_type,
        location=src.location,
        supervisor=src.supervisor,
        days_on_job=src.days_on_job,
        jmp_number=src.jmp_number,
        yesterday_activities=src.yesterday_activities,
        jmp_route_discussed=src.jmp_route_discussed,
        jmp_hours_noted=src.jmp_hours_noted,
        jmp_contact_numbers=src.jmp_contact_numbers,
        jmp_signed_off=src.jmp_signed_off,
        jmp_lead_tail=src.jmp_lead_tail,
        todays_activities=src.todays_activities,
        terrain_discussion=src.terrain_discussion,
        road_condition=src.road_condition,
        muster_point=src.muster_point,
        forecast=src.forecast,
        uv_index=src.uv_index,
        chance_of_rain=src.chance_of_rain,
        min_temp=src.min_temp,
        max_temp=src.max_temp,
        grass_fire=src.grass_fire,
        forest_fire=src.forest_fire,
        wind=src.wind,
        other_topics=src.other_topics,
        include_equipment_allocation=src.include_equipment_allocation,
    )
    new_meeting.save()
    for a in src.attendees.all():
        ToolboxAttendee.objects.create(
            meeting=new_meeting,
            is_separator=a.is_separator,
            name=a.name,
            job_role=a.job_role,
            vehicle=a.vehicle,
            order=a.order,
        )
    return redirect('toolbox_edit', pk=new_meeting.pk)


def toolbox_pdf(request, pk):
    meeting = get_object_or_404(ToolboxMeeting, pk=pk)
    logo_b64 = _logo_b64()
    attendees = list(meeting.attendees.order_by('order', 'pk'))
    has_extended = any(a.job_role or a.vehicle for a in attendees if not a.is_separator)
    has_separator = any(a.is_separator for a in attendees)
    # Pad to even number for standard 2-col layout (only non-separator rows)
    non_sep = [a for a in attendees if not a.is_separator]
    padded = non_sep[:]
    min_slots = max(len(padded), 12)
    if min_slots % 2: min_slots += 1
    padded += [None] * (min_slots - len(padded))
    attendee_pairs = [(padded[i], padded[i+1]) for i in range(0, min_slots, 2)]
    jmp_items = [
        ('Route discussed with all crew', meeting.jmp_route_discussed),
        ('Driving hours / fatigue limits noted', meeting.jmp_hours_noted),
        ('Emergency contact numbers confirmed', meeting.jmp_contact_numbers),
        ('JMP signed off by supervisor', meeting.jmp_signed_off),
        ('Lead / tail vehicle assigned', meeting.jmp_lead_tail),
    ]
    from pathlib import Path
    photos_with_uri = [
        {'photo': p, 'uri': Path(p.image.path).as_uri()}
        for p in meeting.photos.all()
    ]
    html = render_to_string('reports/toolbox_pdf.html', {
        'meeting': meeting,
        'logo_b64': logo_b64,
        'attendees': attendees,
        'attendee_pairs': attendee_pairs,
        'has_extended': has_extended,
        'has_separator': has_separator,
        'jmp_items': jmp_items,
        'photos_with_uri': photos_with_uri,
        'eq_alloc_data': _toolbox_eq_alloc_data(meeting.job) if meeting.include_equipment_allocation else [],
    })
    import tempfile, os as _os
    form_ref = 'VF_GEN_001: Daily Toolbox Meeting Record' if meeting.meeting_type == 'daily' else 'VF_GEN_023: Journey Management Toolbox Meeting'
    footer = f'''
    <div style="font-family:Arial,sans-serif; font-size:6pt; color:#444; width:100%;
                padding:3px 1.5cm 0; border-top:0.5pt solid #bbb;
                display:flex; justify-content:space-between; align-items:flex-start;">
        <div>
            <div style="margin-bottom:1px;">Velseis Form {form_ref}</div>
            <table style="border-collapse:collapse; font-size:6pt;">
                <tr>
                    <th style="border:0.5pt solid #999;padding:1px 4px;background:#ddd;">Version</th>
                    <th style="border:0.5pt solid #999;padding:1px 4px;background:#ddd;">Reviewed by</th>
                    <th style="border:0.5pt solid #999;padding:1px 4px;background:#ddd;">Implementation Date</th>
                    <th style="border:0.5pt solid #999;padding:1px 4px;background:#ddd;">Review Date</th>
                    <th style="border:0.5pt solid #999;padding:1px 4px;background:#ddd;">Authorization</th>
                </tr>
                <tr>
                    <td style="border:0.5pt solid #999;padding:1px 4px;text-align:center;">1.0</td>
                    <td style="border:0.5pt solid #999;padding:1px 4px;text-align:center;">D. Tucker</td>
                    <td style="border:0.5pt solid #999;padding:1px 4px;text-align:center;">22/06/2018</td>
                    <td style="border:0.5pt solid #999;padding:1px 4px;text-align:center;">22/06/2021</td>
                    <td style="border:0.5pt solid #999;padding:1px 4px;text-align:center;">K. Dummett</td>
                </tr>
            </table>
        </div>
        <div style="white-space:nowrap; font-size:7pt; padding-top:2px;">
            Page <span class="pageNumber"></span> of <span class="totalPages"></span>
        </div>
    </div>'''
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            pg = browser.new_page()
            pg.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            pg.wait_for_load_state('networkidle')
            pdf_bytes = pg.pdf(format='A4', print_background=True,
                               display_header_footer=True,
                               header_template='<div></div>',
                               footer_template=footer)
            browser.close()
    finally:
        _os.unlink(tmp.name)
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    label = 'Daily_Toolbox' if meeting.meeting_type == 'daily' else 'JMP_Toolbox'
    date_str = meeting.date.strftime('%d_%m_%Y') if meeting.date else 'undated'
    disposition = 'attachment' if request.GET.get('download') else 'inline'
    response['Content-Disposition'] = f'{disposition}; filename="{label}_{date_str}.pdf"'
    return response


def toolbox_word(request, pk):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    meeting = get_object_or_404(ToolboxMeeting, pk=pk)
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # ── Helpers ───────────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x1a, 0x3a, 0x5c)
    LIGHT_BLUE = RGBColor(0xee, 0xf3, 0xf8)

    def sec_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1A3A5C')
        p._p.get_or_add_pPr().append(shading)
        return p

    def add_row(table, label, value, label_width=None):
        row = table.add_row()
        lc = row.cells[0]
        vc = row.cells[1]
        lc.text = label
        vc.text = str(value) if value else ''
        lc.paragraphs[0].runs[0].bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(8.5)
        vc.paragraphs[0].runs[0].font.size = Pt(8.5)
        lc._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
        return row

    def _cell_shading(fill_hex):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        return shd

    def set_col_width(table, col_idx, width_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width_cm)

    def add_text_box(text):
        p = doc.add_paragraph(text or '')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        if p.runs: p.runs[0].font.size = Pt(8.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), 'BBBBBB')
            pBdr.append(b)
        pPr.append(pBdr)

    def cb_text(value, choices):
        parts = []
        for val, label in choices:
            tick = '☑' if value == val else '☐'
            parts.append(f'{tick} {label}')
        return '   '.join(parts)

    # ── Logo ──────────────────────────────────────────────────────────────
    logo_path = os.path.join(settings.MEDIA_ROOT, 'velseis-logo.png')
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.5))
        p.paragraph_format.space_after = Pt(2)

    # ── Title ─────────────────────────────────────────────────────────────
    title_text = 'Daily Toolbox Meeting Record' if meeting.meeting_type == 'daily' else 'Journey Management Toolbox Meeting'
    tp = doc.add_paragraph(title_text)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.runs[0].bold = True
    tp.runs[0].font.size = Pt(13)
    tp.paragraph_format.space_after = Pt(6)

    # ── Section 1 ─────────────────────────────────────────────────────────
    sec_heading('Section 1: Job Details')
    t1 = doc.add_table(rows=0, cols=4)
    t1.style = 'Table Grid'
    r1 = t1.add_row()
    for i, (lbl, val) in enumerate([('Date', meeting.date.strftime('%d/%m/%Y') if meeting.date else ''),
                                     ('Time', meeting.time.strftime('%H:%M') if meeting.time else '')]):
        r1.cells[i*2].text = lbl
        r1.cells[i*2+1].text = val
        r1.cells[i*2].paragraphs[0].runs[0].bold = True
        r1.cells[i*2]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
    r2 = t1.add_row()
    r2.cells[0].text = 'Location'
    r2.cells[0].paragraphs[0].runs[0].bold = True
    r2.cells[0]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
    r2.cells[1].merge(r2.cells[3])
    r2.cells[1].text = meeting.location or ''
    r3 = t1.add_row()
    r3.cells[0].text = 'Job Number'
    r3.cells[0].paragraphs[0].runs[0].bold = True
    r3.cells[0]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
    r3.cells[1].text = meeting.job.job_number
    r3.cells[2].text = '# Days on Job' if meeting.meeting_type == 'daily' else 'JMP #'
    r3.cells[2].paragraphs[0].runs[0].bold = True
    r3.cells[2]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
    r3.cells[3].text = (meeting.days_on_job or '') if meeting.meeting_type == 'daily' else (meeting.jmp_number or '')
    r4 = t1.add_row()
    r4.cells[0].text = 'Supervisor'
    r4.cells[0].paragraphs[0].runs[0].bold = True
    r4.cells[0]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
    r4.cells[1].merge(r4.cells[3])
    r4.cells[1].text = meeting.supervisor or ''
    for row in t1.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)

    # ── Section 2 ─────────────────────────────────────────────────────────
    if meeting.meeting_type == 'daily':
        sec_heading('Section 2: General Job Information / Discussion for Yesterday\'s Activities')
        add_text_box(meeting.yesterday_activities)
    else:
        sec_heading('Section 2: JMP Checklist')
        items = [
            ('Route discussed with all crew', meeting.jmp_route_discussed),
            ('Driving hours / fatigue limits noted', meeting.jmp_hours_noted),
            ('Emergency contact numbers confirmed', meeting.jmp_contact_numbers),
            ('JMP signed off by supervisor', meeting.jmp_signed_off),
            ('Lead / tail vehicle assigned', meeting.jmp_lead_tail),
        ]
        for label, checked in items:
            p = doc.add_paragraph(f'{"☑" if checked else "☐"}  {label}')
            p.runs[0].font.size = Pt(8.5)

    # ── Section 3 ─────────────────────────────────────────────────────────
    sec_heading('Section 3: Today\'s Activities & Conditions' if meeting.meeting_type == 'daily' else 'Section 3: Road Conditions & Activities')
    if meeting.meeting_type == 'daily':
        p = doc.add_paragraph('What activities do we have today?')
        p.runs[0].bold = True; p.runs[0].font.size = Pt(8.5)
        add_text_box(meeting.todays_activities)
        p2 = doc.add_paragraph('Any terrain discussions? (hand carry, cultural heritage, etc)')
        p2.runs[0].bold = True; p2.runs[0].font.size = Pt(8.5)
        add_text_box(meeting.terrain_discussion)
    else:
        p = doc.add_paragraph('Road Condition Discussion')
        p.runs[0].bold = True; p.runs[0].font.size = Pt(8.5)
        add_text_box(meeting.road_condition)
        p2 = doc.add_paragraph('Today\'s Activities')
        p2.runs[0].bold = True; p2.runs[0].font.size = Pt(8.5)
        add_text_box(meeting.todays_activities)

    tw = doc.add_table(rows=0, cols=2)
    tw.style = 'Table Grid'
    for lbl, val in [
        ('Today\'s Emergency Muster Point', meeting.muster_point or ''),
        ('Forecast', meeting.forecast or ''),
        ('UV Index', cb_text(meeting.uv_index, [('low','Low'),('high','High'),('extreme','Extreme')])),
        ('Chance of Rain', meeting.chance_of_rain or ''),
        ('Min Temp', meeting.min_temp or ''),
        ('Max Temp', meeting.max_temp or ''),
        ('Grass Fire Danger', cb_text(meeting.grass_fire, [('low','Low'),('high','High'),('extreme','Extreme')])),
        ('Forest Fire Danger', cb_text(meeting.forest_fire, [('low','Low'),('high','High'),('extreme','Extreme')])),
        ('Wind Speed & Direction', meeting.wind or ''),
    ]:
        row = tw.add_row()
        row.cells[0].text = lbl
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)
        row.cells[0]._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
        row.cells[1].text = val
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)
    set_col_width(tw, 0, 4.5)

    # ── Section 4 ─────────────────────────────────────────────────────────
    sec_heading('Section 4: Other Topics for Crew Discussion')
    add_text_box(meeting.other_topics)

    # ── Section 5: Attendees ───────────────────────────────────────────────
    sec_heading('Section 5: Attendees')
    if meeting.meeting_type == 'jmp':
        p = doc.add_paragraph('Each attendee acknowledges the JMP has been communicated to them.')
        p.runs[0].font.size = Pt(8)
    attendees = sorted(meeting.attendees.all(), key=lambda a: a.name.lower())
    has_extended = any(a.job_role or a.vehicle for a in attendees)

    def _hdr_cell(cell, txt):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(_cell_shading('1A3A5C'))

    if has_extended:
        ta = doc.add_table(rows=1, cols=5)
        ta.style = 'Table Grid'
        for i, txt in enumerate(['Name', 'Job Role(s)', 'Vehicle', 'BAC', 'Signature']):
            _hdr_cell(ta.rows[0].cells[i], txt)
        min_slots = max(len(attendees), 12)
        padded = attendees + [None] * (min_slots - len(attendees))
        for a in padded:
            row = ta.add_row()
            row.cells[0].text = a.name if a else ''
            row.cells[1].text = a.job_role if a else ''
            row.cells[2].text = a.vehicle if a else ''
            row.cells[3].text = a.bac if a else ''
            row.cells[4].text = ''
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8.5)
        for col_i, w in enumerate([4.5, 3.5, 3.0, 1.5, 4.0]):
            set_col_width(ta, col_i, w)
    else:
        min_slots = max(len(attendees), 12)
        if min_slots % 2: min_slots += 1
        padded = attendees + [None] * (min_slots - len(attendees))
        ta = doc.add_table(rows=1, cols=7)
        ta.style = 'Table Grid'
        for i, txt in enumerate(['Name', 'BAC', 'Signature', '', 'Name', 'BAC', 'Signature']):
            _hdr_cell(ta.rows[0].cells[i], txt)
        for i in range(0, min_slots, 2):
            a1 = padded[i]
            a2 = padded[i+1] if i+1 < len(padded) else None
            row = ta.add_row()
            row.cells[0].text = a1.name if a1 else ''
            row.cells[1].text = a1.bac if a1 else ''
            row.cells[2].text = ''
            row.cells[3].text = ''
            row.cells[4].text = a2.name if a2 else ''
            row.cells[5].text = a2.bac if a2 else ''
            row.cells[6].text = ''
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8.5)
        for col_i, w in enumerate([5.5, 1.5, 3.5, 0.3, 5.5, 1.5, 3.5]):
            set_col_width(ta, col_i, w)

    # ── Photos ────────────────────────────────────────────────────────────
    photos = list(meeting.photos.all())
    if photos:
        sec_heading('Photos')
        cols = meeting.photo_columns
        col_width = Inches(6.5) / cols
        if cols == 1:
            for photo in photos:
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(photo.image.path, width=col_width)
                    p.paragraph_format.space_after = Pt(4)
                    if photo.caption:
                        cp = doc.add_paragraph(photo.caption)
                        cp.runs[0].font.size = Pt(8)
                        cp.paragraph_format.space_after = Pt(8)
                except Exception:
                    pass
        else:
            for i in range(0, len(photos), 2):
                t = doc.add_table(rows=1, cols=2)
                t.style = 'Table Grid'
                for j, photo in enumerate(photos[i:i+2]):
                    try:
                        cell = t.rows[0].cells[j]
                        cell.width = col_width
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(photo.image.path, width=col_width - Inches(0.1))
                        if photo.caption:
                            cp = cell.add_paragraph(photo.caption)
                            cp.runs[0].font.size = Pt(8)
                    except Exception:
                        pass
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Equipment Allocation ───────────────────────────────────────────────
    if meeting.include_equipment_allocation:
        eq_alloc = _toolbox_eq_alloc_data(meeting.job)
        if eq_alloc:
            sec_heading('Equipment Allocation')
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            hdr = t.rows[0].cells
            for cell, label in zip(hdr, ('Person', 'Equipment')):
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)
                cell._tc.get_or_add_tcPr().append(_cell_shading('EEF3F8'))
            for row_data in eq_alloc:
                eq_labels = ', '.join(
                    f"{e['name']} ({e['equipment_type']})" if e['equipment_type'] else e['name']
                    for e in row_data['equipment']
                )
                r = t.add_row()
                r.cells[0].text = row_data['person_name']
                r.cells[1].text = eq_labels
                for cell in r.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    # ── Footer note ───────────────────────────────────────────────────────
    form_ref = 'VF_GEN_001: Daily Toolbox Meeting Record' if meeting.meeting_type == 'daily' else 'VF_GEN_023: Journey Management Toolbox Meeting'
    fp = doc.add_paragraph(f'\nVelseis Form {form_ref}   |   Version 1.0   |   D. Tucker   |   22/06/2018   |   K. Dummett')
    fp.runs[0].font.size = Pt(6)
    fp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Output ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    label = 'Daily_Toolbox' if meeting.meeting_type == 'daily' else 'JMP_Toolbox'
    date_str = meeting.date.strftime('%d_%m_%Y') if meeting.date else 'undated'
    response = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{label}_{date_str}.docx"'
    return response


def job_location_delete(request, pk):
    loc = get_object_or_404(JobLocation, pk=pk)
    if request.method == 'POST':
        loc.delete()
    return JsonResponse({'ok': True})


def job_muster_point_delete(request, pk):
    mp = get_object_or_404(JobMusterPoint, pk=pk)
    if request.method == 'POST':
        mp.delete()
    return JsonResponse({'ok': True})


def toolbox_photo_upload(request, meeting_pk):
    from django.db.models import Max
    meeting = get_object_or_404(ToolboxMeeting, pk=meeting_pk)
    if request.method == 'POST':
        max_order = meeting.photos.aggregate(m=Max('order'))['m'] or 0
        for i, f in enumerate(request.FILES.getlist('images')):
            ToolboxPhoto.objects.create(meeting=meeting, image=f, order=max_order + i + 1)
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_photo_delete(request, pk):
    photo = get_object_or_404(ToolboxPhoto, pk=pk)
    if request.method == 'POST':
        photo.image.delete(save=False)
        photo.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_photo_update(request, pk):
    photo = get_object_or_404(ToolboxPhoto, pk=pk)
    if request.method == 'POST':
        photo.caption = request.POST.get('caption', '').strip()
        photo.border_style = request.POST.get('border_style', 'none')
        photo.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_photo_reorder(request, meeting_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        for item in data:
            ToolboxPhoto.objects.filter(pk=item['pk'], meeting_id=meeting_pk).update(order=item['order'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_photo_columns(request, meeting_pk):
    if request.method == 'POST':
        meeting = get_object_or_404(ToolboxMeeting, pk=meeting_pk)
        cols = int(request.POST.get('columns', 2))
        meeting.photo_columns = cols if cols in (1, 2) else 2
        meeting.save()
        return JsonResponse({'ok': True, 'columns': meeting.photo_columns})
    return JsonResponse({'ok': False}, status=400)


def toolbox_topic_template_add(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        body = request.POST.get('body', '').strip()
        if name:
            ToolboxTopicTemplate.objects.get_or_create(name=name, defaults={'body': body})
    return JsonResponse({'ok': True})


def toolbox_topic_template_delete(request, pk):
    t = get_object_or_404(ToolboxTopicTemplate, pk=pk)
    if request.method == 'POST':
        t.delete()
    return JsonResponse({'ok': True})


def toolbox_topic_template_update(request, pk):
    t = get_object_or_404(ToolboxTopicTemplate, pk=pk)
    if request.method == 'POST':
        body = request.POST.get('body', '').strip()
        t.body = body
        t.save()
    return JsonResponse({'ok': True})


def take5_form(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    import json as _json
    # All suggestions come from across all jobs
    task_suggestions = sorted(set(
        Take5Record.objects.values_list('task_description', flat=True)
    ))
    # task -> [hazards] map (all jobs)
    task_hazards_map = {}
    for r in Take5Record.objects.prefetch_related('hazards').all():
        key = r.task_description
        if key not in task_hazards_map:
            task_hazards_map[key] = []
        for h in r.hazards.all():
            if h.hazard and h.hazard not in task_hazards_map[key]:
                task_hazards_map[key].append(h.hazard)
    # hazard -> [controls] map (all jobs)
    hazard_suggestions = sorted(set(
        Take5Hazard.objects.values_list('hazard', flat=True)
    ))
    hazard_controls_map = {}
    for h in Take5Hazard.objects.prefetch_related('controls').all():
        key = h.hazard
        if key not in hazard_controls_map:
            hazard_controls_map[key] = []
        for c in h.controls.all():
            if c.control and c.control not in hazard_controls_map[key]:
                hazard_controls_map[key].append(c.control)

    error = None
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        pin = request.POST.get('pin', '').strip()
        task = request.POST.get('task', '').strip()
        acknowledged = request.POST.get('acknowledged') == 'on'
        hazards_json = request.POST.get('hazards_json', '[]')

        if not name:
            error = 'Please enter your name.'
        elif job.take5_pin and pin != job.take5_pin:
            error = 'Incorrect PIN. Please try again.'
        elif not task:
            error = 'Please describe the task.'
        elif not acknowledged:
            error = 'You must acknowledge the Take 5 before submitting.'
        else:
            try:
                hazards_data = json.loads(hazards_json)
            except (ValueError, TypeError):
                hazards_data = []
            record = Take5Record.objects.create(
                job=job, submitted_by=name, task_description=task, acknowledged=True,
            )
            for i, item in enumerate(hazards_data):
                hazard_text = item.get('hazard', '').strip()
                controls = [c.strip() for c in item.get('controls', []) if c.strip()]
                if hazard_text or controls:
                    h = Take5Hazard.objects.create(record=record, hazard=hazard_text, order=i)
                    for j, ctrl in enumerate(controls):
                        Take5Control.objects.create(hazard=h, control=ctrl, order=j)
            return render(request, 'reports/take5_success.html', {'job': job, 'record': record})

    return render(request, 'reports/take5_form.html', {
        'job': job,
        'error': error,
        'hazard_suggestions': _json.dumps(hazard_suggestions),
        'hazard_controls_map': _json.dumps(hazard_controls_map),
        'task_suggestions': _json.dumps(task_suggestions),
        'task_hazards_map': _json.dumps(task_hazards_map),
    })


def take5_detail(request, job_pk, pk):
    job = get_object_or_404(Job, pk=job_pk)
    record = get_object_or_404(Take5Record, pk=pk, job=job)
    return render(request, 'reports/take5_detail.html', {'job': job, 'record': record})


def take5_delete(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    record = get_object_or_404(Take5Record, pk=pk)
    record.delete()
    return JsonResponse({'ok': True})


def take5_hazard_delete(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    hazard = get_object_or_404(Take5Hazard, pk=pk)
    hazard.delete()
    return JsonResponse({'ok': True})


def take5_control_delete(request, pk):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    control = get_object_or_404(Take5Control, pk=pk)
    control.delete()
    return JsonResponse({'ok': True})


def take5_qr(request, job_pk):
    import qrcode
    import base64
    job = get_object_or_404(Job, pk=job_pk)
    url = request.build_absolute_uri(reverse('take5_form', args=[job_pk]))
    img = qrcode.make(url)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return render(request, 'reports/take5_qr.html', {
        'job': job,
        'url': url,
        'qr_b64': base64.b64encode(buf.getvalue()).decode(),
    })


def _take5_charts(records):
    """Build base64 chart images from a queryset of Take5Record."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker
    from collections import Counter
    import base64, io as _io

    records = list(records)

    def to_b64(fig):
        buf = _io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', dpi=130)
        buf.seek(0)
        plt.close(fig)
        return 'data:image/png;base64,' + base64.b64encode(buf.read()).decode()

    BLUE = '#1a3a5c'
    LIGHT = '#c5d5e8'
    colors = ['#1a3a5c', '#29b6f6', '#4caf50', '#ff9800', '#e53935',
              '#9c27b0', '#00bcd4', '#ff5722', '#607d8b', '#8bc34a']

    # ── Hazard frequency ────────────────────────────────────────────────────
    hazard_counter = Counter()
    for r in records:
        for h in r.hazards.all():
            if h.hazard:
                hazard_counter[h.hazard] += 1

    # ── Control frequency ───────────────────────────────────────────────────
    control_counter = Counter()
    for r in records:
        for h in r.hazards.all():
            for c in h.controls.all():
                if c.control:
                    control_counter[c.control] += 1

    # ── Submissions per person ───────────────────────────────────────────────
    person_counter = Counter(r.submitted_by for r in records)

    # ── Submissions per day ──────────────────────────────────────────────────
    day_counter = Counter(r.submitted_at.date() for r in records)

    charts = {}

    # Chart 1: Top hazards bar chart
    if hazard_counter:
        top = hazard_counter.most_common(10)
        labels = [t[0] if len(t[0]) <= 30 else t[0][:28] + '…' for t in top]
        vals = [t[1] for t in top]
        fig, ax = plt.subplots(figsize=(7, max(2.5, len(top) * 0.45)))
        bars = ax.barh(labels[::-1], vals[::-1], color=BLUE, height=0.6)
        ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
        ax.set_xlabel('Occurrences', fontsize=8)
        ax.set_title('Top Hazards', fontsize=10, fontweight='bold', color=BLUE)
        ax.tick_params(labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        for bar, val in zip(bars, vals[::-1]):
            ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                    str(val), va='center', fontsize=8)
        fig.tight_layout()
        charts['hazards'] = to_b64(fig)

    # Chart 2: Top controls bar chart
    if control_counter:
        top = control_counter.most_common(10)
        labels = [t[0] if len(t[0]) <= 30 else t[0][:28] + '…' for t in top]
        vals = [t[1] for t in top]
        fig, ax = plt.subplots(figsize=(7, max(2.5, len(top) * 0.45)))
        bars = ax.barh(labels[::-1], vals[::-1], color='#29b6f6', height=0.6)
        ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
        ax.set_xlabel('Occurrences', fontsize=8)
        ax.set_title('Top Controls', fontsize=10, fontweight='bold', color=BLUE)
        ax.tick_params(labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        for bar, val in zip(bars, vals[::-1]):
            ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                    str(val), va='center', fontsize=8)
        fig.tight_layout()
        charts['controls'] = to_b64(fig)

    # Chart 3: Submissions per person (horizontal bar)
    if person_counter:
        top = person_counter.most_common(15)
        labels = [t[0] for t in top]
        vals = [t[1] for t in top]
        fig, ax = plt.subplots(figsize=(7, max(2.5, len(top) * 0.45)))
        bar_colors = [colors[i % len(colors)] for i in range(len(top))]
        bars = ax.barh(labels[::-1], vals[::-1], color=bar_colors[::-1], height=0.6)
        ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
        ax.set_xlabel('Submissions', fontsize=8)
        ax.set_title('Submissions by Person', fontsize=10, fontweight='bold', color=BLUE)
        ax.tick_params(labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        for bar, val in zip(bars, vals[::-1]):
            ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                    str(val), va='center', fontsize=8)
        fig.tight_layout()
        charts['persons'] = to_b64(fig)

    # Chart 4: Submissions per day (only useful if >1 day span)
    if len(day_counter) > 1:
        days_sorted = sorted(day_counter.keys())
        labels = [d.strftime('%d %b') for d in days_sorted]
        vals = [day_counter[d] for d in days_sorted]
        fig, ax = plt.subplots(figsize=(7, 2.8))
        ax.bar(labels, vals, color=BLUE, width=0.6)
        ax.yaxis.set_major_locator(mticker.MaxNLocator(integer=True))
        ax.set_ylabel('Submissions', fontsize=8)
        ax.set_title('Submissions by Day', fontsize=10, fontweight='bold', color=BLUE)
        ax.tick_params(labelsize=8)
        plt.xticks(rotation=45, ha='right')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        fig.tight_layout()
        charts['days'] = to_b64(fig)

    # ── Summary tables data ──────────────────────────────────────────────────
    total = len(records)
    tables = {
        'hazards': [(h, c, round(c/total*100) if total else 0)
                    for h, c in hazard_counter.most_common()],
        'controls': [(c, n) for c, n in control_counter.most_common()],
        'persons': [(p, n) for p, n in person_counter.most_common()],
        'total': total,
    }

    return charts, tables


def _take5_pdf(request, job, records, title, filename):
    import tempfile, os as _os
    records = list(records)
    logo_b64 = _logo_b64()
    charts, tables = _take5_charts(records)
    html = render_to_string('reports/take5_report_pdf.html', {
        'job': job,
        'records': records,
        'title': title,
        'logo_b64': logo_b64,
        'charts': charts,
        'tables': tables,
    })
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            pg = browser.new_page()
            pg.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            pg.wait_for_load_state('networkidle')
            pdf_bytes = pg.pdf(format='A4', print_background=True,
                               display_header_footer=False)
            browser.close()
    finally:
        _os.unlink(tmp.name)
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    disposition = 'attachment' if request.GET.get('download') else 'inline'
    response['Content-Disposition'] = f'{disposition}; filename="{filename}"'
    return response


def take5_report_day(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    date_str = request.GET.get('date', '')
    from datetime import datetime as _dt
    try:
        day = _dt.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        from django.utils.timezone import now
        day = now().date()
    records = (job.take5_records
               .filter(submitted_at__date=day)
               .prefetch_related('hazards__controls'))
    title = f'Take 5 Report — {day.strftime("%d %b %Y")}'
    filename = f'Take5_{job.job_number}_{day.strftime("%d_%m_%Y")}.pdf'
    return _take5_pdf(request, job, records, title, filename)


def take5_report_job(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    records = job.take5_records.prefetch_related('hazards__controls').all()
    title = f'Take 5 Report — {job.job_number} (All Records)'
    filename = f'Take5_{job.job_number}_All.pdf'
    return _take5_pdf(request, job, records, title, filename)


def take5_send_invite(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    from django.core.mail import send_mail
    from django.conf import settings as _settings

    # All personnel with emails in the global list
    all_personnel = Personnel.objects.exclude(email='').order_by('name')

    # Names on this job (for pre-checking)
    job_names = set(
        job.personnel.filter(is_separator=False)
        .values_list('name', flat=True)
    )

    take5_url = request.build_absolute_uri(reverse('take5_form', args=[job_pk]))
    using_console = (getattr(_settings, 'EMAIL_BACKEND', '') ==
                     'django.core.mail.backends.console.EmailBackend')

    sent = None
    errors = []

    if request.method == 'POST':
        selected_pks = request.POST.getlist('personnel')
        extra_raw = request.POST.get('extra_emails', '').strip()

        addresses = []
        name_map = {}
        for p in Personnel.objects.filter(pk__in=selected_pks):
            addresses.append(p.email)
            name_map[p.email] = p.name

        for addr in [e.strip() for e in extra_raw.replace(',', '\n').splitlines() if e.strip()]:
            if addr not in addresses:
                addresses.append(addr)

        if not addresses:
            errors.append('No email addresses selected.')
        else:
            pin_line = f'\n\nJob PIN: {job.take5_pin}' if job.take5_pin else ''
            sent_list = []
            for addr in addresses:
                name = name_map.get(addr, 'Hi')
                try:
                    send_mail(
                        subject=f'Take 5 Pre-Task Check — Job {job.job_number}',
                        message=(
                            f'Hi {name},\n\n'
                            f'Please complete your Take 5 pre-task safety check for '
                            f'Job {job.job_number} — {job.project_name}.\n\n'
                            f'Open your Take 5 form here:\n{take5_url}'
                            f'{pin_line}\n\n'
                            f'This link can be used on your phone or any browser.\n\n'
                            f'Stay safe,\nField Manager'
                        ),
                        from_email=_settings.DEFAULT_FROM_EMAIL,
                        recipient_list=[addr],
                        fail_silently=False,
                    )
                    sent_list.append(addr)
                except Exception as e:
                    errors.append(f'{addr}: {e}')
            sent = sent_list

    return render(request, 'reports/take5_invite.html', {
        'job': job,
        'all_personnel': all_personnel,
        'job_names': job_names,
        'take5_url': take5_url,
        'using_console': using_console,
        'sent': sent,
        'errors': errors,
    })


# ---------------------------------------------------------------------------
# Toolbox V2 — digital signature version (clone of original + signature pads)
# ---------------------------------------------------------------------------

def toolbox_v2_list(request, job_pk):
    return redirect('toolbox_list', job_pk=job_pk)


def _save_toolbox_v2(request, meeting):
    """Save all fields from a V2 toolbox form POST, preserving existing signatures."""
    def d(field): return request.POST.get(field, '').strip() or None

    meeting.date = d('date')
    meeting.time = d('time')
    meeting.location = request.POST.get('location', '').strip()
    meeting.supervisor = request.POST.get('supervisor', '').strip()
    meeting.days_on_job = request.POST.get('days_on_job', '').strip()
    meeting.jmp_number = request.POST.get('jmp_number', '').strip()

    meeting.yesterday_activities = request.POST.get('yesterday_activities', '').strip()
    meeting.jmp_route_discussed = 'jmp_route_discussed' in request.POST
    meeting.jmp_hours_noted = 'jmp_hours_noted' in request.POST
    meeting.jmp_contact_numbers = 'jmp_contact_numbers' in request.POST
    meeting.jmp_signed_off = 'jmp_signed_off' in request.POST
    meeting.jmp_lead_tail = 'jmp_lead_tail' in request.POST

    meeting.todays_activities = request.POST.get('todays_activities', '').strip()
    meeting.terrain_discussion = request.POST.get('terrain_discussion', '').strip()
    meeting.road_condition = request.POST.get('road_condition', '').strip()
    meeting.muster_point = request.POST.get('muster_point', '').strip()
    meeting.forecast = request.POST.get('forecast', '').strip()
    meeting.uv_index = request.POST.get('uv_index', '').strip()
    meeting.chance_of_rain = request.POST.get('chance_of_rain', '').strip()
    meeting.min_temp = request.POST.get('min_temp', '').strip()
    meeting.max_temp = request.POST.get('max_temp', '').strip()
    meeting.grass_fire = request.POST.get('grass_fire', '').strip()
    meeting.forest_fire = request.POST.get('forest_fire', '').strip()
    meeting.wind = request.POST.get('wind', '').strip()
    meeting.other_topics = request.POST.get('other_topics', '').strip()
    meeting.include_equipment_allocation = 'include_equipment_allocation' in request.POST
    meeting.save()

    if meeting.location:
        JobLocation.objects.get_or_create(job=meeting.job, name=meeting.location)
    if meeting.muster_point:
        JobMusterPoint.objects.get_or_create(job=meeting.job, name=meeting.muster_point)

    # Attendees — update by PK to preserve signatures
    names = request.POST.getlist('a_name')
    roles = request.POST.getlist('a_role')
    vehicles = request.POST.getlist('a_vehicle')
    bacs = request.POST.getlist('a_bac')
    pks = request.POST.getlist('a_pk')
    is_seps = request.POST.getlist('a_is_sep')

    submitted_pks = set()
    for i, name in enumerate(names):
        try:
            pk = int(pks[i]) if i < len(pks) else 0
        except (ValueError, TypeError):
            pk = 0
        is_sep = (is_seps[i] == '1') if i < len(is_seps) else False
        role = '' if is_sep else (roles[i].strip() if i < len(roles) else '')
        vehicle = '' if is_sep else (vehicles[i].strip() if i < len(vehicles) else '')
        bac = '' if is_sep else (bacs[i].strip() if i < len(bacs) else '')

        if pk and ToolboxV2Attendee.objects.filter(pk=pk, meeting=meeting).exists():
            a = ToolboxV2Attendee.objects.get(pk=pk)
            a.name = name.strip()
            a.is_separator = is_sep
            a.job_role = role
            a.vehicle = vehicle
            a.bac = bac
            a.order = i
            a.save()
            submitted_pks.add(pk)
        else:
            a = ToolboxV2Attendee.objects.create(
                meeting=meeting,
                name=name.strip(),
                is_separator=is_sep,
                job_role=role,
                vehicle=vehicle,
                bac=bac,
                order=i,
            )
            submitted_pks.add(a.pk)

    # Remove attendees not in this submission
    meeting.attendees.exclude(pk__in=submitted_pks).delete()


def toolbox_v2_create(request, job_pk, meeting_type):
    job = get_object_or_404(Job, pk=job_pk)
    if meeting_type not in ('daily', 'jmp'):
        return redirect('toolbox_v2_list', job_pk=job_pk)
    if request.method == 'POST':
        meeting = ToolboxV2Meeting(job=job, meeting_type=meeting_type)
        meeting.save()
        _save_toolbox_v2(request, meeting)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': meeting.pk,
                                 'edit_url': reverse('toolbox_v2_edit', args=[meeting.pk])})
        return redirect('toolbox_v2_edit', pk=meeting.pk)
    jp = list(job.personnel.filter(is_separator=False).order_by('name'))
    supervisors = list(job.personnel.filter(roles__role__icontains='supervisor').distinct())
    previous = ToolboxV2Meeting.objects.filter(job=job, meeting_type=meeting_type).order_by('-date', '-pk').first()
    prev_data = None
    if previous:
        prev_data = json.dumps({
            'location':             previous.location or '',
            'supervisor':           previous.supervisor or '',
            'muster_point':         previous.muster_point or '',
            'terrain_discussion':   previous.terrain_discussion or '',
            'yesterday_activities': previous.todays_activities or '',
            'other_topics':         previous.other_topics or '',
            'days_on_job':          previous.days_on_job or '',
            'jmp_number':           previous.jmp_number or '',
            'road_condition':       previous.road_condition or '',
            'time':                 previous.time.strftime('%H:%M') if previous.time else '',
        })
    jmp_numbers = list(job.journey_plans.exclude(plan_number='').order_by('-created_at').values_list('plan_number', flat=True)[:20])
    personnel_roles_json = json.dumps({p.name: ', '.join(r.role for r in p.roles.all()) for p in jp})
    job_vehicle_names = list(job.vehicles.values_list('name', flat=True))
    return render(request, 'reports/toolbox_v2_form.html', {
        'job': job,
        'meeting': None,
        'meeting_type': meeting_type,
        'job_personnel': jp,
        'supervisors': supervisors,
        'saved_locations': job.saved_locations.all(),
        'saved_muster_points': job.saved_muster_points.all(),
        'topic_templates': ToolboxTopicTemplate.objects.all(),
        'prev_data': prev_data,
        'jmp_numbers': jmp_numbers,
        'personnel_roles_json': personnel_roles_json,
        'job_vehicle_names': job_vehicle_names,
        'alloc_data_json': json.dumps(_toolbox_alloc_data(job)),
        'eq_alloc_data_json': json.dumps(_toolbox_eq_alloc_data(job)),
    })


def toolbox_v2_edit(request, pk):
    meeting = get_object_or_404(ToolboxV2Meeting, pk=pk)
    if request.method == 'POST':
        _save_toolbox_v2(request, meeting)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': meeting.pk})
        return redirect('toolbox_v2_edit', pk=meeting.pk)
    jp = list(meeting.job.personnel.filter(is_separator=False).order_by('sort_order', 'pk'))
    supervisors = list(meeting.job.personnel.filter(roles__role__icontains='supervisor').distinct())
    personnel_roles_json = json.dumps({p.name: ', '.join(r.role for r in p.roles.all()) for p in jp})
    job_vehicle_names = list(meeting.job.vehicles.values_list('name', flat=True))
    attendees = list(meeting.attendees.order_by('order', 'pk'))
    return render(request, 'reports/toolbox_v2_form.html', {
        'job': meeting.job,
        'meeting': meeting,
        'meeting_type': meeting.meeting_type,
        'job_personnel': jp,
        'supervisors': supervisors,
        'saved_locations': meeting.job.saved_locations.all(),
        'saved_muster_points': meeting.job.saved_muster_points.all(),
        'topic_templates': ToolboxTopicTemplate.objects.all(),
        'personnel_roles_json': personnel_roles_json,
        'job_vehicle_names': job_vehicle_names,
        'alloc_data_json': json.dumps(_toolbox_alloc_data(meeting.job)),
        'eq_alloc_data_json': json.dumps(_toolbox_eq_alloc_data(meeting.job)),
        'jmp_numbers': list(meeting.job.journey_plans.exclude(plan_number='').order_by('-created_at').values_list('plan_number', flat=True)[:20]),
        'attendees': attendees,
        'photos': meeting.photos.all(),
    })


def toolbox_v2_delete(request, pk):
    meeting = get_object_or_404(ToolboxV2Meeting, pk=pk)
    job_pk = meeting.job_id
    if request.method == 'POST':
        meeting.delete()
    return redirect(reverse('toolbox_list', args=[job_pk]) + '?tab=v2')


def toolbox_v2_copy(request, pk):
    src = get_object_or_404(ToolboxV2Meeting, pk=pk)
    new_meeting = ToolboxV2Meeting(
        job=src.job,
        meeting_type=src.meeting_type,
        location=src.location,
        supervisor=src.supervisor,
        days_on_job=src.days_on_job,
        jmp_number=src.jmp_number,
        yesterday_activities=src.yesterday_activities,
        jmp_route_discussed=src.jmp_route_discussed,
        jmp_hours_noted=src.jmp_hours_noted,
        jmp_contact_numbers=src.jmp_contact_numbers,
        jmp_signed_off=src.jmp_signed_off,
        jmp_lead_tail=src.jmp_lead_tail,
        todays_activities=src.todays_activities,
        terrain_discussion=src.terrain_discussion,
        road_condition=src.road_condition,
        muster_point=src.muster_point,
        forecast=src.forecast,
        uv_index=src.uv_index,
        chance_of_rain=src.chance_of_rain,
        min_temp=src.min_temp,
        max_temp=src.max_temp,
        grass_fire=src.grass_fire,
        forest_fire=src.forest_fire,
        wind=src.wind,
        other_topics=src.other_topics,
        include_equipment_allocation=src.include_equipment_allocation,
    )
    new_meeting.save()
    for a in src.attendees.all():
        ToolboxV2Attendee.objects.create(
            meeting=new_meeting,
            is_separator=a.is_separator,
            name=a.name,
            job_role=a.job_role,
            vehicle=a.vehicle,
            order=a.order,
        )
    return redirect('toolbox_v2_edit', pk=new_meeting.pk)


def toolbox_v2_attendee_sign(request, pk):
    """AJAX: save a drawn signature for an attendee."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    attendee = get_object_or_404(ToolboxV2Attendee, pk=pk)
    data = json.loads(request.body)
    sig = data.get('signature', '').strip()
    if not sig:
        return JsonResponse({'error': 'No signature data'}, status=400)
    from datetime import datetime
    attendee.signature = sig
    try:
        attendee.signed_at = datetime.fromisoformat(data.get('signed_at', ''))
    except (ValueError, TypeError):
        from django.utils import timezone
        attendee.signed_at = timezone.now()
    attendee.save()
    return JsonResponse({'ok': True, 'signed_at': attendee.signed_at.strftime('%d %b %Y %H:%M')})


def toolbox_v2_attendee_clear(request, pk):
    """AJAX: clear a signature."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    attendee = get_object_or_404(ToolboxV2Attendee, pk=pk)
    attendee.signature = ''
    attendee.signed_at = None
    attendee.save()
    return JsonResponse({'ok': True})


def toolbox_v2_photo_upload(request, meeting_pk):
    from django.db.models import Max
    meeting = get_object_or_404(ToolboxV2Meeting, pk=meeting_pk)
    if request.method == 'POST':
        max_order = meeting.photos.aggregate(m=Max('order'))['m'] or 0
        for i, f in enumerate(request.FILES.getlist('images')):
            ToolboxV2Photo.objects.create(meeting=meeting, image=f, order=max_order + i + 1)
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_v2_photo_delete(request, pk):
    photo = get_object_or_404(ToolboxV2Photo, pk=pk)
    if request.method == 'POST':
        photo.image.delete(save=False)
        photo.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_v2_photo_update(request, pk):
    photo = get_object_or_404(ToolboxV2Photo, pk=pk)
    if request.method == 'POST':
        photo.caption = request.POST.get('caption', '').strip()
        photo.border_style = request.POST.get('border_style', 'none')
        photo.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_v2_photo_reorder(request, meeting_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        for item in data:
            ToolboxV2Photo.objects.filter(pk=item['pk'], meeting_id=meeting_pk).update(order=item['order'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def toolbox_v2_photo_columns(request, meeting_pk):
    if request.method == 'POST':
        meeting = get_object_or_404(ToolboxV2Meeting, pk=meeting_pk)
        cols = int(request.POST.get('columns', 2))
        meeting.photo_columns = cols if cols in (1, 2) else 2
        meeting.save()
        return JsonResponse({'ok': True, 'columns': meeting.photo_columns})
    return JsonResponse({'ok': False}, status=400)


def toolbox_v2_pdf(request, pk):
    meeting = get_object_or_404(ToolboxV2Meeting, pk=pk)
    all_attendees = list(meeting.attendees.order_by('order', 'pk'))
    has_separators = any(a.is_separator for a in all_attendees)
    if has_separators:
        attendees = all_attendees
        has_extended = True
        attendee_pairs = []
    else:
        attendees = sorted(all_attendees, key=lambda a: a.name.lower())
        has_extended = any(a.job_role or a.vehicle for a in attendees)
        min_slots = max(len(attendees), 12)
        if min_slots % 2:
            min_slots += 1
        padded = attendees[:] + [None] * (min_slots - len(attendees))
        attendee_pairs = [(padded[i], padded[i + 1]) for i in range(0, min_slots, 2)]
    jmp_items = [
        ('Route discussed with all crew', meeting.jmp_route_discussed),
        ('Driving hours / fatigue limits noted', meeting.jmp_hours_noted),
        ('Emergency contact numbers confirmed', meeting.jmp_contact_numbers),
        ('JMP signed off by supervisor', meeting.jmp_signed_off),
        ('Lead / tail vehicle assigned', meeting.jmp_lead_tail),
    ]
    from pathlib import Path
    photos_with_uri = [
        {'photo': p, 'uri': Path(p.image.path).as_uri()}
        for p in meeting.photos.all()
    ]
    html = render_to_string('reports/toolbox_v2_pdf.html', {
        'meeting': meeting,
        'attendees': attendees,
        'attendee_pairs': attendee_pairs,
        'has_extended': has_extended,
        'has_separators': has_separators,
        'jmp_items': jmp_items,
        'logo_b64': _logo_b64(),
        'photos_with_uri': photos_with_uri,
        'eq_alloc_data': _toolbox_eq_alloc_data(meeting.job) if meeting.include_equipment_allocation else [],
    }, request=request)
    import tempfile, os as _os
    from playwright.sync_api import sync_playwright
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            page.wait_for_load_state('networkidle')
            pdf_bytes = page.pdf(format='A4', print_background=True,
                                 margin={'top': '15mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'})
            browser.close()
    finally:
        _os.unlink(tmp.name)
    safe_date = str(meeting.date or 'no-date').replace('-', '')
    filename = f"Toolbox_V2_{meeting.job.job_number}_{safe_date}.pdf"
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    return response


# ─── Journey Management V2 (digital signatures) ───────────────────────────────

def journey_v2_list(request, job_pk):
    return redirect('journey_list', job_pk=job_pk)


def _save_journey_v2(request, plan):
    """Save all form fields onto plan; preserve personnel signatures by pk."""
    def d(field):
        v = request.POST.get(field, '').strip()
        return v or None
    def b(field):
        v = request.POST.get(field)
        return True if v == 'Y' else (False if v == 'N' else None)

    plan.plan_number = request.POST.get('plan_number', '').strip()
    plan.departing_from = request.POST.get('departing_from', '').strip()
    plan.depart_date = d('depart_date')
    plan.depart_time = d('depart_time')
    plan.depart_contact = request.POST.get('depart_contact', '').strip()
    plan.depart_phone = request.POST.get('depart_phone', '').strip()
    plan.overnight_location = request.POST.get('overnight_location', '').strip()
    plan.overnight_arrival_date = d('overnight_arrival_date')
    plan.overnight_arrival_time = d('overnight_arrival_time')
    plan.overnight_departure_date = d('overnight_departure_date')
    plan.overnight_departure_time = d('overnight_departure_time')
    plan.arriving_at = request.POST.get('arriving_at', '').strip()
    plan.arrive_date = d('arrive_date')
    plan.arrive_time = d('arrive_time')
    plan.arrive_contact = request.POST.get('arrive_contact', '').strip()
    plan.arrive_phone = request.POST.get('arrive_phone', '').strip()
    plan.route = request.POST.get('route', '').strip()
    plan.break_journey_at = request.POST.get('break_journey_at', '').strip()
    plan.radio_channel = request.POST.get('radio_channel', '').strip()
    plan.other_instructions = request.POST.get('other_instructions', '').strip()
    plan.route_waypoints = request.POST.get('route_waypoints', '').strip()
    plan.rest_stops_json = request.POST.get('rest_stops_json', '').strip()
    plan.coordinator_name = request.POST.get('coordinator_name', '').strip()
    plan.coordinator_phone = request.POST.get('coordinator_phone', '').strip()
    plan.plan_communicated = b('plan_communicated')
    plan.before_date = d('before_date')
    plan.journey_completed = b('journey_completed')
    plan.after_date = d('after_date')
    plan.action_items = request.POST.get('action_items', '').strip()
    plan.include_map_in_pdf = 'include_map_in_pdf' in request.POST
    plan.map_tile_layer = request.POST.get('map_tile_layer', 'street')
    plan.save()

    # Personnel: preserve signatures by matching submitted p_pk values
    submitted_pks = request.POST.getlist('p_pk')
    regos = request.POST.getlist('p_rego')
    names = request.POST.getlist('p_name')
    drivers = request.POST.getlist('p_driver')
    phones = request.POST.getlist('p_phone')

    existing_by_pk = {p.pk: p for p in plan.personnel.all()}
    submitted_ids = set()

    for i, name in enumerate(names):
        if not name.strip():
            continue
        try:
            p_pk = int(submitted_pks[i]) if i < len(submitted_pks) else 0
        except (ValueError, TypeError):
            p_pk = 0

        rego = regos[i].strip() if i < len(regos) else ''
        is_driver = (drivers[i] == 'Y') if i < len(drivers) else False
        phone = phones[i].strip() if i < len(phones) else ''

        if p_pk and p_pk in existing_by_pk:
            p = existing_by_pk[p_pk]
            p.rego = rego
            p.name = name.strip()
            p.is_driver = is_driver
            p.phone = phone
            p.order = i
            p.save()
            submitted_ids.add(p_pk)
        else:
            new_p = JourneyV2Personnel.objects.create(
                plan=plan, rego=rego, name=name.strip(),
                is_driver=is_driver, phone=phone, order=i,
            )
            submitted_ids.add(new_p.pk)

    # Delete removed rows
    for pk, p in existing_by_pk.items():
        if pk not in submitted_ids:
            p.delete()


def journey_v2_create(request, job_pk):
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        plan = JourneyV2Plan(job=job)
        _save_journey_v2(request, plan)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': plan.pk,
                                 'edit_url': reverse('journey_v2_edit', args=[plan.pk])})
        return redirect('journey_v2_list', job_pk=job_pk)

    personnel_phones = {p.name: p.phone for p in Personnel.objects.exclude(phone='')}
    jp_qs = job.personnel.order_by('name')
    mob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=job, status__iexact='mobilise').values_list('job_personnel__name', flat=True))
    demob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=job, status__iexact='demobilise').values_list('job_personnel__name', flat=True))
    job_personnel_data = [{'name': p.name, 'phone': personnel_phones.get(p.name, ''),
                           'mob': p.name in mob_names, 'demob': p.name in demob_names} for p in jp_qs]
    return render(request, 'reports/journey_v2_form.html', {
        'job': job, 'plan': None,
        'job_personnel': jp_qs,
        'personnel_phones': personnel_phones,
        'job_personnel_data': job_personnel_data,
        'alloc_by_tab': {}, 'alloc_tabs': [],
    })


def journey_v2_edit(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    if request.method == 'POST':
        _save_journey_v2(request, plan)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'ok': True, 'pk': plan.pk})
        return redirect('journey_v2_list', job_pk=plan.job_id)

    personnel_phones = {p.name: p.phone for p in Personnel.objects.exclude(phone='')}
    jp_qs = plan.job.personnel.all()
    mob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=plan.job, status__iexact='mobilise').values_list('job_personnel__name', flat=True))
    demob_names = set(PersonnelScheduleEntry.objects.filter(
        job_personnel__job=plan.job, status__iexact='demobilise').values_list('job_personnel__name', flat=True))
    job_personnel_data = [{'name': p.name, 'phone': personnel_phones.get(p.name, ''),
                           'mob': p.name in mob_names, 'demob': p.name in demob_names} for p in jp_qs]

    fleet_by_name = {v.name: v for v in Vehicle.objects.all()}
    job_vehicles_nv = list(plan.job.vehicles.exclude(vehicle_type='Vibrator'))
    eq_links_for_suffix = JobEquipmentVehicleLink.objects.filter(
        job_vehicle__job=plan.job,
        job_equipment__equipment_type__in=('Trailer', 'Starlink')
    ).select_related('job_equipment')
    trailers_by_veh_tab = {}
    starlinks_by_veh_tab = {}
    for lnk in eq_links_for_suffix:
        key = (lnk.job_vehicle_id, lnk.tab)
        if lnk.job_equipment.equipment_type == 'Trailer':
            trailers_by_veh_tab.setdefault(key, []).append(lnk.job_equipment.name)
        else:
            starlinks_by_veh_tab.setdefault(key, []).append(lnk.job_equipment.name)

    alloc_by_tab = {}
    for tab_key in ('mob', 'job', 'demob', 'other'):
        vehicles = []
        for jv in job_vehicles_nv:
            fleet = fleet_by_name.get(jv.name)
            rego_label = ' '.join(filter(None, [jv.name, fleet.rego if fleet else '']))
            suffix_parts = (
                [f'({t})' for t in trailers_by_veh_tab.get((jv.pk, tab_key), [])] +
                [f'({s})' for s in starlinks_by_veh_tab.get((jv.pk, tab_key), [])]
            )
            suffix = ' '.join(suffix_parts)
            people = [{'name': a.person_name + (' ' + suffix if suffix else ''),
                       'phone': personnel_phones.get(a.person_name, '')}
                      for a in jv.allocations.filter(tab=tab_key)]
            if people:
                vehicles.append({'rego_label': rego_label, 'people': people})
        alloc_by_tab[tab_key] = vehicles

    return render(request, 'reports/journey_v2_form.html', {
        'job': plan.job,
        'plan': plan,
        'job_personnel': jp_qs,
        'personnel_phones': personnel_phones,
        'job_personnel_data': job_personnel_data,
        'alloc_by_tab': alloc_by_tab,
        'alloc_tabs': [k for k, v in alloc_by_tab.items() if v],
    })


def journey_v2_delete(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    job_pk = plan.job_id
    if request.method == 'POST':
        plan.delete()
    return redirect(reverse('journey_list', args=[job_pk]) + '?tab=v2')


def journey_v2_copy(request, pk):
    src = get_object_or_404(JourneyV2Plan, pk=pk)
    new_plan = JourneyV2Plan(
        job=src.job,
        departing_from=src.departing_from,
        depart_time=src.depart_time,
        depart_contact=src.depart_contact,
        depart_phone=src.depart_phone,
        overnight_location=src.overnight_location,
        overnight_arrival_time=src.overnight_arrival_time,
        overnight_departure_time=src.overnight_departure_time,
        arriving_at=src.arriving_at,
        arrive_time=src.arrive_time,
        arrive_contact=src.arrive_contact,
        arrive_phone=src.arrive_phone,
        route=src.route,
        break_journey_at=src.break_journey_at,
        radio_channel=src.radio_channel,
        other_instructions=src.other_instructions,
        route_waypoints=src.route_waypoints,
        rest_stops_json=src.rest_stops_json,
        coordinator_name=src.coordinator_name,
        coordinator_phone=src.coordinator_phone,
        include_map_in_pdf=src.include_map_in_pdf,
        map_tile_layer=src.map_tile_layer,
    )
    new_plan.save()
    for p in src.personnel.all():
        JourneyV2Personnel.objects.create(
            plan=new_plan,
            rego=p.rego,
            name=p.name,
            is_driver=p.is_driver,
            phone=p.phone,
            order=p.order,
        )
    return redirect('journey_v2_edit', pk=new_plan.pk)


def journey_v2_personnel_sign(request, pk):
    """AJAX: save signature PNG (base64) for a personnel row."""
    from datetime import datetime
    p = get_object_or_404(JourneyV2Personnel, pk=pk)
    if request.method == 'POST':
        data = request.POST.get('signature', '')
        if data.startswith('data:image/png;base64,'):
            p.signature = data
            try:
                p.signed_at = datetime.fromisoformat(request.POST.get('signed_at', ''))
            except (ValueError, TypeError):
                from django.utils import timezone
                p.signed_at = timezone.now()
            p.save()
            return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def journey_v2_personnel_clear(request, pk):
    """AJAX: clear signature for a personnel row."""
    p = get_object_or_404(JourneyV2Personnel, pk=pk)
    if request.method == 'POST':
        p.signature = ''
        p.signed_at = None
        p.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def journey_v2_before_sign(request, pk):
    """AJAX: save Before Journey coordinator signature."""
    from datetime import datetime
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    if request.method == 'POST':
        data = request.POST.get('signature', '')
        if data.startswith('data:image/png;base64,'):
            plan.before_signature = data
            try:
                plan.before_signed_at = datetime.fromisoformat(request.POST.get('signed_at', ''))
            except (ValueError, TypeError):
                from django.utils import timezone
                plan.before_signed_at = timezone.now()
            plan.save()
            return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def journey_v2_before_clear(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    if request.method == 'POST':
        plan.before_signature = ''
        plan.before_signed_at = None
        plan.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def journey_v2_after_sign(request, pk):
    """AJAX: save After Journey coordinator signature."""
    from datetime import datetime
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    if request.method == 'POST':
        data = request.POST.get('signature', '')
        if data.startswith('data:image/png;base64,'):
            plan.after_signature = data
            try:
                plan.after_signed_at = datetime.fromisoformat(request.POST.get('signed_at', ''))
            except (ValueError, TypeError):
                from django.utils import timezone
                plan.after_signed_at = timezone.now()
            plan.save()
            return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def journey_v2_after_clear(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    if request.method == 'POST':
        plan.after_signature = ''
        plan.after_signed_at = None
        plan.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def _html_to_jpeg(html, viewport_width=900):
    """Render HTML to a full-page JPEG via Playwright and return the bytes."""
    import tempfile, os as _os
    from playwright.sync_api import sync_playwright
    tmp = tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.close()
        with sync_playwright() as p:
            browser = p.chromium.launch()
            pg = browser.new_page(viewport={'width': viewport_width, 'height': 1080})
            pg.goto(f'file:///{tmp.name.replace(chr(92), "/")}')
            pg.wait_for_load_state('networkidle')
            img_bytes = pg.screenshot(full_page=True, type='jpeg', quality=88)
            browser.close()
    finally:
        _os.unlink(tmp.name)
    return img_bytes


def toolbox_image(request, pk):
    meeting = get_object_or_404(ToolboxMeeting, pk=pk)
    logo_b64 = _logo_b64()
    attendees = sorted(meeting.attendees.all(), key=lambda a: a.name.lower())
    has_extended = any(a.job_role or a.vehicle for a in attendees)
    padded = attendees[:]
    min_slots = max(len(padded), 12)
    if min_slots % 2: min_slots += 1
    padded += [None] * (min_slots - len(padded))
    attendee_pairs = [(padded[i], padded[i+1]) for i in range(0, min_slots, 2)]
    jmp_items = [
        ('Route discussed with all crew', meeting.jmp_route_discussed),
        ('Driving hours / fatigue limits noted', meeting.jmp_hours_noted),
        ('Emergency contact numbers confirmed', meeting.jmp_contact_numbers),
        ('JMP signed off by supervisor', meeting.jmp_signed_off),
        ('Lead / tail vehicle assigned', meeting.jmp_lead_tail),
    ]
    from pathlib import Path
    photos_with_uri = [
        {'photo': p, 'uri': Path(p.image.path).as_uri()}
        for p in meeting.photos.all()
    ]
    html = render_to_string('reports/toolbox_pdf.html', {
        'meeting': meeting, 'logo_b64': logo_b64,
        'attendees': attendees, 'attendee_pairs': attendee_pairs,
        'has_extended': has_extended, 'jmp_items': jmp_items,
        'photos_with_uri': photos_with_uri,
        'eq_alloc_data': _toolbox_eq_alloc_data(meeting.job) if meeting.include_equipment_allocation else [],
    })
    img_bytes = _html_to_jpeg(html)
    label = 'Daily_Toolbox' if meeting.meeting_type == 'daily' else 'JMP_Toolbox'
    date_str = meeting.date.strftime('%d_%m_%Y') if meeting.date else 'undated'
    response = HttpResponse(img_bytes, content_type='image/jpeg')
    response['Content-Disposition'] = f'attachment; filename="{label}_{date_str}.jpg"'
    return response


def toolbox_v2_image(request, pk):
    meeting = get_object_or_404(ToolboxV2Meeting, pk=pk)
    all_attendees = list(meeting.attendees.order_by('order', 'pk'))
    has_separators = any(a.is_separator for a in all_attendees)
    if has_separators:
        attendees = all_attendees
        has_extended = True
        attendee_pairs = []
    else:
        attendees = sorted(all_attendees, key=lambda a: a.name.lower())
        has_extended = any(a.job_role or a.vehicle for a in attendees)
        min_slots = max(len(attendees), 12)
        if min_slots % 2: min_slots += 1
        padded = attendees[:] + [None] * (min_slots - len(attendees))
        attendee_pairs = [(padded[i], padded[i+1]) for i in range(0, min_slots, 2)]
    jmp_items = [
        ('Route discussed with all crew', meeting.jmp_route_discussed),
        ('Driving hours / fatigue limits noted', meeting.jmp_hours_noted),
        ('Emergency contact numbers confirmed', meeting.jmp_contact_numbers),
        ('JMP signed off by supervisor', meeting.jmp_signed_off),
        ('Lead / tail vehicle assigned', meeting.jmp_lead_tail),
    ]
    from pathlib import Path
    photos_with_uri = [
        {'photo': p, 'uri': Path(p.image.path).as_uri()}
        for p in meeting.photos.all()
    ]
    html = render_to_string('reports/toolbox_v2_pdf.html', {
        'meeting': meeting, 'attendees': attendees,
        'attendee_pairs': attendee_pairs, 'has_extended': has_extended,
        'has_separators': has_separators,
        'jmp_items': jmp_items, 'logo_b64': _logo_b64(),
        'photos_with_uri': photos_with_uri,
        'eq_alloc_data': _toolbox_eq_alloc_data(meeting.job) if meeting.include_equipment_allocation else [],
    }, request=request)
    img_bytes = _html_to_jpeg(html)
    safe_date = str(meeting.date or 'no-date').replace('-', '')
    response = HttpResponse(img_bytes, content_type='image/jpeg')
    response['Content-Disposition'] = f'attachment; filename="Toolbox_V2_{meeting.job.job_number}_{safe_date}.jpg"'
    return response


def journey_image(request, pk):
    plan = get_object_or_404(JourneyManagementPlan, pk=pk)
    logo_b64 = _logo_b64()
    route_waypoints = _build_waypoints(plan.route_waypoints)
    route_text_parts = [p.strip() for p in plan.route.split('→') if p.strip()] if plan.route else []
    route_total_km = sum(wp['km'] for wp in route_waypoints if wp.get('km') is not None)
    map_b64 = None
    if plan.include_map_in_pdf and plan.route_waypoints:
        map_b64 = _map_screenshot_b64(plan.route_waypoints, plan.map_tile_layer)
    html = render_to_string('reports/journey_pdf.html', {
        'plan': plan, 'logo_b64': logo_b64,
        'route_waypoints': route_waypoints, 'route_text_parts': route_text_parts,
        'route_total_km': route_total_km, 'include_map': plan.include_map_in_pdf,
        'map_b64': map_b64,
    })
    img_bytes = _html_to_jpeg(html)
    date_str = plan.depart_date.strftime('%d_%m_%Y') if plan.depart_date else 'undated'
    response = HttpResponse(img_bytes, content_type='image/jpeg')
    response['Content-Disposition'] = f'attachment; filename="Journey_Plan_{date_str}.jpg"'
    return response


def journey_v2_image(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    logo_b64 = _logo_b64()
    route_waypoints = _build_waypoints(plan.route_waypoints)
    route_text_parts = [p.strip() for p in plan.route.split('→') if p.strip()] if plan.route else []
    map_b64 = None
    if plan.include_map_in_pdf and plan.route_waypoints:
        map_b64 = _map_screenshot_b64(plan.route_waypoints, plan.map_tile_layer)
    route_total_km = sum(wp['km'] for wp in route_waypoints if wp.get('km') is not None)
    html = render_to_string('reports/journey_v2_pdf.html', {
        'plan': plan, 'logo_b64': logo_b64,
        'route_waypoints': route_waypoints, 'route_text_parts': route_text_parts,
        'route_total_km': route_total_km, 'include_map': plan.include_map_in_pdf,
        'map_b64': map_b64,
    })
    img_bytes = _html_to_jpeg(html)
    date_str = plan.depart_date.strftime('%d_%m_%Y') if plan.depart_date else 'undated'
    response = HttpResponse(img_bytes, content_type='image/jpeg')
    response['Content-Disposition'] = f'attachment; filename="Journey_V2_Plan_{date_str}.jpg"'
    return response


def journey_v2_pdf(request, pk):
    plan = get_object_or_404(JourneyV2Plan, pk=pk)
    logo_b64 = _logo_b64()
    route_waypoints = _build_waypoints(plan.route_waypoints)
    route_text_parts = [p.strip() for p in plan.route.split('→') if p.strip()] if plan.route else []
    # Map screenshot runs in its own Playwright session (before template rendering)
    map_b64 = None
    if plan.include_map_in_pdf and plan.route_waypoints:
        map_b64 = _map_screenshot_b64(plan.route_waypoints, plan.map_tile_layer)

    route_total_km = sum(wp['km'] for wp in route_waypoints if wp.get('km') is not None)

    # Template rendering happens outside Playwright (avoids SynchronousOnlyOperation)
    html = render_to_string('reports/journey_v2_pdf.html', {
        'plan': plan, 'logo_b64': logo_b64,
        'route_waypoints': route_waypoints,
        'route_text_parts': route_text_parts,
        'route_total_km': route_total_km,
        'include_map': plan.include_map_in_pdf,
        'map_b64': map_b64,
    })

    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        pdf_bytes = pg.pdf(format='A4', print_background=True,
                           margin={'top': '1.5cm', 'bottom': '1.5cm', 'left': '1.5cm', 'right': '1.5cm'})
        browser.close()
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    date_str = plan.depart_date.strftime('%d_%m_%Y') if plan.depart_date else 'undated'
    disposition = 'attachment' if request.GET.get('download') else 'inline'
    response['Content-Disposition'] = f'{disposition}; filename="Journey_V2_Plan_{date_str}.pdf"'
    return response


# ---------------------------------------------------------------------------
# Node Stock Take
# ---------------------------------------------------------------------------

def node_stock_take(request, job_pk):
    from .models import NodeStockTakeSession, NodeRecord, DeadNode
    job = get_object_or_404(Job, pk=job_pk)

    pre_sessions = job.node_stock_takes.filter(session_type='pre_job').prefetch_related('mega_bins__crates__nodes')
    post_sessions = job.node_stock_takes.filter(session_type='post_job').prefetch_related('mega_bins__crates__nodes')
    dead_nodes = job.dead_nodes.all()

    pre_serials = set(NodeRecord.objects.filter(
        crate__mega_bin__session__job=job, crate__mega_bin__session__session_type='pre_job'
    ).values_list('serial_number', flat=True))
    post_serials = set(NodeRecord.objects.filter(
        crate__mega_bin__session__job=job, crate__mega_bin__session__session_type='post_job'
    ).values_list('serial_number', flat=True))
    dead_serials = set(dead_nodes.values_list('serial_number', flat=True))

    missing = sorted(pre_serials - post_serials - dead_serials)
    extra = sorted(post_serials - pre_serials)

    recon = {
        'pre_count': len(pre_serials),
        'post_count': len(post_serials),
        'dead_count': len(dead_serials),
        'missing': missing,
        'extra': extra,
    }

    return render(request, 'reports/node_stock_take.html', {
        'job': job,
        'pre_sessions': pre_sessions,
        'post_sessions': post_sessions,
        'dead_nodes': dead_nodes,
        'recon': recon,
        'today': date.today(),
        'dead_reason_choices': [('battery','Battery Failure'),('physical','Physical Damage'),('lost','Lost'),('firmware','Firmware / Software Issue'),('other','Other')],
    })


def node_stock_take_session_create(request, job_pk):
    from .models import NodeStockTakeSession
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        session = NodeStockTakeSession.objects.create(
            job=job,
            session_type=request.POST.get('session_type', 'pre_job'),
            label=request.POST.get('label', '').strip(),
            date=request.POST.get('date') or date.today(),
            notes=request.POST.get('notes', '').strip(),
        )
        return redirect('node_stock_take_session_detail', pk=session.pk)
    return redirect('node_stock_take', job_pk=job_pk)


def node_stock_take_session_detail(request, pk):
    from .models import NodeStockTakeSession
    session = get_object_or_404(NodeStockTakeSession, pk=pk)

    if request.method == 'POST' and request.POST.get('action') == 'save_settings':
        try:
            session.nodes_per_crate = max(1, int(request.POST.get('nodes_per_crate', 20)))
            session.crates_per_mega_bin = max(1, int(request.POST.get('crates_per_mega_bin', 12)))
            session.crate_columns = max(1, int(request.POST.get('crate_columns', 6)))
            raw = request.POST.get('watch_serials', '')
            session.watch_serials = '\n'.join(
                s.strip() for s in raw.replace(',', '\n').splitlines() if s.strip()
            )
            color = request.POST.get('watch_color', '#ff8800').strip()
            if color.startswith('#') and len(color) in (4, 7):
                session.watch_color = color
            session.save(update_fields=['nodes_per_crate', 'crates_per_mega_bin', 'crate_columns', 'watch_serials', 'watch_color'])
        except (ValueError, TypeError):
            pass
        return JsonResponse({'ok': True})

    mega_bins = session.mega_bins.prefetch_related('crates__nodes').all()
    session_data = {
        'id': session.pk,
        'nodes_per_crate': session.nodes_per_crate,
        'crates_per_mega_bin': session.crates_per_mega_bin,
        'crate_columns': session.crate_columns,
        'watch_serials': [s for s in session.watch_serials.splitlines() if s.strip()],
        'watch_color': session.watch_color or '#ff8800',
        'mega_bins': [
            {
                'id': mb.pk,
                'name': mb.name,
                'crates': [
                    {
                        'id': c.pk,
                        'name': c.name,
                        'nodes': {str(n.slot): {'id': n.pk, 'serial': n.serial_number} for n in c.nodes.all()},
                    }
                    for c in mb.crates.all()
                ],
            }
            for mb in mega_bins
        ],
    }

    return render(request, 'reports/node_stock_take_session.html', {
        'session': session,
        'job': session.job,
        'session_json': json.dumps(session_data),
    })


def node_stock_take_session_delete(request, pk):
    from .models import NodeStockTakeSession
    session = get_object_or_404(NodeStockTakeSession, pk=pk)
    job_pk = session.job_id
    if request.method == 'POST':
        session.delete()
    return redirect('node_stock_take', job_pk=job_pk)


def node_stock_take_session_csv(request, pk):
    import csv as _csv
    from .models import NodeStockTakeSession
    from django.http import StreamingHttpResponse
    session = get_object_or_404(NodeStockTakeSession, pk=pk)

    rows = [['Mega Bin', 'Crate', 'Slot', 'Serial Number']]
    for mb in session.mega_bins.prefetch_related('crates__nodes').all():
        for crate in mb.crates.all():
            for node in crate.nodes.all():
                rows.append([mb.name, crate.name, node.slot, node.serial_number])

    filename = f"node_scan_{session.pk}_{session.date}.csv"
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    writer = _csv.writer(response)
    writer.writerows(rows)
    return response


def node_mega_bin_add(request, session_pk):
    import json as _json
    from .models import NodeStockTakeSession, NodeMegaBin, NodeCrate
    session = get_object_or_404(NodeStockTakeSession, pk=session_pk)
    if request.method == 'POST':
        try:
            data = _json.loads(request.body)
        except Exception:
            data = request.POST
        order = session.mega_bins.count()
        name = str(data.get('name', '')).strip() or f'MB-{order + 1:02d}'
        mb = NodeMegaBin.objects.create(session=session, name=name, order=order)
        crates = []
        for i in range(session.crates_per_mega_bin):
            c = NodeCrate.objects.create(mega_bin=mb, name=f'C-{i + 1:02d}', order=i)
            crates.append({'id': c.pk, 'name': c.name, 'nodes': {}})
        return JsonResponse({'ok': True, 'id': mb.pk, 'name': mb.name, 'crates': crates})
    return JsonResponse({'ok': False})


def node_mega_bin_rename(request, pk):
    import json as _json
    from .models import NodeMegaBin
    mb = get_object_or_404(NodeMegaBin, pk=pk)
    if request.method == 'POST':
        try:
            data = _json.loads(request.body)
        except Exception:
            data = request.POST
        name = str(data.get('name', '')).strip()
        if name:
            mb.name = name
            mb.save(update_fields=['name'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False})


def node_crate_rename(request, pk):
    import json as _json
    from .models import NodeCrate
    crate = get_object_or_404(NodeCrate, pk=pk)
    if request.method == 'POST':
        try:
            data = _json.loads(request.body)
        except Exception:
            data = request.POST
        name = str(data.get('name', '')).strip()
        if name:
            crate.name = name
            crate.save(update_fields=['name'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False})


def node_mega_bin_delete(request, pk):
    from .models import NodeMegaBin
    mb = get_object_or_404(NodeMegaBin, pk=pk)
    if request.method == 'POST':
        mb.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False})


def node_crate_add(request, mega_bin_pk):
    from .models import NodeMegaBin, NodeCrate
    mb = get_object_or_404(NodeMegaBin, pk=mega_bin_pk)
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        if name:
            crate = NodeCrate.objects.create(mega_bin=mb, name=name, order=mb.crates.count())
            return JsonResponse({'ok': True, 'id': crate.pk, 'name': crate.name, 'mega_bin_id': mb.pk})
    return JsonResponse({'ok': False})


def node_crate_delete(request, pk):
    from .models import NodeCrate
    crate = get_object_or_404(NodeCrate, pk=pk)
    if request.method == 'POST':
        crate.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False})


def node_scan(request, pk):
    import json as _json
    from .models import NodeCrate, NodeRecord
    crate = get_object_or_404(NodeCrate, pk=pk)
    if request.method == 'POST':
        try:
            data = _json.loads(request.body)
        except Exception:
            data = request.POST
        serial = str(data.get('serial_number', '')).strip()
        slot = int(data.get('slot', 1))

        # Clear slot
        if not serial:
            crate.nodes.filter(slot=slot).delete()
            return JsonResponse({'ok': True, 'cleared': True})

        # Duplicate serial check within this crate (different slot)
        clash = crate.nodes.filter(serial_number=serial).exclude(slot=slot).first()
        if clash:
            return JsonResponse({'ok': False, 'error': f'{serial} already in slot {clash.slot} of this crate'})

        existing = crate.nodes.filter(slot=slot).first()
        if existing:
            existing.serial_number = serial
            existing.save(update_fields=['serial_number'])
            node = existing
        else:
            node = NodeRecord.objects.create(crate=crate, serial_number=serial, slot=slot)

        return JsonResponse({'ok': True, 'node': {'id': node.pk, 'serial': node.serial_number, 'slot': node.slot}})
    return JsonResponse({'ok': False, 'error': 'POST required'})


def node_record_delete(request, pk):
    from .models import NodeRecord
    node = get_object_or_404(NodeRecord, pk=pk)
    crate_pk = node.crate_id
    if request.method == 'POST':
        node.delete()
        from .models import NodeCrate
        count = NodeCrate.objects.get(pk=crate_pk).nodes.count()
        return JsonResponse({'ok': True, 'count': count})
    return JsonResponse({'ok': False})


def dead_node_add(request, job_pk):
    from .models import DeadNode
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        serial = request.POST.get('serial_number', '').strip()
        if serial:
            DeadNode.objects.create(
                job=job,
                serial_number=serial,
                reason=request.POST.get('reason', 'other'),
                notes=request.POST.get('notes', '').strip(),
                date=request.POST.get('date') or date.today(),
                reported_by=request.POST.get('reported_by', '').strip(),
            )
    return redirect('node_stock_take', job_pk=job_pk)


def dead_node_delete(request, pk):
    from .models import DeadNode
    node = get_object_or_404(DeadNode, pk=pk)
    job_pk = node.job_id
    if request.method == 'POST':
        node.delete()
    return redirect('node_stock_take', job_pk=job_pk)


# ---------------------------------------------------------------------------
# Pre-Job Vib Tests
# ---------------------------------------------------------------------------

def pre_job_vib_tests(request, job_pk):
    from .models import PreJobVibFile
    job = get_object_or_404(Job, pk=job_pk)
    vib_files = job.pre_job_vib_files.order_by('uploaded_at')

    METRICS = ['Phase Max', 'Phase Avg', 'Force Max', 'Force Avg', 'THD Max', 'THD Avg']
    COLORS = ['#e63946', '#f4a261', '#2a9d8f', '#457b9d', '#9b2226', '#606c38']

    charts_by_unit = []  # list of {unit_id, datasets: [{label, color, data:[]}], labels:[]}

    for vf in vib_files:
        try:
            df = _read_csv(vf.file.path)
            df.columns = df.columns.str.strip()
            if 'Void' in df.columns:
                df = df[df['Void'].isna()]
            if 'Unit ID' not in df.columns:
                continue

            # Keep only rows that have at least one metric column
            present = [m for m in METRICS if m in df.columns]
            if not present:
                continue

            # Reset index for sequence numbering per file
            df = df.reset_index(drop=True)

            for unit_val in sorted(df['Unit ID'].dropna().unique()):
                unit_df = df[df['Unit ID'] == unit_val].reset_index(drop=True)
                labels = list(range(1, len(unit_df) + 1))

                datasets = []
                for metric, color in zip(METRICS, COLORS):
                    if metric in unit_df.columns:
                        vals = pd.to_numeric(unit_df[metric], errors='coerce').tolist()
                        # Replace NaN with None for JSON
                        vals = [None if (v != v) else round(float(v), 3) for v in vals]
                    else:
                        vals = [None] * len(unit_df)
                    datasets.append({'label': metric, 'color': color, 'data': vals})

                charts_by_unit.append({
                    'unit_id': str(unit_val),
                    'file_name': vf.label or vf.original_name,
                    'labels': labels,
                    'datasets': datasets,
                })
        except Exception:
            pass

    return render(request, 'reports/pre_job_vib_tests.html', {
        'job': job,
        'vib_files': vib_files,
        'charts_by_unit': charts_by_unit,
        'charts_json': json.dumps(charts_by_unit),
    })


def pre_job_vib_upload(request, job_pk):
    import zipfile as _zipfile
    import re as _re
    from django.core.files.base import ContentFile as _ContentFile
    from .models import PreJobVibFile
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        method = request.POST.get('upload_method', 'individual')
        label = request.POST.get('label', '').strip()

        if method == 'zip':
            zf_upload = request.FILES.get('zip_file')
            if zf_upload:
                _PSS_PAT = _re.compile(r'/Reports/PSS_', _re.IGNORECASE)
                try:
                    with _zipfile.ZipFile(zf_upload) as zf:
                        for entry in zf.namelist():
                            norm = '/' + entry.replace('\\', '/')
                            fname = norm.split('/')[-1]
                            if not fname:
                                continue
                            if _PSS_PAT.search(norm):
                                data = zf.read(entry)
                                PreJobVibFile.objects.create(
                                    job=job,
                                    file=_ContentFile(data, name=fname),
                                    original_name=fname,
                                    label=label,
                                )
                except Exception:
                    pass

        elif method == 'folder':
            f = request.FILES.get('pss_file')
            if f:
                PreJobVibFile.objects.create(job=job, file=f, original_name=f.name, label=label)

        else:  # individual
            for f in request.FILES.getlist('pss_file'):
                PreJobVibFile.objects.create(job=job, file=f, original_name=f.name, label=label)

    return redirect('pre_job_vib_tests', job_pk=job_pk)


def pre_job_vib_pdf(request, job_pk):
    from .models import PreJobVibFile
    from playwright.sync_api import sync_playwright
    job = get_object_or_404(Job, pk=job_pk)
    vib_files = job.pre_job_vib_files.order_by('uploaded_at')

    METRICS = ['Phase Max', 'Phase Avg', 'Force Max', 'Force Avg', 'THD Max', 'THD Avg']
    COLORS = ['#e63946', '#f4a261', '#2a9d8f', '#457b9d', '#9b2226', '#606c38']

    charts_by_unit = []
    for vf in vib_files:
        try:
            df = _read_csv(vf.file.path)
            df.columns = df.columns.str.strip()
            if 'Void' in df.columns:
                df = df[df['Void'].isna()]
            if 'Unit ID' not in df.columns:
                continue
            present = [m for m in METRICS if m in df.columns]
            if not present:
                continue
            df = df.reset_index(drop=True)
            for unit_val in sorted(df['Unit ID'].dropna().unique()):
                unit_df = df[df['Unit ID'] == unit_val].reset_index(drop=True)
                labels = list(range(1, len(unit_df) + 1))
                datasets = []
                for metric, color in zip(METRICS, COLORS):
                    if metric in unit_df.columns:
                        vals = pd.to_numeric(unit_df[metric], errors='coerce').tolist()
                        vals = [None if (v != v) else round(float(v), 3) for v in vals]
                    else:
                        vals = [None] * len(unit_df)
                    datasets.append({'label': metric, 'color': color, 'data': vals})
                charts_by_unit.append({
                    'unit_id': str(unit_val),
                    'file_name': vf.label or vf.original_name,
                    'labels': labels,
                    'datasets': datasets,
                })
        except Exception:
            pass

    html = render_to_string('reports/pre_job_vib_pdf.html', {
        'job': job,
        'charts_json': json.dumps(charts_by_unit),
        'generated_date': date.today().strftime('%d %b %Y').lstrip('0'),
    })

    with sync_playwright() as p:
        browser = p.chromium.launch()
        pg = browser.new_page()
        pg.set_content(html, wait_until='networkidle')
        # Wait until Chart.js has rendered all canvases
        pg.wait_for_function('window._chartsReady === true')
        pdf_bytes = pg.pdf(format='A4', print_background=True,
                           margin={'top': '1.5cm', 'bottom': '1.5cm', 'left': '1.5cm', 'right': '1.5cm'})
        browser.close()

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    fname = f"PreJobVibTests_{job.job_number.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'inline; filename="{fname}"'
    return response


def pre_job_vib_file_delete(request, pk):
    from .models import PreJobVibFile
    vf = get_object_or_404(PreJobVibFile, pk=pk)
    job_pk = vf.job_id
    if request.method == 'POST':
        vf.file.delete(save=False)
        vf.delete()
    return redirect('pre_job_vib_tests', job_pk=job_pk)


# ---------------------------------------------------------------------------
# Job Photos
# ---------------------------------------------------------------------------

def job_photos(request, job_pk):
    from django.db.models import Max
    job = get_object_or_404(Job, pk=job_pk)
    if request.method == 'POST':
        max_order = job.photos.aggregate(m=Max('order'))['m'] or 0
        taken_at = request.POST.get('taken_at') or None
        for i, f in enumerate(request.FILES.getlist('images')):
            JobPhoto.objects.create(job=job, image=f, taken_at=taken_at, order=max_order + i + 1)
        return redirect('job_photos', job_pk=job_pk)
    photos = job.photos.all()
    return render(request, 'reports/job_photos.html', {'job': job, 'photos': photos})


def job_photo_delete(request, pk):
    photo = get_object_or_404(JobPhoto, pk=pk)
    job_pk = photo.job_id
    if request.method == 'POST':
        photo.image.delete(save=False)
        photo.delete()
    return redirect('job_photos', job_pk=job_pk)


def job_photo_update(request, pk):
    photo = get_object_or_404(JobPhoto, pk=pk)
    if request.method == 'POST':
        photo.caption = request.POST.get('caption', '').strip()
        taken = request.POST.get('taken_at', '').strip()
        photo.taken_at = taken if taken else None
        photo.save()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def job_photo_reorder(request, job_pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        for item in data:
            JobPhoto.objects.filter(pk=item['pk'], job_id=job_pk).update(order=item['order'])
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=400)


def job_photos_import(request, job_pk):
    import shutil
    from django.core.files import File
    from django.db.models import Max
    job = get_object_or_404(Job, pk=job_pk)
    report_photos = (ReportPhoto.objects
                     .filter(report__job=job)
                     .select_related('report')
                     .order_by('report__date', 'report__report_type', 'order'))
    if request.method == 'POST':
        selected_ids = request.POST.getlist('photo_ids')
        if selected_ids:
            max_order = job.photos.aggregate(m=Max('order'))['m'] or 0
            for i, rp in enumerate(ReportPhoto.objects.filter(pk__in=selected_ids, report__job=job)):
                with rp.image.open('rb') as f:
                    filename = os.path.basename(rp.image.name)
                    jp = JobPhoto(job=job, caption=rp.caption, taken_at=rp.report.date, order=max_order + i + 1)
                    jp.image.save(filename, File(f), save=True)
        return redirect('job_photos', job_pk=job_pk)
    # Group by report for display
    from collections import defaultdict
    groups = defaultdict(list)
    for rp in report_photos:
        groups[rp.report].append(rp)
    return render(request, 'reports/job_photos_import.html', {
        'job': job,
        'groups': dict(groups),
    })
